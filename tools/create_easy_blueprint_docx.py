#!/usr/bin/env python3
"""Create the beginner-friendly MQTT research blueprint."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "BLUEPRINT_PROJECT_MQTT_MUDAH_DIPAHAMI.docx"

BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREEN = "E2F0D9"
LIGHT_ORANGE = "FCE4D6"
LIGHT_GRAY = "F2F2F2"
DARK = RGBColor(31, 31, 31)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color="FFFFFF")
        shade(table.rows[0].cells[index], BLUE)
        set_cell_margins(table.rows[0].cells[index])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for col_index, value in enumerate(values):
            set_cell_text(cells[col_index], value)
            set_cell_margins(cells[col_index])
            if row_index % 2:
                shade(cells[col_index], LIGHT_GRAY)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_steps(doc: Document, items: list[str]) -> None:
    for number, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.65)
        paragraph.paragraph_format.first_line_indent = Cm(-0.65)
        paragraph.add_run(f"{number}.").bold = True
        paragraph.add_run(f"  {item}")


def add_callout(doc: Document, title: str, body: str, fill=LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    title_run = paragraph.add_run(f"{title}\n")
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(BLUE)
    body_run = paragraph.add_run(body)
    body_run.font.size = Pt(10)
    doc.add_paragraph()


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Halaman ")
    run.font.size = Pt(8)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (
        ("Title", 24, BLUE),
        ("Subtitle", 12, "5B6573"),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 12, BLUE),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"

    doc.styles["Heading 1"].paragraph_format.space_before = Pt(8)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(7)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(5)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(4)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        add_page_number(footer)


def add_cover(doc: Document) -> None:
    doc.add_paragraph("\n\n")
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("BLUEPRINT PROJECT MQTT")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Versi Mudah Dipahami")
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    doc.add_paragraph()
    research = doc.add_paragraph()
    research.alignment = WD_ALIGN_PARAGRAPH.CENTER
    research.add_run(
        "Analisis Serangan DoS pada Broker MQTT dan Mitigasinya\n"
        "Menggunakan Rate Limiting"
    ).font.size = Pt(13)

    doc.add_paragraph("\n")
    add_callout(
        doc,
        "Tujuan dokumen",
        "Menjelaskan apa yang dikerjakan oleh Windows, Ubuntu Server A, dan Ubuntu "
        "Server B; bagaimana SYN flood diuji; serta bagaimana hasil penelitian dibaca.",
    )

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(
        "Arsitektur utama: Windows client + Server A broker + Server B attacker"
    ).bold = True
    doc.add_page_break()


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    doc.add_heading("1. Gambaran Paling Sederhana", level=1)
    doc.add_paragraph(
        "Project ini menguji apakah banyak permintaan koneksi palsu dapat mengganggu "
        "layanan MQTT, kemudian menguji apakah rate limiting dapat mengurangi gangguan tersebut."
    )
    add_table(
        doc,
        ["Bagian", "Peran", "Bukan Perannya"],
        [
            [
                "Windows",
                "Mengirim data suhu simulasi dan mengukur layanan MQTT.",
                "Bukan server broker dan bukan pengirim serangan.",
            ],
            [
                "Ubuntu Server A",
                "Menjalankan broker MQTT, merekam trafik, dan menerapkan rate limiting.",
                "Bukan perangkat IoT dan bukan mesin attacker.",
            ],
            [
                "Ubuntu Server B",
                "Mengirim SYN flood terkontrol menuju Server A.",
                "Tidak mengirim data suhu MQTT.",
            ],
        ],
        widths=[3.2, 6.5, 6.5],
    )
    add_callout(
        doc,
        "Inti penelitian",
        "Windows mewakili pengguna normal. Server B mewakili sumber serangan. "
        "Server A adalah objek yang menerima keduanya dan menjadi tempat pengamatan.",
        LIGHT_GREEN,
    )

    doc.add_heading("2. Masalah pada Konsep Lama", level=1)
    doc.add_paragraph(
        "Pada rancangan lama, packet capture dilakukan pada host monitoring yang terpisah. "
        "Masalahnya, host tersebut belum tentu melihat seluruh trafik yang benar-benar masuk "
        "ke broker. Data pengamatan dapat berbeda dari kondisi yang dialami broker."
    )
    doc.add_heading("3. Perubahan pada Konsep Baru", level=1)
    add_bullets(
        doc,
        [
            "Packet capture dipindahkan langsung ke Ubuntu Server A.",
            "Objek yang diamati adalah broker MQTT yang menerima trafik.",
            "Jenis DoS dipersempit menjadi SYN flood.",
            "Mitigasi difokuskan pada rate limiting paket SYN.",
            "Istilah sniffer dan prober terpisah tidak lagi menjadi inti arsitektur.",
        ],
    )
    add_callout(
        doc,
        "Yang tidak berubah",
        "Topik utama tetap analisis DoS terhadap broker MQTT dan pengujian efektivitas "
        "rate limiting. Perubahan hanya memperjelas jenis serangan dan posisi pengambilan data.",
    )
    doc.add_page_break()

    doc.add_heading("4. Arsitektur Tiga Sistem", level=1)
    add_table(
        doc,
        ["WINDOWS", "UBUNTU SERVER A", "UBUNTU SERVER B"],
        [
            ["Client MQTT", "Broker Mosquitto", "Attacker pengujian"],
            ["Mengirim data suhu", "Menerima trafik", "Mengirim paket SYN"],
            ["Mengukur latency", "Menjalankan TShark", "Menjalankan hping3"],
            ["Menyimpan hasil client", "Menjalankan nftables", "Menyimpan attack.log"],
            ["Menjalankan analisis", "Menyimpan PCAP/raw flow", "Target hanya Server A"],
        ],
        widths=[5.4, 5.4, 5.4],
    )

    doc.add_heading("5. Jaringan yang Digunakan", level=1)
    doc.add_paragraph(
        "Mode awal menggunakan VirtualBox Host-Only. Jaringan ini dibuat di dalam laptop "
        "agar Windows dan kedua VM dapat berkomunikasi tanpa mengirim SYN flood ke Wi-Fi."
    )
    add_table(
        doc,
        ["Perangkat", "Contoh IP", "Kegunaan"],
        [
            ["Windows", "192.168.56.1", "Client dan analisis"],
            ["Server A", "192.168.56.10", "Alamat broker MQTT"],
            ["Server B", "192.168.56.20", "Mesin pengujian serangan"],
        ],
        widths=[5.0, 4.0, 7.2],
    )
    add_callout(
        doc,
        "Peran Wi-Fi",
        "Wi-Fi hanya dipakai untuk internet, misalnya mengunduh package. Trafik penelitian "
        "melewati Host-Only Adapter. Ketiga sistem tidak perlu berada sebagai tiga perangkat "
        "fisik pada Wi-Fi yang sama.",
        LIGHT_ORANGE,
    )

    doc.add_heading("6. Mengapa Server A Menjadi Pusat", level=1)
    add_bullets(
        doc,
        [
            "Broker MQTT berjalan di Server A.",
            "Trafik normal dari Windows masuk ke Server A.",
            "Paket SYN dari Server B juga masuk ke Server A.",
            "Rate limiting bekerja pada jalur masuk Server A.",
            "Capture di Server A menunjukkan paket yang benar-benar mencapai broker.",
        ],
    )
    doc.add_page_break()

    doc.add_heading("7. Mekanisme Trafik Normal", level=1)
    add_steps(
        doc,
        [
            "Program Windows membuat nilai suhu simulasi.",
            "Windows melakukan publish ke topik unram/iot/suhu.",
            "Pesan dikirim melalui koneksi TCP ke port MQTT 1883.",
            "Mosquitto pada Server A menerima dan meneruskan pesan kepada subscriber.",
            "Windows menerima kembali pesan lalu menghitung waktu pengiriman.",
        ],
    )
    add_callout(
        doc,
        "Apa yang dikirim Windows?",
        "Windows mengirim pesan MQTT berisi data suhu simulasi. Data suhu bukan serangan. "
        "Data ini menjadi beban normal untuk memeriksa apakah layanan broker tetap bekerja.",
        LIGHT_GREEN,
    )

    doc.add_heading("8. Mekanisme SYN Flood", level=1)
    doc.add_paragraph(
        "Sebelum MQTT dapat digunakan, client harus membentuk koneksi TCP. Paket SYN adalah "
        "permintaan awal untuk membuka koneksi TCP. Server B mengirim banyak paket SYN menuju "
        "IP Server A dan port 1883."
    )
    add_steps(
        doc,
        [
            "Server B menentukan target 192.168.56.10:1883.",
            "Hping3 mengirim paket SYN berulang kali.",
            "Server A menerima banyak permintaan awal koneksi.",
            "TShark merekam kenaikan jumlah paket SYN.",
            "Kinerja pesan normal Windows dibandingkan dengan kondisi tanpa serangan.",
        ],
    )
    add_table(
        doc,
        ["Istilah", "Lapisan", "Arti dalam project"],
        [
            ["MQTT", "Application layer", "Protokol pesan yang dipakai client."],
            ["TCP", "Transport layer", "Jalur koneksi yang digunakan MQTT."],
            ["SYN", "Bagian dari TCP", "Permintaan awal pembentukan koneksi."],
            ["SYN flood", "Serangan pada TCP", "Banyak SYN dikirim untuk membebani target."],
        ],
        widths=[3.2, 4.0, 9.0],
    )
    doc.add_page_break()

    doc.add_heading("9. Cara Kerja Rate Limiting", level=1)
    doc.add_paragraph(
        "Nftables pada Server A menghitung laju paket SYN yang masuk. Paket masih diterima "
        "selama berada dalam batas. Paket berlebih dibuang agar broker tidak menangani seluruh "
        "permintaan koneksi."
    )
    add_table(
        doc,
        ["Nilai", "Contoh", "Arti"],
        [
            ["RL_RATE", "50/second", "Laju SYN yang diizinkan sebelum dianggap berlebih."],
            ["RL_BURST", "100 paket", "Toleransi lonjakan singkat sebelum paket dibuang."],
            ["ATTACK_RATE", "1000 paket/detik", "Laju SYN yang dibuat Server B."],
        ],
        widths=[3.5, 4.0, 8.7],
    )
    add_callout(
        doc,
        "Mengapa masih perlu capture?",
        "Capture membuktikan bahwa trafik SYN memang datang. Counter nftables membuktikan "
        "bahwa paket berlebih dibuang. KPI MQTT menunjukkan apakah layanan menjadi lebih baik.",
    )

    doc.add_heading("10. Tiga Skenario Pengujian", level=1)
    add_table(
        doc,
        ["Skenario", "Client Windows", "SYN Flood", "Rate Limiting", "Tujuan"],
        [
            ["normal", "Aktif", "Tidak aktif", "Tidak aktif", "Mendapatkan kondisi dasar."],
            ["syn_flood", "Aktif", "Aktif", "Tidak aktif", "Mengukur dampak serangan."],
            [
                "syn_flood_\nrate_limit",
                "Aktif",
                "Aktif",
                "Aktif",
                "Mengukur efek mitigasi.",
            ],
        ],
        widths=[3.5, 2.8, 2.8, 3.0, 4.2],
    )
    doc.add_paragraph(
        "Client harus tetap aktif pada ketiga skenario. Jika client tidak berjalan saat serangan, "
        "penelitian hanya mengetahui jumlah paket SYN dan tidak mengetahui dampaknya terhadap MQTT."
    )
    doc.add_page_break()

    doc.add_heading("11. Data yang Dihasilkan", level=1)
    add_table(
        doc,
        ["File", "Dibuat oleh", "Isi"],
        [
            ["mqtt_client.csv", "Windows", "Status pesan dan latency MQTT."],
            ["capture.pcapng", "Server A", "Rekaman paket asli."],
            ["raw_flow.csv", "Server A", "Paket penting dalam bentuk tabel."],
            ["metadata.env", "Server A", "Konfigurasi satu run."],
            ["nft_after_capture.txt", "Server A", "Counter rate limiting."],
            ["attack.log", "Server B", "Catatan proses SYN flood."],
            ["metrics.json / metrics.csv", "Windows", "KPI hasil satu run."],
        ],
        widths=[4.4, 4.0, 7.8],
    )
    doc.add_paragraph(
        "Untuk analisis, mqtt_client.csv dari Windows digabungkan dengan metadata.env dan "
        "raw_flow.csv dari Server A dalam folder RUN_ID yang sama."
    )

    doc.add_heading("12. KPI dan Cara Membacanya", level=1)
    add_table(
        doc,
        ["KPI", "Makna", "Arah yang baik"],
        [
            ["mqtt_success_rate", "Persentase pesan MQTT berhasil.", "Lebih tinggi."],
            ["mqtt_latency_median_ms", "Waktu tengah pengiriman pesan.", "Lebih rendah."],
            ["mqtt_latency_p95_ms", "Latency pada bagian pesan yang lambat.", "Lebih rendah."],
            ["syn_rate_mean", "Rata-rata SYN per detik.", "Naik saat serangan."],
            ["syn_rate_peak", "Puncak SYN dalam satu detik.", "Membuktikan intensitas."],
        ],
        widths=[5.0, 7.0, 4.2],
    )
    add_callout(
        doc,
        "Cara menyimpulkan",
        "Pertama pastikan SYN rate meningkat. Kedua lihat apakah MQTT terganggu. Ketiga "
        "bandingkan serangan tanpa mitigasi dengan serangan yang memakai rate limiting.",
        LIGHT_GREEN,
    )
    doc.add_page_break()

    doc.add_heading("13. Pembagian Direktori", level=1)
    add_table(
        doc,
        ["Direktori", "Diberikan kepada", "Program utama"],
        [
            [
                "01_windows_client",
                "Windows",
                "run_client.ps1, analyze_run.ps1, analyze_all.ps1",
            ],
            [
                "02_ubuntu_server_a_broker",
                "Ubuntu Server A",
                "run_broker.sh, finalize_broker.sh, rate_limit.sh",
            ],
            [
                "03_ubuntu_server_b_attacker",
                "Ubuntu Server B",
                "run_attacker.sh",
            ],
        ],
        widths=[5.5, 4.2, 6.5],
    )

    doc.add_heading("14. Alur Kerja Satu Pengujian", level=1)
    add_steps(
        doc,
        [
            "Tentukan RUN_ID dan SCENARIO.",
            "Samakan kedua nilai tersebut pada ketiga sistem.",
            "Mulai capture dan perlakuan rate limiting pada Server A.",
            "Jalankan client MQTT pada Windows.",
            "Jalankan Server B; pada skenario normal program tidak mengirim serangan.",
            "Finalisasi capture pada Server A menjadi raw_flow.csv.",
            "Pindahkan metadata.env dan raw_flow.csv ke folder run Windows.",
            "Jalankan analisis satu run.",
            "Setelah semua run tersedia, jalankan analisis perbandingan.",
        ],
    )
    add_callout(
        doc,
        "Mengapa capture 75 detik?",
        "Client dan attacker berjalan 60 detik. Capture dibuat 75 detik karena Server A "
        "dijalankan lebih dahulu dan pengguna memerlukan waktu berpindah ke Windows dan Server B.",
    )

    doc.add_heading("15. Pengulangan", level=1)
    doc.add_paragraph(
        "Satu run cukup untuk demonstrasi, tetapi tidak cukup untuk kesimpulan penelitian. "
        "Setiap skenario perlu diulang dengan konfigurasi yang sama, misalnya lima kali."
    )
    add_bullets(
        doc,
        [
            "normal_01 sampai normal_05",
            "attack_01 sampai attack_05",
            "mitigation_01 sampai mitigation_05",
        ],
    )
    doc.add_page_break()

    doc.add_heading("16. Mode Local, Campus, dan Public", level=1)
    add_table(
        doc,
        ["Mode", "Lokasi target", "Kegunaan", "Syarat"],
        [
            ["local", "VM Server A", "Setup, demo, dan validasi metode.", "Testbed sendiri."],
            [
                "campus",
                "Jaringan internal/VPN",
                "Implementasi pada lingkungan kampus.",
                "Izin pengelola.",
            ],
            [
                "public",
                "Broker yang dapat diakses internet",
                "Implementasi melalui alamat publik.",
                "Izin tertulis dan kontrol server.",
            ],
        ],
        widths=[2.5, 4.4, 5.0, 4.3],
    )
    add_callout(
        doc,
        "Batas klaim penelitian",
        "Mode local membuktikan bahwa metode dan program bekerja pada testbed. Jika skripsi "
        "menyatakan pengujian dilakukan pada Server Universitas Mataram, data final harus "
        "diambil pada server atau jaringan Universitas Mataram yang benar-benar diberi izin.",
        LIGHT_ORANGE,
    )

    doc.add_heading("17. BROKER_HOST dan BROKER_CAPTURE_HOST", level=1)
    add_table(
        doc,
        ["Variabel", "Digunakan oleh", "Arti"],
        [
            [
                "BROKER_HOST",
                "Windows dan Server B",
                "Alamat yang dipakai untuk mencapai broker.",
            ],
            [
                "BROKER_CAPTURE_HOST",
                "Server A",
                "Alamat broker yang terlihat pada interface capture.",
            ],
        ],
        widths=[4.8, 5.0, 6.4],
    )
    doc.add_paragraph(
        "Pada mode local keduanya biasanya sama, misalnya 192.168.56.10. Pada jaringan "
        "dengan NAT, BROKER_HOST dapat berupa IP publik sedangkan BROKER_CAPTURE_HOST berupa IP privat."
    )

    doc.add_heading("18. Batas Teknis Saat Ini", level=1)
    add_bullets(
        doc,
        [
            "Tutorial lokal menggunakan MQTT tanpa TLS pada port 1883.",
            "Port 8883 memerlukan dukungan sertifikat TLS yang belum menjadi bagian client saat ini.",
            "allow_anonymous=true hanya layak pada Host-Only testbed, bukan server nyata.",
            "SYN flood tidak selalu membuat broker mati; dampaknya bergantung pada kapasitas server.",
            "Keberhasilan mitigasi harus dilihat dari KPI MQTT dan counter drop, bukan asumsi.",
        ],
    )
    doc.add_page_break()

    doc.add_heading("19. Ringkasan yang Harus Dipahami", level=1)
    add_table(
        doc,
        ["Pertanyaan", "Jawaban"],
        [
            ["Siapa pengguna normal?", "Program client MQTT pada Windows."],
            ["Apa data normalnya?", "Pesan MQTT berisi suhu simulasi."],
            ["Di mana broker berjalan?", "Ubuntu Server A."],
            ["Siapa yang mengirim serangan?", "Ubuntu Server B."],
            ["Apa bentuk serangannya?", "Banyak paket SYN menuju TCP port 1883."],
            ["Di mana trafik direkam?", "Langsung pada interface Server A."],
            ["Apa mitigasinya?", "Rate limiting nftables pada Server A."],
            ["Apa yang dibandingkan?", "Normal, SYN flood, dan SYN flood dengan rate limiting."],
            ["Di mana analisis dilakukan?", "Windows setelah file hasil digabungkan."],
        ],
        widths=[7.0, 9.2],
    )

    add_callout(
        doc,
        "Kesimpulan blueprint",
        "Windows menguji apakah MQTT masih dapat digunakan. Server B menciptakan kondisi SYN "
        "flood. Server A menerima, merekam, dan membatasi trafik. Hasil ketiga skenario "
        "dibandingkan untuk menilai dampak serangan dan efektivitas rate limiting.",
        LIGHT_GREEN,
    )
    doc.add_heading("20. Dokumen Pendamping", level=1)
    doc.add_paragraph(
        "Setelah memahami blueprint ini, gunakan file "
        "TUTORIAL_SETUP_DAN_PENGUJIAN_END_TO_END.txt untuk langkah instalasi dan perintah."
    )

    return doc


def main() -> int:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
