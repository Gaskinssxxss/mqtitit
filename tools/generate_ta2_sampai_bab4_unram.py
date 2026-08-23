#!/usr/bin/env python3
"""Generate draft TA 2 sampai BAB IV untuk bimbingan.

Draft ini memakai:
- substansi proposal TA 1 MQTT DoS rate limiting;
- format umum TA 2 dari contoh dokumen Indra;
- data awal pengujian UNRAM yang sudah tersedia di output/unram_experiments.
"""

from __future__ import annotations

import csv
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "doc"
TMP_DIR = ROOT / "tmp" / "docs" / "ta2_bab4"
FIG_DIR = TMP_DIR / "figures"
OUT_DOCX = OUT_DIR / "DRAFT_TA2_SAMPAI_BAB4_MQTT_DOS_RATE_LIMITING_UNRAM.docx"

TITLE = (
    "ANALISIS SERANGAN DENIAL OF SERVICE PADA BROKER MQTT DAN "
    "MITIGASINYA MENGGUNAKAN RATE LIMITING PADA SERVER UNIVERSITAS MATARAM"
)
AUTHOR = "Kalsum Rahmawati"
NIM = "F1D022126"
LOGO = ROOT / "tmp" / "docs" / "proposal_assets" / "word" / "media" / "image1.jpg"
COMPARISON = ROOT / "output" / "unram_experiments" / "_report_bab4_bab5" / "comparison_baseline_vs_synflood.csv"


def font(size: int = 28, bold: bool = False):
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size)
    return ImageFont.load_default()


def read_comparison() -> list[dict[str, str]]:
    with COMPARISON.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def metric(rows: list[dict[str, str]], name: str, col: str) -> str:
    for row in rows:
        if row["indikator"] == name:
            return row[col]
    return "-"


def make_bar_chart(rows: list[dict[str, str]], filename: str, title: str, metrics: list[str]) -> Path:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1500, 900), "#ffffff")
    d = ImageDraw.Draw(img)
    title_font = font(42, True)
    label_font = font(25, False)
    small_font = font(22, False)
    d.text((90, 50), title, fill="#111827", font=title_font)

    values = []
    for m in metrics:
        b = float(metric(rows, m, "baseline").replace("%", ""))
        s = float(metric(rows, m, "synflood_low").replace("%", ""))
        values.append((m, b, s))
    max_v = max(max(b, s) for _, b, s in values) or 1

    left, top = 170, 160
    chart_w, chart_h = 1120, 560
    d.line((left, top, left, top + chart_h), fill="#111827", width=4)
    d.line((left, top + chart_h, left + chart_w, top + chart_h), fill="#111827", width=4)

    group_w = chart_w / len(values)
    for i, (name, b, s) in enumerate(values):
        x0 = left + i * group_w + 50
        bar_w = 70
        b_h = int((b / max_v) * (chart_h - 50))
        s_h = int((s / max_v) * (chart_h - 50))
        base_y = top + chart_h
        d.rectangle((x0, base_y - b_h, x0 + bar_w, base_y), fill="#2563eb")
        d.rectangle((x0 + 95, base_y - s_h, x0 + 95 + bar_w, base_y), fill="#dc2626")
        d.text((x0 - 20, base_y + 15), name[:24], fill="#111827", font=small_font)
        d.text((x0 - 5, base_y - b_h - 35), f"{b:g}", fill="#111827", font=small_font)
        d.text((x0 + 90, base_y - s_h - 35), f"{s:g}", fill="#111827", font=small_font)

    d.rectangle((1030, 95, 1070, 125), fill="#2563eb")
    d.text((1085, 94), "Baseline", fill="#111827", font=label_font)
    d.rectangle((1030, 135, 1070, 165), fill="#dc2626")
    d.text((1085, 134), "SYN flood terkendali", fill="#111827", font=label_font)
    out = FIG_DIR / filename
    img.save(out)
    return out


def make_architecture() -> Path:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    out = FIG_DIR / "arsitektur_pengujian_mqtt.png"
    img = Image.new("RGB", (1500, 850), "#f8fafc")
    d = ImageDraw.Draw(img)
    d.text((360, 40), "Arsitektur Pengujian Broker MQTT Universitas Mataram", fill="#111827", font=font(38, True))

    def box(xy, text, fill, outline):
        d.rounded_rectangle(xy, radius=20, fill=fill, outline=outline, width=4)
        lines = text.split("\n")
        y = xy[1] + 35
        for line in lines:
            bbox = d.textbbox((0, 0), line, font=font(25, True if line == lines[0] else False))
            d.text((xy[0] + ((xy[2] - xy[0]) - (bbox[2] - bbox[0])) / 2, y), line, fill="#111827", font=font(25, True if line == lines[0] else False))
            y += 42

    def arr(start, end, label):
        d.line((start, end), fill="#111827", width=4)
        ex, ey = end
        d.polygon([(ex, ey), (ex - 18, ey - 10), (ex - 18, ey + 10)], fill="#111827")
        d.text(((start[0] + end[0]) / 2 - 90, (start[1] + end[1]) / 2 - 40), label, fill="#111827", font=font(21, False))

    box((70, 190, 410, 430), "Laptop Client / PC A\nVPN WireGuard\nProber MQTT\nPacket capture", "#dbeafe", "#1d4ed8")
    box((580, 190, 930, 430), "Server UNRAM\nBroker MQTT\nMosquitto\nPort 1883", "#dcfce7", "#15803d")
    box((1080, 190, 1430, 430), "Attacker Terkontrol\nhping3\nSYN ke port 1883\nRate rendah", "#fee2e2", "#b91c1c")
    arr((410, 285), (580, 285), "publish MQTT")
    d.line((1080, 350, 930, 350), fill="#111827", width=4)
    d.polygon([(930, 350), (948, 340), (948, 360)], fill="#111827")
    d.text((960, 305), "SYN flood", fill="#111827", font=font(21, False))
    d.text((100, 575), "Output data: prober.csv, capture.pcapng, capture_raw_flow.csv, metrics.csv, summary.txt", fill="#111827", font=font(28, True))
    d.text((100, 635), "Catatan: data awal yang dipakai pada BAB IV berasal dari pengujian melalui VPN WireGuard dan masih perlu dilengkapi skenario rate limiting.", fill="#92400e", font=font(24, True))
    img.save(out)
    return out


def make_flow() -> Path:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    out = FIG_DIR / "alur_data_pengujian.png"
    img = Image.new("RGB", (1500, 850), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((520, 45), "Alur Data Pengujian", fill="#111827", font=font(42, True))

    items = [
        ((70, 230, 330, 390), "Capture\npcapng"),
        ((420, 230, 680, 390), "Ekstraksi\nraw_flow.csv"),
        ((770, 230, 1030, 390), "Kompilasi\nmetrics.csv"),
        ((1120, 230, 1430, 390), "Pembahasan\nBAB IV"),
        ((420, 540, 680, 700), "Prober\nprober.csv"),
    ]
    for xy, text in items:
        d.rounded_rectangle(xy, radius=18, fill="#eef2ff", outline="#4338ca", width=4)
        y = xy[1] + 35
        for line in text.split("\n"):
            bbox = d.textbbox((0, 0), line, font=font(28, True))
            d.text((xy[0] + ((xy[2] - xy[0]) - (bbox[2] - bbox[0])) / 2, y), line, fill="#111827", font=font(28, True))
            y += 45
    for start, end in [((330, 310), (420, 310)), ((680, 310), (770, 310)), ((1030, 310), (1120, 310)), ((680, 620), (1120, 620))]:
        d.line((start, end), fill="#111827", width=4)
        ex, ey = end
        d.polygon([(ex, ey), (ex - 18, ey - 10), (ex - 18, ey + 10)], fill="#111827")
    d.line((1275, 540, 1275, 390), fill="#111827", width=4)
    d.polygon([(1275, 390), (1265, 408), (1285, 408)], fill="#111827")
    img.save(out)
    return out


def doc_style(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3)
    sec.bottom_margin = Cm(3)
    sec.left_margin = Cm(4)
    sec.right_margin = Cm(3)
    for st_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        st = doc.styles[st_name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(12)
        st.font.color.rgb = RGBColor(0, 0, 0)
        if st_name.startswith("Heading"):
            st.font.bold = True
    normal = doc.styles["Normal"]
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def p(doc, text="", align=None, bold=False, italic=False, indent=True):
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.space_after = Pt(0)
    if indent:
        par.paragraph_format.first_line_indent = Cm(1.25)
    par.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return par


def center(doc, text="", bold=False):
    return p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=bold, indent=False)


def h1(doc, bab, title):
    doc.add_page_break()
    center(doc, bab, bold=True)
    center(doc, title, bold=True)
    p(doc, "", indent=False)


def h2(doc, number, title):
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.space_before = Pt(6)
    run = par.add_run(f"{number} {title}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def h3(doc, number, title):
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 1.5
    run = par.add_run(f"{number} {title}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def numlist(doc, items):
    for i, item in enumerate(items, 1):
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.75)
        par.paragraph_format.first_line_indent = Cm(-0.35)
        par.paragraph_format.line_spacing = 1.5
        run = par.add_run(f"{i}. {item}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def set_cell(cell, text, bold=False):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    par = cell.paragraphs[0]
    par.paragraph_format.line_spacing = 1.15
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = par.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def shade(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def table(doc, caption, headers, rows):
    center(doc, caption, bold=True)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(headers):
        set_cell(t.rows[0].cells[i], head, True)
        shade(t.rows[0].cells[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    p(doc, "", indent=False)


def fig(doc, path: Path, caption: str, width: float = 5.8):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(str(path), width=Inches(width))
    center(doc, caption, bold=True)


def cover(doc):
    center(doc, TITLE, bold=True)
    p(doc, "", indent=False)
    center(doc, "Tugas Akhir", bold=False)
    center(doc, "Untuk memenuhi sebagian persyaratan", bold=False)
    center(doc, "Mencapai derajat Sarjana S-1 Program Studi Teknik Informatika", bold=False)
    p(doc, "", indent=False)
    if LOGO.exists():
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(str(LOGO), width=Inches(2.35))
    else:
        for _ in range(4):
            p(doc, "", indent=False)
    p(doc, "", indent=False)
    center(doc, "Oleh:", bold=True)
    center(doc, AUTHOR, bold=True)
    center(doc, NIM, bold=True)
    for _ in range(4):
        p(doc, "", indent=False)
    center(doc, "PROGRAM STUDI TEKNIK INFORMATIKA", bold=True)
    center(doc, "FAKULTAS TEKNIK", bold=True)
    center(doc, "UNIVERSITAS MATARAM", bold=True)
    center(doc, "2026", bold=True)


def front_matter(doc):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3)
    sec.bottom_margin = Cm(3)
    sec.left_margin = Cm(4)
    sec.right_margin = Cm(3)
    sec.footer.is_linked_to_previous = False
    add_page_number(sec)

    center(doc, "KATA PENGANTAR", bold=True)
    p(doc, "Puji syukur penulis panjatkan ke hadirat Allah SWT karena atas rahmat dan karunia-Nya draft tugas akhir ini dapat disusun. Draft ini membahas analisis serangan Denial of Service pada broker MQTT dan mitigasinya menggunakan rate limiting pada server Universitas Mataram.")
    p(doc, "Dokumen ini disusun sebagai bahan bimbingan awal TA 2. Isi yang disajikan memuat struktur skripsi, dasar teori, metodologi, serta pembahasan hasil pengujian awal sampai BAB IV. Data BAB IV yang digunakan masih bersifat sementara karena pengujian mitigasi rate limiting belum lengkap.")
    p(doc, "Penulis menyadari bahwa naskah ini masih memerlukan penyempurnaan berdasarkan arahan dosen pembimbing, hasil pengujian lanjutan, dan format resmi fakultas.")
    p(doc, "Mataram, Agustus 2026", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
    p(doc, AUTHOR, align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)

    doc.add_page_break()
    center(doc, "ABSTRAK", bold=True)
    p(doc, "Penelitian ini membahas analisis serangan Denial of Service pada broker Message Queuing Telemetry Transport dan mitigasinya menggunakan rate limiting pada server Universitas Mataram. MQTT banyak digunakan pada komunikasi Internet of Things karena ringan dan mendukung pola publish-subscribe. Dalam arsitektur MQTT, broker menjadi pusat komunikasi antara client publisher dan subscriber. Posisi broker yang sentral membuat aspek ketersediaan layanan perlu diuji, terutama terhadap serangan yang menargetkan proses koneksi TCP.")
    p(doc, "Penelitian dirancang menggunakan konsep sniffer dan prober. Sniffer digunakan untuk menangkap lalu lintas jaringan yang berhubungan dengan broker MQTT, sedangkan prober digunakan sebagai client normal yang mengirim pesan MQTT secara periodik ke broker. Serangan yang diuji difokuskan pada SYN flood menuju port MQTT 1883. Data pengujian berupa capture.pcapng, capture_raw_flow.csv, metrics.csv, prober.csv, dan summary.txt. Data tersebut dianalisis untuk melihat perubahan success rate, round trip time, laju paket TCP, laju paket SYN, estimasi half-open connection, dan rasio SYN/SYN-ACK.")
    p(doc, "Hasil pengujian awal melalui VPN WireGuard menunjukkan bahwa broker MQTT Universitas Mataram dapat diakses pada port 1883. Pada skenario baseline, prober mencatat success rate 100 persen. Pada skenario SYN flood terkendali, success rate prober juga tetap 100 persen, tetapi indikator trafik serangan meningkat. Rata-rata syn_rate_target naik dari 0,8197 menjadi 4,3729 paket per detik, sedangkan rata-rata half_open_conn_target naik dari 0,2258 menjadi 3,6333. Hasil sementara ini menunjukkan bahwa pola SYN flood berhasil terlihat pada data capture, tetapi intensitas serangan rendah belum menyebabkan kegagalan publish MQTT. Pengujian rate limiting perlu dilengkapi untuk menghasilkan kesimpulan akhir mengenai efektivitas mitigasi.")
    p(doc, "Kata kunci: MQTT, broker, Denial of Service, SYN flood, packet capture, rate limiting.", indent=False)

    doc.add_page_break()
    center(doc, "DAFTAR ISI", bold=True)
    for item in [
        "KATA PENGANTAR", "ABSTRAK", "DAFTAR ISI", "DAFTAR GAMBAR", "DAFTAR TABEL",
        "BAB I PENDAHULUAN", "1.1 Latar Belakang", "1.2 Rumusan Masalah", "1.3 Batasan Masalah",
        "1.4 Tujuan Penelitian", "1.5 Manfaat Penelitian", "1.6 Sistematika Penulisan",
        "BAB II TINJAUAN PUSTAKA", "2.1 Penelitian Terkait", "2.2 Teori Penunjang",
        "BAB III METODOLOGI PENELITIAN", "BAB IV HASIL DAN PEMBAHASAN", "DAFTAR PUSTAKA",
    ]:
        p(doc, item, indent=False)

    doc.add_page_break()
    center(doc, "DAFTAR GAMBAR", bold=True)
    for item in [
        "Gambar 3.1 Arsitektur Pengujian Broker MQTT Universitas Mataram",
        "Gambar 3.2 Alur Data Pengujian",
        "Gambar 4.1 Perbandingan Indikator Trafik Baseline dan SYN Flood",
        "Gambar 4.2 Perbandingan Performa Prober Baseline dan SYN Flood",
    ]:
        p(doc, item, indent=False)

    doc.add_page_break()
    center(doc, "DAFTAR TABEL", bold=True)
    for item in [
        "Tabel 2.1 Penelitian Terkait",
        "Tabel 3.1 Alat dan Bahan",
        "Tabel 3.2 File Project dan Fungsinya",
        "Tabel 3.3 Skenario Pengujian",
        "Tabel 3.4 Metrik Pengujian",
        "Tabel 4.1 Lingkungan Pengujian Awal",
        "Tabel 4.2 Rekap Perbandingan Baseline dan SYN Flood",
    ]:
        p(doc, item, indent=False)


def chapter1(doc):
    h1(doc, "BAB I", "PENDAHULUAN")
    h2(doc, "1.1", "Latar Belakang")
    for text in [
        "Perkembangan sistem digital berbasis jaringan telah menjadi bagian penting dalam lingkungan pendidikan tinggi. Universitas Mataram sebagai institusi pendidikan memiliki kebutuhan terhadap layanan teknologi informasi yang stabil, termasuk layanan yang mendukung komunikasi perangkat Internet of Things. Dalam sistem IoT, perangkat dapat mengirimkan data sensor, status perangkat, atau pesan kontrol secara berkala menuju server.",
        "Salah satu protokol komunikasi yang banyak digunakan pada sistem IoT adalah Message Queuing Telemetry Transport atau MQTT. MQTT dirancang sebagai protokol ringan yang bekerja dengan model publish-subscribe. Pada model ini, client yang mengirim data disebut publisher, client yang menerima data disebut subscriber, sedangkan broker MQTT berperan sebagai pusat pertukaran pesan.",
        "Broker MQTT memiliki peran penting karena seluruh komunikasi antar client melewati broker. Apabila broker berjalan normal, pesan dari publisher dapat diterima dan diteruskan kepada subscriber sesuai topic. Sebaliknya, apabila broker mengalami gangguan, proses komunikasi dapat melambat, gagal, atau tidak dapat digunakan oleh client normal. Oleh karena itu, ketersediaan broker menjadi aspek penting dalam layanan MQTT.",
        "Salah satu ancaman terhadap ketersediaan layanan adalah Denial of Service. DoS bertujuan membuat layanan sulit diakses atau tidak tersedia bagi pengguna sah. Pada layanan berbasis TCP, salah satu bentuk DoS yang umum dianalisis adalah SYN flood. Serangan ini mengirim banyak paket SYN ke server sehingga server harus menangani banyak permintaan koneksi awal.",
        "MQTT tanpa TLS umumnya berjalan pada port 1883 dan menggunakan TCP sebagai protokol transport. Karena itu, broker MQTT dapat dianalisis terhadap serangan SYN flood yang diarahkan ke port tersebut. Dampak serangan dapat diamati dari dua sisi, yaitu perubahan trafik jaringan dan perubahan performa client normal yang tetap mencoba mengirim pesan MQTT.",
        "Penelitian ini menggunakan konsep sniffer dan prober. Sniffer menangkap trafik jaringan yang berkaitan dengan broker MQTT, sedangkan prober bertindak sebagai client normal yang mengirim pesan MQTT secara berkala. Dengan dua komponen ini, penelitian tidak hanya melihat bahwa trafik serangan meningkat, tetapi juga melihat apakah layanan broker masih dapat digunakan oleh client normal.",
        "Mitigasi yang dibahas adalah rate limiting. Rate limiting merupakan pembatasan laju trafik atau koneksi dalam periode tertentu. Dalam konteks broker MQTT, rate limiting diharapkan dapat mengurangi dampak trafik SYN yang berlebihan agar broker tetap lebih stabil. Penelitian ini membandingkan kondisi normal, kondisi SYN flood, dan kondisi SYN flood setelah mitigasi rate limiting.",
        "Berdasarkan uraian tersebut, penelitian ini diarahkan untuk menganalisis serangan DoS pada broker MQTT dan mitigasinya menggunakan rate limiting pada server Universitas Mataram. Hasil penelitian diharapkan dapat memberikan gambaran mengenai karakteristik trafik serangan dan dasar evaluasi terhadap mekanisme mitigasi yang diterapkan.",
    ]:
        p(doc, text)

    h2(doc, "1.2", "Rumusan Masalah")
    numlist(doc, [
        "Bagaimana karakteristik trafik broker MQTT pada kondisi normal?",
        "Bagaimana perubahan trafik ketika broker MQTT menerima serangan SYN flood?",
        "Bagaimana dampak SYN flood terhadap keberhasilan publish dan RTT client MQTT normal?",
        "Bagaimana efektivitas rate limiting dalam mengurangi dampak serangan DoS pada broker MQTT?",
    ])

    h2(doc, "1.3", "Batasan Masalah")
    numlist(doc, [
        "Objek penelitian adalah broker MQTT pada server Universitas Mataram atau lingkungan uji yang diberi izin.",
        "Port MQTT yang dianalisis adalah port 1883.",
        "Serangan DoS yang diuji difokuskan pada SYN flood berbasis TCP.",
        "Komponen sniffer digunakan untuk menangkap packet capture, sedangkan prober digunakan untuk mengirim pesan MQTT normal.",
        "Metrik performa client dibatasi pada success rate publish dan RTT.",
        "Metrik trafik dibatasi pada laju paket TCP, laju paket SYN, estimasi half-open connection, rasio SYN/SYN-ACK, dan burstiness.",
        "Mitigasi yang dibahas adalah rate limiting pada sisi server atau firewall.",
        "Pengujian dilakukan secara terkontrol dan tidak ditujukan untuk mengganggu layanan produksi di luar izin penelitian.",
    ])

    h2(doc, "1.4", "Tujuan Penelitian")
    numlist(doc, [
        "Menganalisis karakteristik trafik MQTT pada kondisi normal.",
        "Menganalisis perubahan trafik saat terjadi serangan SYN flood terhadap broker MQTT.",
        "Mengukur dampak serangan terhadap success rate dan RTT client MQTT normal.",
        "Mengevaluasi rate limiting sebagai mekanisme mitigasi terhadap DoS pada broker MQTT.",
    ])

    h2(doc, "1.5", "Manfaat Penelitian")
    h3(doc, "1.5.1", "Manfaat Teoritis")
    p(doc, "Penelitian ini dapat memperkaya kajian mengenai keamanan dan ketersediaan layanan MQTT, khususnya analisis serangan SYN flood menggunakan data packet capture dan log client.")
    h3(doc, "1.5.2", "Manfaat Praktis")
    p(doc, "Penelitian ini dapat menjadi acuan awal bagi pengelola server dalam memahami indikator trafik DoS dan penerapan rate limiting untuk menjaga stabilitas layanan broker MQTT.")
    h3(doc, "1.5.3", "Manfaat Akademik")
    p(doc, "Penelitian ini dapat menjadi dasar bagi penelitian lanjutan mengenai monitoring, mitigasi, dan evaluasi keamanan layanan IoT di lingkungan Universitas Mataram.")

    h2(doc, "1.6", "Sistematika Penulisan")
    p(doc, "BAB I menjelaskan latar belakang, rumusan masalah, batasan masalah, tujuan, manfaat, dan sistematika penulisan. BAB II menjelaskan teori penunjang dan penelitian terkait. BAB III menjelaskan metodologi, arsitektur, alat dan bahan, skenario, metrik, serta metode analisis. BAB IV menjelaskan hasil implementasi, hasil pengujian awal, dan pembahasan berdasarkan data yang tersedia.")


def chapter2(doc):
    h1(doc, "BAB II", "TINJAUAN PUSTAKA")
    h2(doc, "2.1", "Penelitian Terkait")
    table(doc, "Tabel 2.1 Penelitian Terkait", ["No", "Topik Penelitian", "Fokus", "Keterkaitan"], [
        ["1", "Keamanan MQTT pada IoT", "Autentikasi, TLS, dan keamanan komunikasi MQTT", "Menjelaskan posisi MQTT sebagai protokol IoT yang perlu diamankan"],
        ["2", "Kerentanan MQTT terhadap DoS", "Analisis broker ketika menerima trafik berlebih", "Menjadi dasar bahwa broker MQTT dapat terdampak serangan ketersediaan"],
        ["3", "SYN flood pada layanan TCP", "Peningkatan paket SYN dan koneksi half-open", "Menjadi dasar indikator serangan yang digunakan"],
        ["4", "Packet capture dengan Wireshark/tshark", "Perekaman dan ekstraksi paket jaringan", "Menjadi dasar penggunaan sniffer dalam penelitian"],
        ["5", "Rate limiting pada server", "Pembatasan jumlah koneksi atau paket", "Menjadi dasar mitigasi yang diuji"],
    ])
    p(doc, "Penelitian ini mengambil posisi pada analisis eksperimental terhadap broker MQTT. Perbedaannya dengan penelitian yang hanya membahas deteksi adalah penelitian ini menghubungkan data trafik dengan performa client MQTT normal dan mengevaluasi mitigasi rate limiting.")

    h2(doc, "2.2", "Internet of Things")
    p(doc, "Internet of Things adalah konsep yang menghubungkan perangkat fisik ke jaringan agar perangkat dapat mengirimkan data dan menerima perintah secara otomatis. Perangkat IoT umumnya memiliki keterbatasan sumber daya sehingga membutuhkan protokol komunikasi yang ringan.")
    h2(doc, "2.3", "MQTT")
    p(doc, "MQTT adalah protokol komunikasi ringan yang menggunakan model publish-subscribe. Client publisher mengirim pesan ke topic tertentu pada broker, sedangkan subscriber menerima pesan dari topic yang diikuti. MQTT banyak digunakan pada sistem IoT karena sederhana dan hemat bandwidth.")
    h2(doc, "2.4", "Broker MQTT")
    p(doc, "Broker MQTT adalah server pusat yang menerima koneksi client, menerima pesan publish, dan meneruskan pesan kepada subscriber. Karena broker menjadi pusat komunikasi, ketersediaannya berpengaruh langsung terhadap sistem MQTT secara keseluruhan.")
    h2(doc, "2.5", "TCP dan Three-Way Handshake")
    p(doc, "TCP membentuk koneksi melalui three-way handshake yang terdiri dari SYN, SYN-ACK, dan ACK. MQTT berjalan di atas TCP, sehingga proses komunikasi MQTT bergantung pada keberhasilan pembentukan koneksi TCP.")
    h2(doc, "2.6", "Denial of Service")
    p(doc, "Denial of Service adalah serangan yang bertujuan mengganggu ketersediaan layanan. Serangan ini dapat dilakukan dengan mengirim trafik dalam jumlah besar atau mengeksploitasi proses tertentu sehingga server kewalahan melayani permintaan.")
    h2(doc, "2.7", "SYN Flood")
    p(doc, "SYN flood adalah DoS pada lapisan transport. Serangan ini mengirim banyak paket SYN menuju server. Indikatornya antara lain peningkatan laju SYN, peningkatan koneksi half-open, dan perubahan rasio SYN terhadap SYN-ACK.")
    h2(doc, "2.8", "Packet Capture dan Tshark")
    p(doc, "Packet capture adalah proses merekam paket jaringan pada interface tertentu. Tshark digunakan untuk menangkap dan mengekstrak paket secara command-line. Hasil capture dapat dianalisis ulang melalui file pcapng dan CSV raw flow.")
    h2(doc, "2.9", "Sniffer dan Prober")
    p(doc, "Sniffer adalah komponen yang menangkap trafik jaringan. Prober adalah client uji yang mengirim pesan MQTT normal secara berkala. Dalam penelitian ini, sniffer dan prober digunakan bersama agar data trafik dan performa client dapat dibandingkan.")
    h2(doc, "2.10", "Rate Limiting")
    p(doc, "Rate limiting adalah mekanisme pembatasan laju trafik atau koneksi dalam periode waktu tertentu. Dalam penelitian ini, rate limiting digunakan sebagai mitigasi terhadap serangan SYN flood yang menuju port MQTT broker.")


def chapter3(doc, arch: Path, flow: Path):
    h1(doc, "BAB III", "METODOLOGI PENELITIAN")
    h2(doc, "3.1", "Jenis Penelitian")
    p(doc, "Penelitian ini merupakan penelitian eksperimental kuantitatif. Eksperimental karena pengujian dilakukan melalui skenario yang dikendalikan, dan kuantitatif karena hasil dianalisis menggunakan nilai numerik seperti jumlah paket, success rate, RTT, dan metrik trafik.")
    h2(doc, "3.2", "Arsitektur Sistem")
    p(doc, "Arsitektur sistem terdiri dari client/host monitoring, broker MQTT pada server Universitas Mataram, dan attacker terkontrol. Client menjalankan prober untuk mengirim pesan MQTT normal serta menjalankan packet capture untuk merekam trafik. Broker MQTT menjadi target layanan, sedangkan attacker mengirim paket SYN menuju port 1883 pada skenario serangan.")
    fig(doc, arch, "Gambar 3.1 Arsitektur Pengujian Broker MQTT Universitas Mataram")
    h2(doc, "3.3", "Alur Data")
    p(doc, "Alur data dimulai dari proses packet capture dan pengiriman pesan MQTT normal. Capture menghasilkan file pcapng. File tersebut diekstrak menjadi raw flow dan dikompilasi menjadi metrics.csv. Prober menghasilkan prober.csv. Kedua kelompok data kemudian dibandingkan pada BAB IV.")
    fig(doc, flow, "Gambar 3.2 Alur Data Pengujian")
    h2(doc, "3.4", "Alat dan Bahan")
    table(doc, "Tabel 3.1 Alat dan Bahan", ["No", "Alat/Bahan", "Fungsi"], [
        ["1", "Laptop client/host monitoring", "Menjalankan prober, capture, dan analisis data"],
        ["2", "Server Universitas Mataram", "Menjalankan broker MQTT sebagai target pengujian"],
        ["3", "WireGuard VPN", "Jalur akses resmi ke server pengujian"],
        ["4", "Mosquitto", "Broker MQTT dan client publish"],
        ["5", "tshark", "Menangkap dan mengekstrak trafik jaringan"],
        ["6", "hping3", "Mengirim SYN flood terkontrol pada skenario serangan"],
        ["7", "Python", "Mengolah raw flow dan menghitung metrik"],
        ["8", "nftables/rate limiting", "Mitigasi yang akan diuji pada tahap lanjutan"],
    ])
    h2(doc, "3.5", "File Project dan Fungsinya")
    table(doc, "Tabel 3.2 File Project dan Fungsinya", ["No", "File", "Fungsi"], [
        ["1", "sniffer_capture.sh", "Menjalankan tshark untuk menghasilkan capture.pcapng"],
        ["2", "mqtt_prober.sh", "Mengirim publish MQTT normal dan mencatat RTT"],
        ["3", "extract_raw_flow.sh", "Mengekstrak pcapng menjadi capture_raw_flow.csv"],
        ["4", "compile.py", "Menghitung metrics.csv dari raw flow"],
        ["5", "lab_env.sh", "Menyimpan konfigurasi target, port, interface, dan folder output"],
        ["6", "run_all_cli.sh", "Menjalankan proses pengujian secara interaktif"],
    ])
    h2(doc, "3.6", "Skenario Pengujian")
    table(doc, "Tabel 3.3 Skenario Pengujian", ["Skenario", "Kondisi", "Tujuan"], [
        ["Baseline", "Broker berjalan tanpa serangan", "Menentukan kondisi normal sebagai pembanding"],
        ["SYN flood terkendali", "Broker menerima trafik SYN dengan rate rendah", "Melihat perubahan indikator trafik serangan"],
        ["SYN flood + rate limiting", "Rate limiting aktif saat serangan dijalankan", "Menguji efektivitas mitigasi"],
    ])
    h2(doc, "3.7", "Metrik Pengujian")
    table(doc, "Tabel 3.4 Metrik Pengujian", ["Metrik", "Sumber", "Makna"], [
        ["success rate", "prober.csv", "Persentase publish MQTT yang berhasil"],
        ["rtt_ms", "prober.csv", "Waktu proses publish MQTT dari client"],
        ["tcp_pkt_rate_target", "metrics.csv", "Laju paket TCP menuju broker"],
        ["syn_rate_target", "metrics.csv", "Laju paket SYN menuju broker"],
        ["half_open_conn_target", "metrics.csv", "Estimasi koneksi TCP yang belum selesai"],
        ["syn_ack_ratio_target", "metrics.csv", "Perbandingan paket SYN terhadap SYN-ACK"],
        ["burstiness_target", "metrics.csv", "Tingkat lonjakan trafik dalam window pengamatan"],
    ])
    h2(doc, "3.8", "Metode Analisis Data")
    p(doc, "Analisis dilakukan dengan membandingkan metrik baseline dan skenario serangan. Baseline digunakan sebagai kondisi acuan. Apabila skenario serangan menunjukkan peningkatan syn_rate_target dan half_open_conn_target, maka serangan dapat dikatakan terlihat pada data capture. Dampak terhadap layanan MQTT dilihat dari success rate dan RTT prober.")


def chapter4(doc, rows: list[dict[str, str]], chart_traffic: Path, chart_prober: Path):
    h1(doc, "BAB IV", "HASIL DAN PEMBAHASAN")
    h2(doc, "4.1", "Gambaran Umum Pengujian")
    p(doc, "Bab ini menyajikan hasil pengujian awal yang telah dilakukan pada broker MQTT Universitas Mataram. Data yang digunakan berasal dari dua skenario, yaitu baseline dan SYN flood terkendali. Skenario mitigasi rate limiting belum dimasukkan sebagai hasil final karena data pengujian mitigasi belum lengkap. Oleh karena itu, pembahasan pada BAB IV ini diposisikan sebagai bahan bimbingan dan dasar pengujian lanjutan.")
    table(doc, "Tabel 4.1 Lingkungan Pengujian Awal", ["Parameter", "Nilai"], [
        ["Broker MQTT", "Server Universitas Mataram"],
        ["IP broker", "172.16.10.44"],
        ["Port broker", "1883"],
        ["Jalur akses", "VPN WireGuard"],
        ["Interface capture", "unram"],
        ["Skenario baseline", "Publish MQTT normal tanpa serangan"],
        ["Skenario SYN flood", "hping3, rate rendah sekitar 10 paket SYN/detik selama 30 detik"],
        ["Output data", "capture.pcapng, capture_raw_flow.csv, metrics.csv, prober.csv, summary.txt"],
    ])
    p(doc, "Pengujian awal sengaja dilakukan dengan rate rendah agar tidak menimbulkan gangguan berlebihan pada server dan tidak memicu pemblokiran akses VPN. Dengan pendekatan ini, data yang dihasilkan aman digunakan sebagai bukti awal bahwa pola SYN flood dapat terlihat pada capture, meskipun belum menunjukkan gangguan availability yang besar.")

    h2(doc, "4.2", "Hasil Pengujian Baseline")
    p(doc, "Pada skenario baseline, prober mengirim pesan MQTT normal ke broker tanpa menjalankan serangan. Hasil baseline menunjukkan bahwa broker dapat diakses dan dapat menerima publish MQTT dari client penguji. Jumlah publish prober adalah 41 pesan, seluruhnya berhasil, sehingga success rate mencapai 100 persen.")
    p(doc, "Nilai rata-rata RTT pada baseline adalah 498,71 ms, dengan RTT P95 sebesar 1255 ms dan RTT maksimum sebesar 2083 ms. Dari sisi trafik, rata-rata laju paket TCP menuju broker adalah 5,2258 paket per detik, sedangkan rata-rata laju paket SYN adalah 0,8197 paket per detik. Nilai ini menjadi acuan untuk membandingkan kondisi serangan.")

    h2(doc, "4.3", "Hasil Pengujian SYN Flood Terkendali")
    p(doc, "Pada skenario SYN flood terkendali, attacker mengirim paket SYN ke port 1883 broker MQTT dengan rate rendah. Prober tetap dijalankan bersamaan untuk melihat apakah client normal masih dapat melakukan publish. Hasil menunjukkan 42 publish berhasil dan 0 publish gagal, sehingga success rate tetap 100 persen.")
    p(doc, "Walaupun success rate tetap stabil, indikator trafik berubah cukup jelas. Rata-rata laju paket TCP menuju broker meningkat dari 5,2258 menjadi 12,4000 paket per detik. Rata-rata laju paket SYN meningkat dari 0,8197 menjadi 4,3729 paket per detik. Rata-rata estimasi half-open connection juga meningkat dari 0,2258 menjadi 3,6333. Peningkatan ini sesuai dengan karakteristik SYN flood karena serangan memang menambah permintaan koneksi awal TCP menuju broker.")

    h2(doc, "4.4", "Rekap Perbandingan Baseline dan SYN Flood")
    important = [
        "Jumlah publish prober",
        "Publish berhasil",
        "Publish gagal",
        "Success rate prober",
        "Rata-rata RTT prober (ms)",
        "RTT P95 prober (ms)",
        "Rata-rata laju paket TCP ke broker (pkt/s)",
        "Maksimum laju paket TCP ke broker (pkt/s)",
        "Rata-rata laju paket SYN ke broker (pkt/s)",
        "Maksimum laju paket SYN ke broker (pkt/s)",
        "Rata-rata estimasi half-open connection",
        "Maksimum estimasi half-open connection",
        "Rata-rata rasio SYN/SYN-ACK",
        "Maksimum rasio SYN/SYN-ACK",
        "Jumlah baris raw flow",
        "Ukuran file pcap",
    ]
    table(doc, "Tabel 4.2 Rekap Perbandingan Baseline dan SYN Flood", ["Indikator", "Baseline", "SYN Flood", "Selisih"], [
        [name, metric(rows, name, "baseline"), metric(rows, name, "synflood_low"), metric(rows, name, "selisih_syn_minus_baseline")]
        for name in important
    ])
    fig(doc, chart_traffic, "Gambar 4.1 Perbandingan Indikator Trafik Baseline dan SYN Flood")
    fig(doc, chart_prober, "Gambar 4.2 Perbandingan Performa Prober Baseline dan SYN Flood")

    h2(doc, "4.5", "Pembahasan Indikator Trafik")
    p(doc, "Perubahan paling kuat terlihat pada metrik syn_rate_target dan half_open_conn_target. Peningkatan syn_rate_target menunjukkan bahwa jumlah paket SYN menuju broker meningkat selama skenario serangan. Peningkatan half_open_conn_target menunjukkan bahwa terdapat lebih banyak permintaan koneksi TCP yang berada pada tahap awal koneksi.")
    p(doc, "Maksimum syn_rate_target pada baseline adalah 1 paket per detik, sedangkan pada SYN flood terkendali meningkat menjadi 11 paket per detik. Maksimum half_open_conn_target juga meningkat dari 1 menjadi 11. Angka ini menunjukkan bahwa skenario serangan berhasil memunculkan pola trafik yang berbeda dari baseline, meskipun rate serangan masih rendah.")
    p(doc, "Jumlah baris raw flow meningkat dari 586 menjadi 1221, sedangkan ukuran file pcap meningkat dari 53832 byte menjadi 100908 byte. Peningkatan ini mendukung temuan bahwa jumlah trafik yang tertangkap pada skenario SYN flood lebih besar dibandingkan baseline.")

    h2(doc, "4.6", "Pembahasan Performa Prober MQTT")
    p(doc, "Dari sisi prober MQTT, kedua skenario menunjukkan success rate 100 persen. Pada baseline terdapat 41 publish berhasil, sedangkan pada SYN flood terkendali terdapat 42 publish berhasil. Tidak ditemukan publish gagal pada dua skenario tersebut.")
    p(doc, "Rata-rata RTT pada skenario SYN flood terkendali justru lebih rendah dibanding baseline, yaitu dari 498,71 ms menjadi 399,95 ms. Hal ini menunjukkan bahwa pada rate pengujian rendah, serangan belum cukup kuat untuk menurunkan availability broker dari sudut pandang client normal. Dengan demikian, kesimpulan yang tepat bukan bahwa broker berhasil dibuat down, melainkan bahwa indikator trafik serangan sudah terlihat sementara layanan MQTT masih dapat merespons client normal.")

    h2(doc, "4.7", "Kesesuaian Hasil dengan Konsep SYN Flood")
    p(doc, "SYN flood bekerja dengan menambah banyak permintaan koneksi awal TCP. Karena MQTT berjalan di atas TCP, peningkatan paket SYN menuju port 1883 relevan untuk dianalisis sebagai indikasi serangan terhadap broker MQTT. Hasil pengujian awal menunjukkan peningkatan paket SYN dan estimasi half-open connection, sehingga data sesuai dengan konsep dasar SYN flood.")
    p(doc, "Namun, skenario yang digunakan masih bersifat terkendali dan rate rendah. Karena itu, hasil ini belum dapat digunakan untuk menyatakan bahwa broker mengalami kegagalan layanan. Hasil ini lebih tepat digunakan untuk menunjukkan bahwa metode capture dan analisis metrik sudah mampu membedakan kondisi baseline dan kondisi serangan awal.")

    h2(doc, "4.8", "Status Pengujian Rate Limiting")
    p(doc, "Sesuai rancangan penelitian, rate limiting merupakan skenario penting untuk menjawab efektivitas mitigasi. Pada draft ini, hasil rate limiting belum dimasukkan sebagai hasil final karena data yang tersedia baru mencakup baseline dan SYN flood terkendali. Oleh karena itu, bagian rate limiting masih menjadi pekerjaan lanjutan sebelum kesimpulan akhir skripsi disusun.")
    p(doc, "Pengujian rate limiting perlu dilakukan dengan pola yang sama seperti skenario SYN flood, yaitu menggunakan durasi, target, interface capture, dan format output yang konsisten. Hasil yang diharapkan adalah adanya perbandingan antara SYN flood tanpa rate limiting dan SYN flood dengan rate limiting. Perbandingan tersebut akan menunjukkan apakah metrik seperti syn_rate_target, half_open_conn_target, RTT, dan success rate menjadi lebih terkendali setelah mitigasi diterapkan.")

    h2(doc, "4.9", "Keterbatasan Pengujian Awal")
    p(doc, "Keterbatasan pertama adalah jumlah skenario yang tersedia belum lengkap. Data awal baru mencakup baseline dan SYN flood terkendali, sedangkan skenario rate limiting belum memiliki hasil final. Keterbatasan kedua adalah intensitas serangan masih rendah untuk menjaga keamanan pengujian dan menghindari pemblokiran akses VPN.")
    p(doc, "Keterbatasan ketiga adalah packet capture dilakukan pada jalur akses VPN dari client penguji. Data tersebut valid untuk menunjukkan trafik pengujian dari client menuju broker melalui VPN, tetapi interpretasinya perlu dibatasi pada titik observasi tersebut. Jika penelitian ingin memperoleh gambaran seluruh trafik yang benar-benar masuk ke broker, capture idealnya dijalankan langsung pada server broker atau pada titik jaringan yang dilewati seluruh trafik menuju broker.")

    h2(doc, "4.10", "Ringkasan BAB IV")
    p(doc, "Berdasarkan hasil awal, broker MQTT Universitas Mataram dapat diakses melalui VPN WireGuard pada port 1883. Baseline menunjukkan success rate 100 persen. Skenario SYN flood terkendali juga menunjukkan success rate 100 persen, tetapi indikator trafik serangan meningkat, terutama syn_rate_target dan half_open_conn_target.")
    p(doc, "Dengan demikian, hasil sementara menunjukkan bahwa metode pengujian mampu menangkap perubahan karakteristik trafik SYN flood. Akan tetapi, data belum cukup untuk menyimpulkan efektivitas rate limiting karena skenario mitigasi belum selesai. Draft BAB IV ini dapat digunakan sebagai bahan bimbingan untuk menunjukkan arah analisis, data awal, dan kebutuhan pengujian lanjutan.")


def references(doc):
    doc.add_page_break()
    center(doc, "DAFTAR PUSTAKA", bold=True)
    refs = [
        "Abdi, N. S. (2025). Analisis sistematis literatur tentang penerapan sistem informasi dalam transformasi digital. Jurnal Sistem Informasi dan Teknik Komputer, 10(1), 1-5.",
        "Susanto, F., Prasiani, N. K., dan Darmawan, P. (2022). Implementasi Internet of Things dalam kehidupan sehari-hari. Jurnal Imagine, 2(1), 35-40.",
        "OASIS. MQTT Version 5.0 Specification.",
        "Eclipse Foundation. Eclipse Mosquitto Documentation.",
        "The Wireshark Foundation. Wireshark and Tshark Documentation.",
        "Hartanto, S. (2024). Pengujian kehandalan jaringan terhadap TCP/SYN Flood Attack dengan metode simulasi serangan lokal. Jurnal Ilmiah Teknik Elektro.",
        "Purwoko, M., dan Hilal, H. (2019). Analisis penerapan firewall nftables sebagai sistem keamanan server pada mesin virtualisasi. InComTech, 9(1), 1-22.",
        "Vaccari, I., Aiello, M., dan Cambiaso, E. (2020). SlowITe, a novel Denial of Service attack affecting MQTT. Sensors, 20(10), 2932.",
        "NETSCOUT. (2024). DDoS Threat Intelligence Report.",
        "Bangare, P. S., dan Patil, K. P. (2024). Enhancing MQTT security for Internet of Things: Lightweight two-way authorization and authentication with advanced security measures. Measurement: Sensors.",
    ]
    for ref in refs:
        p(doc, ref, indent=False)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    rows = read_comparison()
    arch = make_architecture()
    flow = make_flow()
    chart_traffic = make_bar_chart(
        rows,
        "chart_trafik.png",
        "Perbandingan Indikator Trafik",
        [
            "Rata-rata laju paket TCP ke broker (pkt/s)",
            "Rata-rata laju paket SYN ke broker (pkt/s)",
            "Rata-rata estimasi half-open connection",
            "Rata-rata rasio SYN/SYN-ACK",
        ],
    )
    chart_prober = make_bar_chart(
        rows,
        "chart_prober.png",
        "Perbandingan Performa Prober",
        [
            "Jumlah publish prober",
            "Publish berhasil",
            "Publish gagal",
            "Rata-rata RTT prober (ms)",
        ],
    )

    doc = Document()
    doc_style(doc)
    cover(doc)
    front_matter(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc, arch, flow)
    chapter4(doc, rows, chart_traffic, chart_prober)
    references(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
