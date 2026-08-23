#!/usr/bin/env python3
"""
Generate draft skripsi DOCX untuk project MQTT DoS rate limiting konsep lama.

Konsep dokumen:
- PC A / host monitoring menjalankan sniffer dan prober.
- Broker MQTT menjadi target layanan.
- Attacker berjalan terpisah pada skenario DoS.
- Data utama: log prober, pcapng, raw_flow.csv, metrics.csv.
"""

from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "doc"
IMG_DIR = ROOT / "tmp" / "docs" / "draft_skripsi_legacy_images"
OUT_DOCX = OUT_DIR / "DRAFT_SKRIPSI_MQTT_DOS_RATE_LIMITING_UNRAM.docx"


TITLE = (
    "ANALISIS SERANGAN DENIAL OF SERVICE PADA BROKER MQTT DAN "
    "MITIGASINYA MENGGUNAKAN RATE LIMITING PADA SERVER UNIVERSITAS MATARAM"
)


def font(size: int = 28, bold: bool = False):
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size)
    return ImageFont.load_default()


def rounded_box(draw, xy, text, fill="#ffffff", outline="#111827", width=3, text_fill="#111827", fnt=None):
    if fnt is None:
        fnt = font(26, True)
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=width)
    x1, y1, x2, y2 = xy
    lines = text.split("\n")
    heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        widths.append(bbox[2] - bbox[0])
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + (len(lines) - 1) * 8
    y = y1 + ((y2 - y1) - total_h) / 2
    for line, w, h in zip(lines, widths, heights):
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, fill=text_fill, font=fnt)
        y += h + 8


def arrow(draw, start, end, label=None):
    draw.line([start, end], fill="#111827", width=4)
    sx, sy = start
    ex, ey = end
    if ex >= sx:
        pts = [(ex, ey), (ex - 16, ey - 9), (ex - 16, ey + 9)]
    else:
        pts = [(ex, ey), (ex + 16, ey - 9), (ex + 16, ey + 9)]
    draw.polygon(pts, fill="#111827")
    if label:
        f = font(20, False)
        bbox = draw.textbbox((0, 0), label, font=f)
        tx = (sx + ex) / 2 - (bbox[2] - bbox[0]) / 2
        ty = (sy + ey) / 2 - 34
        draw.rounded_rectangle((tx - 8, ty - 5, tx + (bbox[2] - bbox[0]) + 8, ty + 27), radius=8, fill="#ffffff")
        draw.text((tx, ty), label, fill="#111827", font=f)


def make_diagrams():
    IMG_DIR.mkdir(parents=True, exist_ok=True)

    # Diagram arsitektur
    img = Image.new("RGB", (1500, 850), "#f8fafc")
    d = ImageDraw.Draw(img)
    title_font = font(38, True)
    d.text((430, 45), "Arsitektur Pengujian Konsep Sniffer dan Prober", fill="#111827", font=title_font)
    rounded_box(d, (80, 180, 430, 430), "PC A / Host Monitoring\nSniffer: tshark\nProber: mosquitto_pub\nAnalisis: Python", fill="#e0f2fe", outline="#0369a1")
    rounded_box(d, (610, 180, 960, 430), "Broker MQTT\nMosquitto\nPort 1883\nTarget Pengujian", fill="#dcfce7", outline="#15803d")
    rounded_box(d, (1080, 180, 1430, 430), "Attacker\nMesin terpisah\nSYN flood\nSaat skenario DoS", fill="#fee2e2", outline="#b91c1c")
    arrow(d, (430, 285), (610, 285), "trafik MQTT normal")
    arrow(d, (1080, 335), (960, 335), "trafik serangan")
    d.line([(255, 430), (255, 560), (760, 560), (760, 430)], fill="#0369a1", width=4)
    d.polygon([(760, 430), (751, 446), (769, 446)], fill="#0369a1")
    d.text((420, 590), "Sniffer menangkap trafik yang terlihat pada interface PC A", fill="#111827", font=font(24, False))
    d.text((130, 700), "Catatan validitas:", fill="#b45309", font=font(24, True))
    d.text((130, 735), "capture di PC A hanya lengkap jika trafik attacker-broker terlihat", fill="#b45309", font=font(24, True))
    d.text((130, 770), "pada interface yang dicapture.", fill="#b45309", font=font(24, True))
    img.save(IMG_DIR / "arsitektur.png")

    # Diagram alur penelitian
    img = Image.new("RGB", (1500, 1050), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((575, 45), "Alur Penelitian", fill="#111827", font=font(40, True))
    steps = [
        ("Identifikasi Masalah", "Broker MQTT dapat terganggu oleh DoS"),
        ("Studi Literatur", "MQTT, DoS, SYN flood, sniffing, rate limiting"),
        ("Perancangan Testbed", "PC A, broker MQTT, attacker"),
        ("Implementasi Tool", "sniffer, prober, ekstraksi, compile metrics"),
        ("Pengujian Normal", "baseline trafik dan performa MQTT"),
        ("Pengujian DoS", "SYN flood terhadap port broker"),
        ("Pengujian Mitigasi", "rate limiting lalu serangan ulang"),
        ("Analisis Data", "bandingkan log, raw flow, metrics"),
        ("Kesimpulan", "dampak serangan dan efektivitas mitigasi"),
    ]
    x, y = 300, 130
    for i, (head, body) in enumerate(steps):
        rounded_box(d, (x, y, x + 900, y + 75), f"{i+1}. {head}\n{body}", fill="#f1f5f9", outline="#334155", fnt=font(22, True))
        if i < len(steps) - 1:
            arrow(d, (750, y + 75), (750, y + 112))
        y += 105
    img.save(IMG_DIR / "alur_penelitian.png")

    # Diagram alur data
    img = Image.new("RGB", (1500, 850), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((550, 45), "Alur Data Eksperimen", fill="#111827", font=font(40, True))
    boxes = [
        ((80, 210, 350, 370), "sniffer_capture.sh\nfile pcapng"),
        ((430, 210, 700, 370), "extract_raw_flow.sh\nraw_flow.csv"),
        ((780, 210, 1050, 370), "compile.py\nmetrics.csv"),
        ((1130, 210, 1420, 370), "Analisis\nperbandingan skenario"),
        ((430, 520, 700, 680), "mqtt_prober.sh\nprobe_log.csv"),
    ]
    for xy, text in boxes:
        rounded_box(d, xy, text, fill="#eef2ff", outline="#4338ca", fnt=font(24, True))
    arrow(d, (350, 290), (430, 290))
    arrow(d, (700, 290), (780, 290))
    arrow(d, (1050, 290), (1130, 290))
    arrow(d, (700, 600), (1130, 600), "performa client")
    d.line([(1275, 520), (1275, 370)], fill="#111827", width=4)
    d.polygon([(1275, 370), (1266, 386), (1284, 386)], fill="#111827")
    img.save(IMG_DIR / "alur_data.png")


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def style_doc(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(4)
    sec.left_margin = Cm(4)
    sec.bottom_margin = Cm(3)
    sec.right_margin = Cm(3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for name, size in [("Heading 1", 12), ("Heading 2", 12), ("Heading 3", 12)]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(0)


def p(doc, text="", align=None, bold=False, italic=False, first_indent=True):
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.space_after = Pt(0)
    if first_indent:
        par.paragraph_format.first_line_indent = Cm(1.25)
    if align:
        par.alignment = align
    else:
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return par


def center(doc, text="", bold=False):
    return p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=bold, first_indent=False)


def h1(doc, bab, title):
    doc.add_page_break()
    center(doc, bab, bold=True)
    center(doc, title, bold=True)
    p(doc, "", first_indent=False)


def h2(doc, number, title):
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after = Pt(0)
    run = par.add_run(f"{number} {title}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return par


def h3(doc, number, title):
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.space_before = Pt(3)
    par.paragraph_format.space_after = Pt(0)
    run = par.add_run(f"{number} {title}")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return par


def bullet(doc, items):
    for item in items:
        par = doc.add_paragraph(style=None)
        par.paragraph_format.left_indent = Cm(0.75)
        par.paragraph_format.first_line_indent = Cm(-0.35)
        par.paragraph_format.line_spacing = 1.5
        run = par.add_run(f"- {item}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def number_list(doc, items):
    for i, item in enumerate(items, 1):
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.75)
        par.paragraph_format.first_line_indent = Cm(-0.35)
        par.paragraph_format.line_spacing = 1.5
        run = par.add_run(f"{i}. {item}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def table(doc, caption, headers, rows, widths=None):
    center(doc, caption, bold=True)
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for i, head in enumerate(headers):
        set_cell_text(hdr.cells[i], head, bold=True)
        set_cell_shading(hdr.cells[i], "D9EAF7")
        if widths:
            hdr.cells[i].width = Cm(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), bold=False)
            if widths:
                cells[i].width = Cm(widths[i])
    p(doc, "", first_indent=False)
    return t


def figure(doc, path, caption, width=5.9):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    run.add_picture(str(path), width=Inches(width))
    center(doc, caption, bold=True)


def cover(doc):
    for _ in range(2):
        p(doc, "", first_indent=False)
    center(doc, TITLE, bold=True)
    p(doc, "", first_indent=False)
    center(doc, "DRAFT SKRIPSI", bold=True)
    p(doc, "", first_indent=False)
    center(doc, "Diajukan untuk memenuhi sebagian persyaratan memperoleh gelar Sarjana Komputer", bold=False)
    p(doc, "", first_indent=False)
    center(doc, "Oleh:", bold=False)
    center(doc, "NAMA MAHASISWA", bold=True)
    center(doc, "NIM: F1D0XXXXX", bold=True)
    for _ in range(5):
        p(doc, "", first_indent=False)
    center(doc, "PROGRAM STUDI TEKNIK INFORMATIKA", bold=True)
    center(doc, "FAKULTAS TEKNIK", bold=True)
    center(doc, "UNIVERSITAS MATARAM", bold=True)
    center(doc, "2026", bold=True)


def front_matter(doc):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(4)
    sec.left_margin = Cm(4)
    sec.bottom_margin = Cm(3)
    sec.right_margin = Cm(3)
    sec.footer.is_linked_to_previous = False
    add_page_number(sec)
    center(doc, "KATA PENGANTAR", bold=True)
    p(doc, "Puji syukur penulis panjatkan ke hadirat Tuhan Yang Maha Esa karena atas rahmat dan karunia-Nya draft skripsi ini dapat disusun. Draft ini membahas analisis serangan Denial of Service pada broker MQTT dan mitigasinya menggunakan rate limiting pada server Universitas Mataram.")
    p(doc, "Dokumen ini masih berupa draft awal yang perlu disesuaikan kembali dengan arahan dosen pembimbing, data hasil pengujian, format resmi fakultas, serta identitas mahasiswa. Walaupun demikian, struktur utama, alur penelitian, teori, metodologi, rancangan pengujian, dan kerangka pembahasan hasil telah disusun agar dapat menjadi dasar penyusunan skripsi lengkap.")
    p(doc, "Penulis menyadari bahwa draft ini masih memerlukan penyempurnaan. Oleh karena itu, kritik dan saran sangat diperlukan agar penelitian dapat disusun dengan lebih baik.")
    p(doc, "Mataram, 2026", align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)
    p(doc, "Penulis", align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)

    doc.add_page_break()
    center(doc, "ABSTRAK", bold=True)
    p(doc, "Message Queuing Telemetry Transport (MQTT) merupakan protokol komunikasi ringan yang banyak digunakan pada sistem Internet of Things. Dalam arsitektur MQTT, broker berperan sebagai pusat komunikasi antara publisher dan subscriber. Peran broker yang sentral membuat ketersediaan broker menjadi aspek penting, karena gangguan pada broker dapat memengaruhi proses pengiriman data oleh client normal. Salah satu ancaman terhadap ketersediaan layanan adalah Denial of Service, khususnya SYN flood yang menargetkan proses pembentukan koneksi TCP.")
    p(doc, "Penelitian ini bertujuan menganalisis dampak serangan SYN flood terhadap broker MQTT dan mengevaluasi rate limiting sebagai mitigasi. Konsep pengujian yang digunakan adalah konsep sniffer dan prober. PC A atau host monitoring menjalankan sniffer untuk menangkap trafik MQTT menggunakan tshark dan menjalankan prober untuk mengirim pesan MQTT normal menggunakan mosquitto_pub. Broker MQTT menjadi target layanan, sedangkan attacker dijalankan pada mesin terpisah saat skenario DoS. Data yang dikumpulkan meliputi file capture pcapng, raw flow CSV, log prober, dan metrik trafik yang dihitung menggunakan Python.")
    p(doc, "Pengujian dirancang dalam tiga skenario, yaitu kondisi normal, kondisi SYN flood, dan kondisi SYN flood dengan rate limiting. Analisis dilakukan dengan membandingkan laju paket TCP, laju paket SYN, estimasi koneksi half-open, rasio SYN terhadap SYN-ACK, keberhasilan pengiriman pesan MQTT, dan latency prober. Hasil yang diharapkan adalah terlihatnya perubahan karakteristik trafik pada saat SYN flood serta penurunan dampak serangan setelah rate limiting diterapkan. Penelitian ini memberikan rancangan pengujian yang dapat digunakan untuk memahami dampak DoS pada broker MQTT dan mengevaluasi mitigasi berbasis pembatasan laju trafik.")
    p(doc, "Kata kunci: MQTT, broker, Denial of Service, SYN flood, packet capture, rate limiting.", first_indent=False)

    doc.add_page_break()
    center(doc, "DAFTAR ISI", bold=True)
    entries = [
        "KATA PENGANTAR",
        "ABSTRAK",
        "BAB I PENDAHULUAN",
        "1.1 Latar Belakang",
        "1.2 Rumusan Masalah",
        "1.3 Batasan Masalah",
        "1.4 Tujuan Penelitian",
        "1.5 Manfaat Penelitian",
        "1.6 Sistematika Penulisan",
        "BAB II TINJAUAN PUSTAKA",
        "2.1 Penelitian Terkait",
        "2.2 Internet of Things",
        "2.3 MQTT",
        "2.4 Broker MQTT",
        "2.5 TCP dan SYN Flood",
        "2.6 Packet Capture",
        "2.7 Rate Limiting",
        "BAB III METODOLOGI PENELITIAN",
        "BAB IV HASIL DAN PEMBAHASAN",
        "BAB V PENUTUP",
        "DAFTAR PUSTAKA",
        "LAMPIRAN",
    ]
    for item in entries:
        p(doc, item, first_indent=False)

    doc.add_page_break()
    center(doc, "DAFTAR GAMBAR", bold=True)
    for item in [
        "Gambar 3.1 Arsitektur Pengujian Konsep Sniffer dan Prober",
        "Gambar 3.2 Alur Penelitian",
        "Gambar 3.3 Alur Data Eksperimen",
    ]:
        p(doc, item, first_indent=False)

    doc.add_page_break()
    center(doc, "DAFTAR TABEL", bold=True)
    for item in [
        "Tabel 2.1 Penelitian Terkait",
        "Tabel 3.1 Alat dan Bahan",
        "Tabel 3.2 File Project dan Fungsinya",
        "Tabel 3.3 Skenario Pengujian",
        "Tabel 3.4 Metrik Pengujian",
        "Tabel 4.1 Template Rekap Hasil Pengujian",
        "Tabel 4.2 Template Perbandingan Metrik Trafik",
    ]:
        p(doc, item, first_indent=False)


def chapter1(doc):
    h1(doc, "BAB I", "PENDAHULUAN")
    h2(doc, "1.1", "Latar Belakang")
    paragraphs = [
        "Perkembangan jaringan komputer dan Internet of Things mendorong semakin banyak sistem yang membutuhkan komunikasi data ringan, stabil, dan dapat berjalan secara terus-menerus. Dalam lingkungan kampus seperti Universitas Mataram, kebutuhan komunikasi data dapat muncul pada sistem monitoring ruangan, perangkat sensor, praktikum laboratorium, sistem otomasi, dan pengembangan layanan berbasis smart campus. Sistem seperti ini membutuhkan jalur komunikasi yang dapat mengirim data secara periodik dari perangkat ke server.",
        "Salah satu protokol yang banyak digunakan pada sistem Internet of Things adalah Message Queuing Telemetry Transport atau MQTT. MQTT dirancang sebagai protokol yang ringan dan menggunakan pola publish-subscribe. Pada pola ini, perangkat pengirim data tidak langsung mengirim pesan kepada perangkat penerima, tetapi mengirim pesan ke broker. Broker kemudian meneruskan pesan tersebut kepada subscriber yang berlangganan topic tertentu.",
        "Broker MQTT memiliki posisi penting karena menjadi pusat lalu lintas pesan. Jika broker berjalan normal, client dapat mengirim data dan subscriber dapat menerima data sesuai topic. Namun jika broker mengalami gangguan, proses komunikasi MQTT dapat ikut terganggu. Gangguan pada broker dapat menyebabkan pesan terlambat, koneksi gagal, atau layanan tidak dapat digunakan oleh client normal. Oleh karena itu, ketersediaan broker MQTT menjadi salah satu aspek yang perlu dianalisis.",
        "Salah satu ancaman terhadap ketersediaan layanan adalah Denial of Service. Serangan DoS bertujuan membuat layanan menjadi lambat, sulit diakses, atau tidak dapat digunakan. Pada layanan yang berjalan di atas TCP, salah satu jenis serangan DoS yang umum dibahas adalah SYN flood. Serangan SYN flood dilakukan dengan mengirim banyak paket SYN ke server, sehingga server harus menangani banyak permintaan koneksi awal.",
        "MQTT umumnya berjalan di atas TCP, terutama pada port 1883 untuk MQTT tanpa TLS. Karena itu, broker MQTT dapat menjadi target pengujian SYN flood pada port tersebut. Ketika jumlah paket SYN menuju broker meningkat tajam, karakteristik trafik dapat berubah. Client normal yang mencoba mengirim pesan juga dapat mengalami perubahan performa, misalnya latency meningkat atau pengiriman pesan gagal.",
        "Penelitian ini menggunakan konsep sniffer dan prober untuk menganalisis kondisi broker MQTT. Sniffer adalah bagian yang menangkap paket jaringan menggunakan tshark, sedangkan prober adalah bagian yang mengirim pesan MQTT normal menggunakan mosquitto_pub. Keduanya dijalankan pada PC A atau host monitoring. Broker MQTT menjadi target layanan yang diuji, sedangkan attacker dijalankan pada mesin terpisah pada skenario DoS.",
        "Data yang dikumpulkan pada penelitian ini berupa file packet capture, log prober, data raw flow, dan metrik trafik. Packet capture digunakan sebagai bukti mentah lalu lintas jaringan yang tertangkap. Log prober digunakan untuk melihat keberhasilan pengiriman pesan MQTT normal. Raw flow digunakan untuk membaca field penting dari paket, sedangkan metrik trafik digunakan untuk membandingkan kondisi normal, kondisi serangan, dan kondisi serangan dengan mitigasi.",
        "Mitigasi yang diuji adalah rate limiting. Rate limiting merupakan mekanisme pembatasan laju trafik dalam periode waktu tertentu. Pada konteks penelitian ini, rate limiting digunakan untuk membatasi trafik SYN yang menuju layanan MQTT. Tujuannya adalah mengurangi dampak trafik berlebihan agar layanan broker tetap lebih terkendali bagi client normal.",
        "Berdasarkan uraian tersebut, penelitian ini membahas analisis serangan DoS pada broker MQTT dan mitigasinya menggunakan rate limiting pada server Universitas Mataram. Penelitian dilakukan pada lingkungan uji terkontrol agar proses pengujian aman, dapat diulang, dan dapat dibandingkan antar skenario.",
    ]
    for text in paragraphs:
        p(doc, text)

    h2(doc, "1.2", "Rumusan Masalah")
    p(doc, "Berdasarkan latar belakang tersebut, rumusan masalah utama penelitian ini adalah bagaimana dampak serangan DoS jenis SYN flood terhadap broker MQTT dan bagaimana efektivitas rate limiting dalam mengurangi dampak serangan tersebut.")
    number_list(doc, [
        "Bagaimana karakteristik trafik MQTT pada kondisi normal?",
        "Bagaimana perubahan trafik saat broker MQTT menerima serangan SYN flood?",
        "Bagaimana dampak serangan terhadap performa prober sebagai client MQTT normal?",
        "Bagaimana perbandingan kondisi serangan sebelum dan sesudah rate limiting diterapkan?",
    ])

    h2(doc, "1.3", "Batasan Masalah")
    number_list(doc, [
        "Penelitian berfokus pada broker MQTT sebagai layanan yang diuji.",
        "Protokol MQTT yang diuji adalah MQTT tanpa TLS pada port 1883.",
        "Serangan DoS yang dianalisis dibatasi pada SYN flood.",
        "Konsep pengujian menggunakan PC A sebagai host monitoring yang menjalankan sniffer dan prober.",
        "Attacker dijalankan terpisah dan hanya digunakan pada lingkungan uji yang diizinkan.",
        "Data serangan hanya dianalisis berdasarkan trafik yang tertangkap oleh interface sniffer pada PC A.",
        "Mitigasi yang diuji adalah rate limiting atau pembatasan laju trafik.",
        "Penelitian tidak membahas autentikasi MQTT, TLS, IDS machine learning, atau DDoS skala internet.",
    ])

    h2(doc, "1.4", "Tujuan Penelitian")
    number_list(doc, [
        "Menganalisis karakteristik trafik MQTT pada kondisi normal.",
        "Menganalisis perubahan trafik saat broker MQTT menerima SYN flood.",
        "Mengukur performa prober berdasarkan keberhasilan publish MQTT dan latency.",
        "Menguji rate limiting sebagai mitigasi terhadap dampak SYN flood.",
        "Membandingkan hasil pengujian pada skenario normal, DoS, dan DoS dengan rate limiting.",
    ])

    h2(doc, "1.5", "Manfaat Penelitian")
    h3(doc, "1.5.1", "Manfaat Akademik")
    p(doc, "Penelitian ini dapat menjadi referensi mengenai analisis ketersediaan broker MQTT terhadap serangan DoS, khususnya SYN flood, menggunakan data packet capture dan log client normal.")
    h3(doc, "1.5.2", "Manfaat Praktis")
    p(doc, "Penelitian ini menyediakan rancangan pengujian sederhana yang dapat digunakan untuk melihat hubungan antara trafik serangan, performa client MQTT, dan penerapan rate limiting.")
    h3(doc, "1.5.3", "Manfaat bagi Universitas Mataram")
    p(doc, "Penelitian ini dapat menjadi gambaran awal mengenai risiko gangguan ketersediaan layanan broker MQTT dan pentingnya pemantauan trafik pada server yang digunakan untuk komunikasi IoT.")

    h2(doc, "1.6", "Sistematika Penulisan")
    p(doc, "BAB I berisi pendahuluan yang menjelaskan latar belakang, rumusan masalah, batasan masalah, tujuan, manfaat, dan sistematika penulisan. BAB II berisi tinjauan pustaka mengenai IoT, MQTT, broker MQTT, TCP, SYN flood, packet capture, dan rate limiting. BAB III berisi metodologi penelitian, arsitektur pengujian, alat dan bahan, skenario pengujian, metrik, serta metode analisis data. BAB IV berisi rancangan hasil dan pembahasan yang digunakan untuk menafsirkan data pengujian. BAB V berisi kesimpulan dan saran.")


def chapter2(doc):
    h1(doc, "BAB II", "TINJAUAN PUSTAKA")
    h2(doc, "2.1", "Penelitian Terkait")
    p(doc, "Penelitian terkait digunakan untuk menempatkan penelitian ini terhadap kajian sebelumnya. Beberapa penelitian membahas MQTT pada sistem IoT, keamanan broker MQTT, serangan DoS, SYN flood, packet capture, firewall, dan rate limiting. Topik-topik tersebut menjadi dasar karena penelitian ini berada pada irisan antara komunikasi MQTT, keamanan ketersediaan layanan, dan analisis trafik jaringan.")
    table(doc, "Tabel 2.1 Penelitian Terkait", ["No", "Penelitian", "Fokus", "Keterkaitan"], [
        ["1", "Azwar dan Gary (2025)", "Keamanan komunikasi IoT berbasis MQTT", "Dasar penggunaan MQTT pada IoT"],
        ["2", "Anggoro dan Widiasari (2024)", "MQTT pada smarthome ESP32", "Mendukung konteks client IoT"],
        ["3", "Hardening MQTT Broker Server", "Pengamanan broker MQTT", "Dasar broker sebagai titik kritis"],
        ["4", "Simulasi DoS Menggunakan Hping3", "Pembangkit trafik DoS", "Dasar skenario SYN flood"],
        ["5", "Deteksi SYN Flood pada Server", "Karakteristik SYN flood", "Dasar indikator paket SYN"],
        ["6", "Analisis Trafik dengan Wireshark", "Packet capture", "Dasar pengamatan trafik"],
        ["7", "Firewall Iptables untuk DDoS", "Filtering trafik", "Dasar mitigasi server"],
        ["8", "Rate Limiting untuk Flood Attack", "Pembatasan laju trafik", "Dasar mitigasi rate limiting"],
    ], widths=[1, 4, 4, 5])
    p(doc, "Berdasarkan penelitian terkait tersebut, penelitian ini difokuskan pada analisis DoS terhadap broker MQTT dengan menggunakan data yang berasal dari sniffer dan prober. Penelitian ini tidak hanya melihat jumlah trafik, tetapi juga menghubungkan perubahan trafik dengan performa client MQTT normal.")

    sections = [
        ("2.2", "Internet of Things", [
            "Internet of Things adalah konsep penghubungan perangkat fisik ke jaringan agar perangkat dapat mengirim, menerima, dan memproses data. Perangkat IoT dapat berupa sensor, aktuator, mikrokontroler, atau perangkat monitoring. Pada sistem IoT, data biasanya dikirim secara berkala ke server atau aplikasi pemantauan.",
            "Sistem IoT membutuhkan protokol komunikasi yang ringan karena perangkat IoT sering memiliki keterbatasan sumber daya. Selain itu, komunikasi data harus stabil agar data yang dikirim tidak terlambat atau hilang. Dalam penelitian ini, konsep IoT digunakan sebagai konteks penggunaan MQTT dan broker.",
        ]),
        ("2.3", "Message Queuing Telemetry Transport", [
            "Message Queuing Telemetry Transport atau MQTT adalah protokol komunikasi ringan yang menggunakan model publish-subscribe. Publisher mengirim pesan ke broker, subscriber menerima pesan dari broker, dan topic digunakan sebagai kanal pesan.",
            "MQTT umumnya berjalan di atas TCP. Port standar MQTT tanpa TLS adalah 1883, sedangkan MQTT dengan TLS umumnya menggunakan port 8883. Karena MQTT berjalan di atas TCP, gangguan pada proses koneksi TCP dapat memengaruhi komunikasi MQTT.",
        ]),
        ("2.4", "Broker MQTT", [
            "Broker MQTT adalah server pusat yang menerima koneksi client, menerima pesan dari publisher, dan meneruskan pesan kepada subscriber. Broker menjadi titik penting karena semua komunikasi MQTT melewati broker.",
            "Jika broker lambat atau tidak dapat menerima koneksi, client normal dapat mengalami kegagalan publish atau peningkatan latency. Oleh sebab itu, broker MQTT perlu dianalisis dari sisi ketersediaan, terutama saat menerima trafik tidak normal.",
        ]),
        ("2.5", "Transmission Control Protocol", [
            "Transmission Control Protocol atau TCP adalah protokol transport berbasis koneksi. Sebelum data dikirim, TCP melakukan proses three-way handshake. Proses ini terdiri dari paket SYN, SYN-ACK, dan ACK.",
            "Pada komunikasi MQTT, client perlu membentuk koneksi TCP terlebih dahulu sebelum dapat mengirim pesan. Jika proses awal koneksi terganggu, maka proses publish MQTT dapat ikut terganggu.",
        ]),
        ("2.6", "Denial of Service", [
            "Denial of Service adalah serangan yang bertujuan membuat layanan menjadi tidak tersedia atau mengalami penurunan kualitas. Dampak DoS dapat berupa latency meningkat, koneksi gagal, timeout, atau layanan menjadi tidak responsif.",
            "Pada broker MQTT, DoS dapat mengganggu proses koneksi dan pengiriman pesan. Karena broker menjadi pusat komunikasi, gangguan terhadap broker dapat memengaruhi client yang sebenarnya normal.",
        ]),
        ("2.7", "SYN Flood", [
            "SYN flood adalah jenis DoS pada lapisan transport TCP. Serangan ini dilakukan dengan mengirim paket SYN dalam jumlah besar ke server. Server kemudian harus menangani banyak permintaan koneksi awal.",
            "Indikator SYN flood dapat dilihat dari meningkatnya paket SYN menuju target, meningkatnya estimasi koneksi half-open, dan meningkatnya rasio SYN terhadap SYN-ACK. Pada penelitian ini, SYN flood diarahkan ke port MQTT broker.",
        ]),
        ("2.8", "Packet Capture dan Tshark", [
            "Packet capture adalah proses merekam paket jaringan yang melewati interface tertentu. Data capture dapat disimpan dalam format pcap atau pcapng dan dianalisis menggunakan Wireshark atau tshark.",
            "Tshark adalah versi command-line dari Wireshark. Dalam penelitian ini, tshark digunakan oleh sniffer untuk merekam trafik MQTT dan mengekstrak field penting menjadi raw_flow.csv.",
        ]),
        ("2.9", "Sniffer dan Prober", [
            "Sniffer adalah komponen yang menangkap paket jaringan. Pada project ini, sniffer dijalankan oleh file sniffer_capture.sh dan menghasilkan file pcapng. Sniffer dijalankan pada PC A atau host monitoring.",
            "Prober adalah komponen yang mengirim pesan MQTT normal ke broker. Pada project ini, prober dijalankan oleh mqtt_prober.sh menggunakan mosquitto_pub. Prober menghasilkan log berisi timestamp, nomor urut, exit_code, dan rtt_ms.",
        ]),
        ("2.10", "Rate Limiting", [
            "Rate limiting adalah pembatasan laju trafik dalam periode waktu tertentu. Mekanisme ini dapat digunakan untuk membatasi paket yang masuk ke server agar trafik berlebihan tidak langsung diproses seluruhnya oleh layanan.",
            "Dalam penelitian ini, rate limiting digunakan sebagai mitigasi terhadap SYN flood. Efektivitasnya dilihat dari perubahan metrik trafik dan performa prober setelah mitigasi diterapkan.",
        ]),
        ("2.11", "Metrik Analisis Trafik", [
            "Metrik analisis trafik digunakan untuk mengubah data capture menjadi angka yang dapat dibandingkan. Metrik yang digunakan meliputi tcp_pkt_rate_target, tcp_byte_rate_target, syn_rate_target, rst_rate_target, half_open_conn_target, syn_ack_ratio_target, unique_src_ip_tcp_target, src_ip_entropy_tcp_target, top_src_ip_fraction_tcp_target, avg_pkt_size_target, small_pkt_ratio_target, mean_iat_target, dan burstiness_target.",
            "Metrik tersebut membantu membedakan kondisi normal dan kondisi serangan. Misalnya, syn_rate_target digunakan untuk melihat peningkatan paket SYN, sedangkan rtt_ms pada log prober digunakan untuk melihat dampak terhadap client normal.",
        ]),
        ("2.12", "Kerangka Pemikiran", [
            "Kerangka pemikiran penelitian ini dimulai dari broker MQTT yang memiliki peran penting dalam komunikasi IoT. Karena broker berjalan di atas TCP, broker berpotensi terdampak oleh SYN flood. Dampak tersebut perlu diamati melalui sniffer dan prober.",
            "Sniffer menghasilkan bukti trafik, prober menghasilkan data performa client, dan rate limiting diuji sebagai mitigasi. Hasil dari tiga skenario kemudian dibandingkan untuk menjawab rumusan masalah.",
        ]),
    ]
    for num, title, paras in sections:
        h2(doc, num, title)
        for text in paras:
            p(doc, text)


def chapter3(doc):
    h1(doc, "BAB III", "METODOLOGI PENELITIAN")
    h2(doc, "3.1", "Jenis Penelitian")
    p(doc, "Penelitian ini menggunakan pendekatan eksperimental kuantitatif. Pendekatan eksperimental digunakan karena sistem diuji pada beberapa skenario yang telah ditentukan. Pendekatan kuantitatif digunakan karena data yang dianalisis berupa angka, seperti jumlah paket, laju paket SYN, jumlah byte, latency, exit_code, dan metrik trafik.")
    p(doc, "Pengujian dilakukan pada lingkungan uji terkontrol. Lingkungan ini digunakan agar serangan tidak diarahkan ke sistem produksi publik dan agar hasil pengujian dapat diulang dengan kondisi yang relatif sama.")

    h2(doc, "3.2", "Arsitektur Sistem")
    p(doc, "Arsitektur pengujian menggunakan konsep sniffer dan prober. Komponen utama terdiri dari PC A atau host monitoring, broker MQTT, dan attacker. PC A menjalankan sniffer untuk menangkap trafik yang terlihat pada interface jaringan dan menjalankan prober untuk mengirim trafik MQTT normal. Broker MQTT menjadi target layanan. Attacker dijalankan terpisah pada skenario DoS.")
    figure(doc, IMG_DIR / "arsitektur.png", "Gambar 3.1 Arsitektur Pengujian Konsep Sniffer dan Prober")
    p(doc, "Pada arsitektur ini, PC A bukan broker dan bukan attacker. PC A berperan sebagai host monitoring. Artinya, PC A menjalankan alat pemantauan dan client normal. Broker MQTT menerima koneksi dari prober dan menerima trafik serangan dari attacker pada skenario DoS.")
    p(doc, "Catatan penting pada arsitektur ini adalah posisi sniffer berada pada PC A. Dengan demikian, data serangan akan lengkap apabila trafik attacker menuju broker terlihat pada interface yang dicapture oleh PC A. Jika trafik attacker langsung menuju broker tanpa melewati interface tersebut, maka capture pada PC A dapat tidak merekam seluruh trafik serangan. Batasan ini dicatat sebagai bagian dari validitas pengujian.")

    h2(doc, "3.3", "Alur Penelitian")
    figure(doc, IMG_DIR / "alur_penelitian.png", "Gambar 3.2 Alur Penelitian")
    p(doc, "Penelitian diawali dengan identifikasi masalah ketersediaan broker MQTT. Setelah itu dilakukan studi literatur mengenai MQTT, DoS, SYN flood, packet capture, dan rate limiting. Tahap berikutnya adalah perancangan testbed, implementasi tool, pengujian normal, pengujian serangan, pengujian mitigasi, analisis data, dan penarikan kesimpulan.")

    h2(doc, "3.4", "Alat dan Bahan")
    table(doc, "Tabel 3.1 Alat dan Bahan", ["No", "Alat/Bahan", "Fungsi"], [
        ["1", "PC A / host monitoring", "Menjalankan sniffer, prober, ekstraksi, dan analisis"],
        ["2", "Broker MQTT", "Layanan target pengujian"],
        ["3", "Mesin attacker", "Mengirim trafik SYN flood pada skenario DoS"],
        ["4", "Mosquitto broker", "Implementasi broker MQTT"],
        ["5", "mosquitto_pub", "Mengirim pesan MQTT normal dari prober"],
        ["6", "tshark", "Packet capture dan ekstraksi field trafik"],
        ["7", "Python 3", "Menghitung metrik trafik dari raw_flow.csv"],
        ["8", "pandas dan numpy", "Pengolahan data metrik"],
        ["9", "Firewall/rate limiting", "Mitigasi trafik berlebihan pada skenario DoS + rate limiting"],
    ], widths=[1, 4, 9])

    h2(doc, "3.5", "File Project")
    table(doc, "Tabel 3.2 File Project dan Fungsinya", ["No", "File", "Fungsi"], [
        ["1", "lab_env.sh", "Menyimpan konfigurasi TARGET_IP, MQTT_PORTS, NET_IFACE, LOG_DIR, dan PCAP_DIR"],
        ["2", "run_all.sh", "Menjalankan sniffer_capture.sh dan mqtt_prober.sh secara bersamaan"],
        ["3", "run_all_cli.sh", "Versi interaktif agar user dapat memasukkan IP broker, interface, dan port tanpa mengedit file konfigurasi"],
        ["4", "sniffer_capture.sh", "Menjalankan tshark untuk menyimpan trafik MQTT ke file pcapng"],
        ["5", "mqtt_prober.sh", "Mengirim pesan MQTT berkala ke broker dan mencatat exit_code serta rtt_ms"],
        ["6", "extract_raw_flow.sh", "Mengekstrak field penting dari pcapng menjadi raw_flow.csv"],
        ["7", "compile.py", "Menghitung metrik trafik dari raw_flow.csv"],
        ["8", "requirements.txt", "Daftar library Python yang dibutuhkan"],
    ], widths=[1, 4, 9])

    h2(doc, "3.6", "Alur Data Eksperimen")
    figure(doc, IMG_DIR / "alur_data.png", "Gambar 3.3 Alur Data Eksperimen")
    p(doc, "Data eksperimen dimulai dari proses capture dan proses publish MQTT normal. Sniffer menghasilkan pcapng. File pcapng kemudian diekstrak menjadi raw_flow.csv menggunakan extract_raw_flow.sh. Selanjutnya compile.py menghitung metrik trafik dan menghasilkan metrics.csv. Di sisi lain, prober menghasilkan probe_log.csv yang digunakan untuk melihat performa client MQTT normal.")

    h2(doc, "3.7", "Skenario Pengujian")
    table(doc, "Tabel 3.3 Skenario Pengujian", ["Skenario", "Kondisi Broker", "Kondisi Attacker", "Tujuan"], [
        ["Normal", "Broker berjalan normal", "Tidak aktif", "Menentukan baseline trafik dan performa prober"],
        ["DoS / SYN flood", "Broker menerima trafik normal dan serangan", "Aktif", "Melihat dampak SYN flood terhadap trafik dan performa prober"],
        ["DoS + rate limiting", "Broker menerima serangan dengan mitigasi aktif", "Aktif", "Menilai apakah rate limiting menurunkan dampak serangan"],
    ], widths=[3, 4, 3, 5])
    h3(doc, "3.7.1", "Skenario Normal")
    p(doc, "Pada skenario normal, broker MQTT berjalan tanpa serangan. PC A menjalankan sniffer dan prober. Attacker tidak dijalankan. Hasil skenario ini menjadi baseline untuk perbandingan.")
    h3(doc, "3.7.2", "Skenario DoS / SYN Flood")
    p(doc, "Pada skenario DoS, PC A tetap menjalankan sniffer dan prober. Attacker mengirim SYN flood ke broker MQTT. Tujuan skenario ini adalah melihat perubahan trafik dan performa prober saat broker menerima trafik serangan.")
    h3(doc, "3.7.3", "Skenario DoS dengan Rate Limiting")
    p(doc, "Pada skenario ini, rate limiting diterapkan sebelum atau saat serangan dijalankan sesuai rancangan pengujian. Setelah itu attacker mengirim SYN flood ke broker dan PC A tetap menjalankan sniffer serta prober. Hasilnya dibandingkan dengan skenario DoS tanpa mitigasi.")

    h2(doc, "3.8", "Metrik Pengujian")
    table(doc, "Tabel 3.4 Metrik Pengujian", ["No", "Metrik", "Sumber Data", "Makna"], [
        ["1", "exit_code", "probe_log.csv", "Status keberhasilan publish MQTT"],
        ["2", "rtt_ms", "probe_log.csv", "Waktu proses publish MQTT"],
        ["3", "tcp_pkt_rate_target", "metrics.csv", "Laju paket TCP menuju port broker"],
        ["4", "tcp_byte_rate_target", "metrics.csv", "Laju byte TCP menuju port broker"],
        ["5", "syn_rate_target", "metrics.csv", "Laju paket SYN menuju broker"],
        ["6", "half_open_conn_target", "metrics.csv", "Estimasi koneksi yang belum selesai"],
        ["7", "syn_ack_ratio_target", "metrics.csv", "Perbandingan paket SYN dan SYN-ACK"],
        ["8", "unique_src_ip_tcp_target", "metrics.csv", "Jumlah IP sumber unik"],
        ["9", "burstiness_target", "metrics.csv", "Tingkat lonjakan trafik"],
    ], widths=[1, 4, 3, 6])

    h2(doc, "3.9", "Perancangan Sniffer")
    p(doc, "Sniffer dirancang untuk menangkap trafik TCP yang berhubungan dengan layanan MQTT. File sniffer_capture.sh membaca konfigurasi dari lab_env.sh, membuat nama file capture berdasarkan timestamp, menyusun filter berdasarkan port MQTT, lalu menjalankan tshark pada interface yang telah dipilih. Output utama sniffer adalah file pcapng.")
    p(doc, "Filter capture dibuat berdasarkan port MQTT agar data yang ditangkap tetap fokus pada trafik yang relevan. Pada konfigurasi dasar, port yang digunakan adalah 1883. Jika penelitian ingin mengamati port lain, nilai MQTT_PORTS dapat diubah. Namun pada draft ini pengujian utama tetap diarahkan pada MQTT tanpa TLS pada port 1883.")
    p(doc, "Kualitas data sniffer sangat bergantung pada pemilihan interface. Jika interface yang dipilih salah, file capture dapat kosong atau tidak memuat trafik yang diperlukan. Karena itu, sebelum eksperimen dijalankan, PC A perlu memastikan interface jaringan aktif menggunakan perintah seperti ip addr atau tshark -D.")

    h2(doc, "3.10", "Perancangan Prober")
    p(doc, "Prober dirancang sebagai client MQTT normal. File mqtt_prober.sh mengirim pesan ke broker secara periodik menggunakan mosquitto_pub. Setiap pesan memiliki nomor urut sehingga proses pengiriman dapat dicatat secara runtut. Prober tidak dimaksudkan sebagai sensor fisik, melainkan simulasi client normal yang mewakili perangkat pengirim data pada sistem MQTT.")
    p(doc, "Setiap proses publish dicatat dalam file log dengan kolom timestamp, seq, exit_code, dan rtt_ms. Kolom exit_code digunakan untuk mengetahui apakah mosquitto_pub berhasil atau gagal. Kolom rtt_ms digunakan untuk melihat waktu yang diperlukan pada proses publish. Jika broker terganggu, nilai rtt_ms dapat meningkat atau exit_code dapat menunjukkan kegagalan.")

    h2(doc, "3.11", "Perancangan Ekstraksi dan Analisis Metrik")
    p(doc, "File extract_raw_flow.sh digunakan untuk mengubah file pcapng menjadi raw_flow.csv. Ekstraksi dilakukan menggunakan tshark dengan field seperti frame.time_epoch, frame.len, ip.src, ip.dst, tcp.srcport, tcp.dstport, tcp.flags.syn, tcp.flags.ack, tcp.flags.reset, tcp.stream, mqtt.msgtype, mqtt.clientid, dan mqtt.topic.")
    p(doc, "File compile.py digunakan untuk menghitung metrik per jendela waktu satu detik. Metrik ini membantu melihat perubahan trafik secara lebih objektif. Misalnya, syn_rate_target menunjukkan jumlah paket SYN menuju broker per detik, sedangkan half_open_conn_target memberikan estimasi koneksi yang belum selesai berdasarkan stream TCP.")
    p(doc, "Analisis tidak hanya melihat satu metrik. Jika hanya melihat jumlah paket, kesimpulan bisa kurang kuat. Oleh karena itu, penelitian menggabungkan metrik trafik dari capture dengan log prober. Kombinasi ini digunakan untuk melihat apakah perubahan trafik benar-benar berkaitan dengan perubahan performa client MQTT normal.")

    h2(doc, "3.12", "Prosedur Pengujian")
    number_list(doc, [
        "Menyiapkan broker MQTT dan memastikan port 1883 dapat diakses dari PC A.",
        "Menyiapkan PC A dengan tshark, mosquitto-clients, Python, pandas, dan numpy.",
        "Menentukan interface capture yang benar pada PC A.",
        "Menjalankan pengujian normal menggunakan run_all_cli.sh tanpa attacker.",
        "Menyimpan file pcapng, probe_log.csv, raw_flow.csv, dan metrics.csv dari skenario normal.",
        "Menjalankan pengujian DoS dengan attacker aktif pada testbed yang diizinkan.",
        "Menyimpan seluruh output skenario DoS.",
        "Mengaktifkan rate limiting sesuai rancangan mitigasi.",
        "Menjalankan kembali pengujian DoS dengan kondisi rate limiting aktif.",
        "Membandingkan hasil ketiga skenario berdasarkan metrik dan log prober.",
    ])

    h2(doc, "3.13", "KPI Keberhasilan Pengujian")
    table(doc, "Tabel 3.5 KPI Keberhasilan Pengujian", ["No", "KPI", "Indikator Berhasil"], [
        ["1", "Broker dapat diakses", "Prober dapat melakukan publish pada skenario normal"],
        ["2", "Capture berjalan", "File pcapng terbentuk dan memuat trafik MQTT"],
        ["3", "Ekstraksi berhasil", "File raw_flow.csv terbentuk dari pcapng"],
        ["4", "Metrik terbentuk", "compile.py menghasilkan metrics.csv"],
        ["5", "Skenario DoS terlihat", "Terdapat peningkatan indikator SYN dibanding normal"],
        ["6", "Mitigasi dapat dievaluasi", "Ada perbedaan metrik antara DoS tanpa mitigasi dan DoS + rate limiting"],
        ["7", "Data dapat dianalisis", "Setiap skenario memiliki log prober dan data trafik yang lengkap"],
    ], widths=[1, 5, 8])

    h2(doc, "3.14", "Metode Analisis Data")
    number_list(doc, [
        "Membandingkan hasil prober pada skenario normal, DoS, dan DoS + rate limiting.",
        "Membandingkan laju paket TCP dan paket SYN pada setiap skenario.",
        "Melihat apakah SYN flood menyebabkan peningkatan half-open connection dan rasio SYN terhadap SYN-ACK.",
        "Melihat apakah rate limiting menurunkan indikator trafik serangan atau menjaga performa prober.",
        "Mencatat keterbatasan capture apabila trafik attacker tidak terlihat penuh oleh interface sniffer.",
    ])

    h2(doc, "3.15", "Validitas dan Etika Pengujian")
    p(doc, "Validitas pengujian dijaga dengan menggunakan lingkungan uji yang sama untuk setiap skenario, mencatat konfigurasi IP dan interface, menggunakan durasi pengujian yang konsisten, dan menyimpan file mentah seperti pcapng serta log prober. Etika pengujian dijaga dengan tidak mengarahkan serangan ke sistem produksi publik tanpa izin. Serangan hanya dilakukan pada testbed yang memang disiapkan untuk penelitian.")


def chapter4(doc):
    h1(doc, "BAB IV", "HASIL DAN PEMBAHASAN")
    h2(doc, "4.1", "Gambaran Umum Implementasi")
    p(doc, "Bab ini disusun sebagai draft pembahasan hasil. Nilai pada tabel hasil perlu diganti dengan data final setelah pengujian dilakukan. Struktur pembahasan dibuat agar hasil eksperimen dapat ditulis secara runtut dari implementasi sistem, lingkungan pengujian, hasil skenario normal, hasil skenario DoS, hasil skenario DoS dengan rate limiting, hingga evaluasi keterbatasan.")
    p(doc, "Implementasi project terdiri dari sniffer, prober, ekstraksi raw flow, dan perhitungan metrik. Sniffer berjalan menggunakan tshark untuk menghasilkan file pcapng. Prober berjalan menggunakan mosquitto_pub untuk menghasilkan log performa client normal. File pcapng diekstrak menjadi raw_flow.csv, kemudian diolah menggunakan compile.py untuk menghasilkan metrics.csv.")

    h2(doc, "4.2", "Hasil Implementasi Komponen")
    table(doc, "Tabel 4.1 Hasil Implementasi Komponen", ["No", "Komponen", "Status", "Output"], [
        ["1", "Sniffer", "Berjalan", "pcap/mqtt_traffic_<timestamp>.pcapng"],
        ["2", "Prober", "Berjalan", "logs/mqtt_1883_probe_log.csv"],
        ["3", "Ekstraksi raw flow", "Berjalan", "pcap/mqtt_traffic_<timestamp>_raw_flow.csv"],
        ["4", "Perhitungan metrik", "Berjalan", "metrics.csv"],
        ["5", "Rate limiting", "Disesuaikan pada sisi mitigasi", "Aturan pembatasan trafik SYN"],
    ], widths=[1, 4, 4, 6])

    h2(doc, "4.3", "Lingkungan Pengujian")
    p(doc, "Lingkungan pengujian terdiri dari PC A sebagai host monitoring, broker MQTT sebagai target, dan attacker sebagai pembangkit trafik serangan. Seluruh pengujian dilakukan pada jaringan uji terkontrol. Konfigurasi seperti IP broker, port MQTT, dan interface capture dicatat pada lab_env.sh atau dimasukkan melalui run_all_cli.sh.")
    table(doc, "Tabel 4.2 Template Konfigurasi Lingkungan Pengujian", ["Parameter", "Nilai"], [
        ["IP broker MQTT", "[isi IP broker]"],
        ["Port MQTT", "1883"],
        ["Interface capture PC A", "[isi interface]"],
        ["Durasi pengujian", "[isi durasi]"],
        ["Tool capture", "tshark"],
        ["Tool prober", "mosquitto_pub"],
        ["Tool analisis", "Python, pandas, numpy"],
    ], widths=[5, 9])

    h2(doc, "4.4", "Hasil Pengujian Skenario Normal")
    p(doc, "Skenario normal digunakan sebagai baseline. Pada kondisi ini, attacker tidak berjalan. Prober mengirim pesan MQTT secara berkala ke broker, sedangkan sniffer menangkap trafik yang terlihat pada interface PC A. Hasil yang diharapkan adalah mayoritas exit_code bernilai 0 dan rtt_ms relatif stabil.")
    table(doc, "Tabel 4.3 Template Hasil Skenario Normal", ["Indikator", "Nilai", "Interpretasi"], [
        ["Jumlah pesan prober", "[isi nilai]", "Jumlah percobaan publish MQTT"],
        ["Publish berhasil", "[isi nilai]", "Pesan dengan exit_code 0"],
        ["Publish gagal", "[isi nilai]", "Pesan dengan exit_code selain 0"],
        ["Rata-rata rtt_ms", "[isi nilai]", "Latency rata-rata client normal"],
        ["Rata-rata syn_rate_target", "[isi nilai]", "Baseline paket SYN"],
        ["Rata-rata tcp_pkt_rate_target", "[isi nilai]", "Baseline paket TCP"],
    ], widths=[5, 3, 7])
    p(doc, "Jika hasil normal menunjukkan publish berhasil dan metrik trafik rendah atau stabil, maka baseline dapat digunakan sebagai pembanding untuk skenario serangan.")

    h2(doc, "4.5", "Hasil Pengujian Skenario SYN Flood")
    p(doc, "Pada skenario SYN flood, attacker mengirim paket SYN ke broker MQTT. PC A tetap menjalankan prober agar performa client normal tetap dapat diamati. Perubahan utama yang dicari adalah peningkatan syn_rate_target, peningkatan tcp_pkt_rate_target, peningkatan half_open_conn_target, dan perubahan rtt_ms atau exit_code pada prober.")
    table(doc, "Tabel 4.4 Template Hasil Skenario SYN Flood", ["Indikator", "Normal", "SYN Flood", "Perubahan"], [
        ["Publish gagal", "[isi]", "[isi]", "[naik/turun/tetap]"],
        ["Rata-rata rtt_ms", "[isi]", "[isi]", "[naik/turun/tetap]"],
        ["syn_rate_target", "[isi]", "[isi]", "[naik/turun/tetap]"],
        ["tcp_pkt_rate_target", "[isi]", "[isi]", "[naik/turun/tetap]"],
        ["half_open_conn_target", "[isi]", "[isi]", "[naik/turun/tetap]"],
        ["burstiness_target", "[isi]", "[isi]", "[naik/turun/tetap]"],
    ], widths=[4, 3, 3, 4])
    p(doc, "Apabila syn_rate_target dan half_open_conn_target meningkat tajam dibanding kondisi normal, maka data tersebut menunjukkan adanya perubahan trafik yang sesuai dengan karakteristik SYN flood. Jika rtt_ms meningkat atau exit_code gagal bertambah, maka serangan juga berdampak pada client MQTT normal.")

    h2(doc, "4.6", "Hasil Pengujian Skenario SYN Flood dengan Rate Limiting")
    p(doc, "Skenario ini digunakan untuk menilai apakah rate limiting dapat mengurangi dampak SYN flood. Analisis dilakukan dengan membandingkan hasil DoS tanpa mitigasi dan DoS dengan rate limiting. Fokus evaluasi bukan hanya apakah paket serangan hilang, tetapi apakah dampaknya terhadap trafik dan client normal menjadi lebih terkendali.")
    table(doc, "Tabel 4.5 Template Hasil Mitigasi Rate Limiting", ["Indikator", "DoS Tanpa Mitigasi", "DoS + Rate Limiting", "Interpretasi"], [
        ["Publish gagal", "[isi]", "[isi]", "[membaik/tidak]"],
        ["Rata-rata rtt_ms", "[isi]", "[isi]", "[membaik/tidak]"],
        ["syn_rate_target", "[isi]", "[isi]", "[membaik/tidak]"],
        ["half_open_conn_target", "[isi]", "[isi]", "[membaik/tidak]"],
        ["syn_ack_ratio_target", "[isi]", "[isi]", "[membaik/tidak]"],
        ["burstiness_target", "[isi]", "[isi]", "[membaik/tidak]"],
    ], widths=[4, 4, 4, 3])
    p(doc, "Rate limiting dapat dinilai efektif apabila indikator serangan menjadi lebih rendah atau performa prober lebih stabil dibanding skenario DoS tanpa mitigasi. Namun jika capture dilakukan pada PC A, interpretasi tetap harus memperhatikan apakah trafik attacker terlihat lengkap pada interface yang dicapture.")

    h2(doc, "4.7", "Pembahasan Perbandingan Antar Skenario")
    p(doc, "Perbandingan antar skenario dilakukan dengan menempatkan normal sebagai baseline. Skenario DoS menunjukkan dampak serangan terhadap broker dan client normal. Skenario DoS dengan rate limiting menunjukkan perubahan setelah mitigasi diterapkan. Dengan pola ini, pembahasan dapat menjawab rumusan masalah secara langsung.")
    table(doc, "Tabel 4.6 Template Ringkasan Perbandingan", ["Aspek", "Normal", "DoS", "DoS + Rate Limiting"], [
        ["Kondisi trafik TCP", "Stabil", "[isi hasil]", "[isi hasil]"],
        ["Paket SYN", "Rendah", "[isi hasil]", "[isi hasil]"],
        ["Performa prober", "Stabil", "[isi hasil]", "[isi hasil]"],
        ["Indikasi half-open", "Rendah", "[isi hasil]", "[isi hasil]"],
        ["Kesimpulan sementara", "Baseline", "[isi hasil]", "[isi hasil]"],
    ], widths=[4, 4, 4, 4])

    h2(doc, "4.8", "Pembahasan Berdasarkan Log Prober")
    p(doc, "Log prober digunakan untuk membaca kondisi dari sudut pandang client normal. Pada skenario normal, publish MQTT seharusnya berhasil secara konsisten. Pada skenario SYN flood, perubahan yang perlu diperhatikan adalah peningkatan rtt_ms dan munculnya exit_code gagal. Jika perubahan tersebut terjadi bersamaan dengan peningkatan indikator SYN pada metrics.csv, maka hubungan antara serangan dan gangguan client menjadi lebih kuat.")
    p(doc, "Jika rtt_ms tidak banyak berubah meskipun syn_rate meningkat, pembahasan perlu menjelaskan bahwa serangan yang diberikan belum cukup berdampak pada client normal atau broker masih mampu menangani beban. Dengan demikian, hasil yang tidak ekstrem tetap dapat dibahas secara ilmiah selama data dan kondisi pengujian dijelaskan dengan jujur.")

    h2(doc, "4.9", "Pembahasan Berdasarkan Packet Capture")
    p(doc, "Packet capture menjadi bukti mentah lalu lintas jaringan. Dari file pcapng, peneliti dapat menunjukkan bahwa trafik MQTT benar-benar tertangkap selama pengujian. File raw_flow.csv memudahkan pembacaan field penting seperti IP sumber, IP tujuan, port, flag TCP, dan timestamp paket.")
    p(doc, "Pada skenario DoS, indikator yang paling penting adalah peningkatan paket SYN menuju port broker. Jika data menunjukkan banyak paket SYN dengan ACK bernilai 0, maka pola tersebut sesuai dengan karakteristik SYN flood. Jika data juga menunjukkan peningkatan burstiness, maka trafik serangan dapat dijelaskan sebagai trafik yang datang secara meledak dalam waktu singkat.")

    h2(doc, "4.10", "Pembahasan Efektivitas Rate Limiting")
    p(doc, "Efektivitas rate limiting dinilai dengan membandingkan skenario DoS tanpa mitigasi dan skenario DoS dengan mitigasi. Jika rate limiting efektif, nilai tertentu seperti syn_rate_target, half_open_conn_target, atau rtt_ms dapat menjadi lebih rendah atau lebih stabil. Namun hasil akhir harus mengikuti data eksperimen, bukan asumsi.")
    p(doc, "Apabila rate limiting tidak menurunkan seluruh metrik, pembahasan tetap dapat diarahkan pada metrik mana yang membaik dan metrik mana yang tidak. Dalam penelitian jaringan, mitigasi tidak selalu menghilangkan serangan sepenuhnya. Mitigasi dapat dianggap bermanfaat jika mampu menurunkan dampak terhadap layanan atau membuat performa client normal lebih stabil.")

    h2(doc, "4.11", "Keterbatasan Hasil")
    p(doc, "Keterbatasan utama penelitian ini adalah posisi sniffer yang berada pada PC A. Apabila trafik attacker menuju broker tidak melewati interface PC A, maka capture dapat tidak merekam seluruh trafik serangan. Oleh karena itu, hasil pcapng perlu dipahami sebagai trafik yang terlihat dari titik pengamatan PC A. Keterbatasan ini harus disampaikan secara eksplisit agar analisis tetap jujur dan dapat dipertanggungjawabkan.")
    p(doc, "Keterbatasan lain adalah pengujian dilakukan pada lingkungan terkontrol, sehingga hasilnya tidak dapat langsung disamakan dengan DDoS skala internet. Penelitian ini lebih tepat diposisikan sebagai analisis eksperimental pada testbed untuk memahami pola trafik, dampak terhadap client normal, dan potensi rate limiting sebagai mitigasi.")


def chapter5(doc):
    h1(doc, "BAB V", "PENUTUP")
    h2(doc, "5.1", "Kesimpulan")
    p(doc, "Berdasarkan rancangan dan kerangka pengujian yang telah disusun, penelitian ini menganalisis broker MQTT sebagai layanan yang dapat terdampak oleh serangan DoS jenis SYN flood. Broker MQTT memiliki peran penting karena menjadi pusat komunikasi antara publisher dan subscriber. Gangguan terhadap broker dapat memengaruhi client normal yang mengirim pesan MQTT.")
    p(doc, "Konsep sniffer dan prober digunakan untuk mengamati kondisi pengujian. Sniffer menghasilkan file capture, sedangkan prober menghasilkan log performa client MQTT normal. Data tersebut kemudian diekstrak dan dihitung menjadi metrik trafik agar kondisi normal, kondisi SYN flood, dan kondisi SYN flood dengan rate limiting dapat dibandingkan.")
    p(doc, "Kesimpulan final harus disesuaikan dengan hasil eksperimen. Secara umum, penelitian diharapkan dapat menunjukkan apakah SYN flood meningkatkan indikator trafik seperti syn_rate_target dan half_open_conn_target, apakah serangan memengaruhi rtt_ms atau keberhasilan publish MQTT, dan apakah rate limiting dapat menurunkan dampak serangan.")

    h2(doc, "5.2", "Saran")
    number_list(doc, [
        "Pengujian sebaiknya dilakukan dengan durasi dan konfigurasi yang konsisten pada setiap skenario.",
        "Data mentah seperti pcapng, raw_flow.csv, probe_log.csv, dan metrics.csv perlu disimpan sebagai bukti hasil pengujian.",
        "Jika memungkinkan, pengujian berikutnya dapat memindahkan titik capture langsung ke broker agar data trafik masuk ke broker terekam lebih lengkap.",
        "Penelitian lanjutan dapat membahas MQTT dengan TLS pada port 8883, autentikasi broker, atau integrasi monitoring real-time.",
        "Pengujian skala lebih besar perlu tetap dilakukan pada lingkungan yang memiliki izin dan bukan pada sistem produksi publik.",
    ])

    h2(doc, "5.3", "Penutup")
    p(doc, "Draft skripsi ini disusun sebagai dasar pengembangan naskah penelitian. Bagian hasil dan pembahasan perlu diperbarui setelah data eksperimen final tersedia. Dengan penyusunan data yang rapi, penelitian ini dapat menjelaskan hubungan antara serangan SYN flood, kondisi broker MQTT, performa client normal, dan efektivitas rate limiting secara lebih terukur.")


def references(doc):
    doc.add_page_break()
    center(doc, "DAFTAR PUSTAKA", bold=True)
    refs = [
        "Azwar, H., dan Gary, F. (2025). Implementasi Transport Layer Security dengan Algoritma AES untuk Meningkatkan Keamanan Komunikasi pada Jaringan IoT Berbasis MQTT Menggunakan ESP32. Jurnal Elektro dan Mesin Terapan, 11(1), 60-66.",
        "Anggoro, W. W., dan Widiasari, I. R. (2024). MQTT Protocol-Based ESP-32 Smarthome with Multi-sensor. Journal of Electrical, Electronic, Information, and Communication Technology.",
        "Design and Hardening of an MQTT Broker Server on AWS EC2 for IoT Communication Security. (2025). Jurnal Jartel.",
        "Simulasi Serangan Denial of Service Menggunakan Hping3 melalui Kali Linux. (2024). Journal of Internet and Software Engineering.",
        "Erlangga, M. F. E., Fahriani, N., dan Tantri, A. H. (2024). Deteksi Serangan Syn Flood pada Server Menggunakan Metode Algoritma K-Nearest Neighbor. SEMASTER: Seminar Nasional Teknologi Informasi dan Ilmu Komputer, 2(1), 68-72.",
        "Pengujian Kehandalan Jaringan terhadap TCP/SYN Flood Attack dengan Metode Simulasi Serangan Lokal.",
        "Analisis Forensik terhadap Serangan DDoS Ping of Death Menggunakan Tools NMAP dan HPING3.",
        "Implementation of Rate Limiting and Telegram Bot for HTTP GET Flood Attack Mitigation. (2025). Journal of Information System and Application Development.",
        "Implementasi Firewall Menggunakan Iptables untuk Melindungi Server dari Serangan DDoS. (2024). Journal of Internet and Software Engineering.",
        "Putri, A. F., Hadi, A., dan Rusdiana, L. (2025). Analisis Trafik Jaringan Menggunakan Wireshark untuk Deteksi Serangan Deauthentication pada Perangkat Kamera Wi-Fi. Jurnal Saintekom, 15(2), 165-176.",
        "Implementasi Packet Sniffing dengan Wireshark untuk Analisis Lalu Lintas Jaringan.",
        "Analisis Trafik Jaringan Menggunakan Wireshark untuk Monitoring Aktivitas Jaringan.",
        "OASIS. MQTT Version 5.0 Specification.",
        "The Tcpdump Group. Wireshark and Tshark Documentation.",
        "Eclipse Foundation. Eclipse Mosquitto Documentation.",
    ]
    for ref in refs:
        p(doc, ref, first_indent=False)


def appendices(doc):
    doc.add_page_break()
    center(doc, "LAMPIRAN", bold=True)
    h2(doc, "Lampiran A", "Contoh Konfigurasi lab_env.sh")
    p(doc, "TARGET_IP=192.168.100.50", first_indent=False)
    p(doc, "MQTT_PORTS=1883", first_indent=False)
    p(doc, "NET_IFACE=eth0", first_indent=False)
    p(doc, "LOG_DIR=./logs", first_indent=False)
    p(doc, "PCAP_DIR=./pcap", first_indent=False)

    h2(doc, "Lampiran B", "Contoh Alur Menjalankan Project")
    number_list(doc, [
        "Pastikan broker MQTT berjalan.",
        "Pastikan PC A dapat terhubung ke broker.",
        "Jalankan run_all_cli.sh pada PC A.",
        "Isi IP broker, interface capture, dan port MQTT.",
        "Pada skenario DoS, jalankan attacker pada mesin terpisah.",
        "Tekan CTRL+C setelah durasi pengujian selesai.",
        "Jalankan compile.py untuk menghasilkan metrics.csv.",
        "Bandingkan hasil normal, DoS, dan DoS + rate limiting.",
    ])

    h2(doc, "Lampiran C", "Catatan Revisi yang Perlu Diisi Setelah Pengujian")
    number_list(doc, [
        "Isi nama mahasiswa, NIM, pembimbing, dan format resmi kampus.",
        "Isi konfigurasi lingkungan pengujian yang benar-benar digunakan.",
        "Ganti tabel template BAB IV dengan angka hasil pengujian.",
        "Tambahkan screenshot Wireshark atau grafik hasil jika diminta pembimbing.",
        "Sesuaikan daftar pustaka dengan gaya sitasi yang diminta program studi.",
    ])


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_diagrams()
    doc = Document()
    style_doc(doc)
    cover(doc)
    front_matter(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    references(doc)
    appendices(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
