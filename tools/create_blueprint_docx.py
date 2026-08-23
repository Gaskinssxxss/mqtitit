#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "BLUEPRINT_SISTEM_LENGKAP.docx"


def add_code(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Blueprint Sistem Lengkap\nProject Analisis DoS Pada Broker MQTT")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Konsep baru: SYN flood, packet capture di broker, dan mitigasi rate limiting"
    ).italic = True

    doc.add_paragraph()
    doc.add_heading("1. Gambaran Besar Project", level=1)
    doc.add_paragraph(
        "Project ini dibuat untuk membantu penelitian tentang layanan broker MQTT pada server "
        "Universitas Mataram. Inti penelitian tidak berubah dari proposal awal, yaitu melihat "
        "dampak serangan DoS terhadap broker MQTT dan mengevaluasi rate limiting sebagai mitigasi."
    )
    doc.add_paragraph(
        "Perubahan utamanya adalah ruang lingkup dibuat lebih fokus. Jenis serangan dipersempit "
        "menjadi SYN flood, dan pengambilan data dipindahkan langsung ke server broker. Dengan "
        "cara ini, data yang dianalisis adalah trafik yang benar-benar diterima oleh broker."
    )

    doc.add_heading("2. Tujuan Sistem", level=1)
    add_bullets(
        doc,
        [
            "Menguji kondisi broker MQTT saat berjalan normal.",
            "Menguji perubahan performa broker MQTT saat menerima SYN flood.",
            "Menguji apakah rate limiting dapat mengurangi dampak SYN flood.",
            "Menghasilkan data KPI yang mudah dijelaskan: success rate, latency, dan SYN packet rate.",
            "Menyediakan mode lokal, kampus, dan publik tanpa mengubah kode utama.",
        ],
    )

    doc.add_heading("3. Konsep Dasar Yang Harus Dipahami", level=1)
    add_table(
        doc,
        ["Istilah", "Penjelasan sederhana"],
        [
            [
                "Broker MQTT",
                "Pusat penerima dan penerus pesan MQTT. Dalam project ini broker adalah objek yang diuji.",
            ],
            [
                "Perangkat IoT simulasi",
                "Client yang mengirim data normal ke broker, misalnya data suhu.",
            ],
            [
                "SYN",
                "Permintaan awal untuk membuka koneksi TCP ke server.",
            ],
            [
                "SYN flood",
                "Serangan yang mengirim banyak permintaan koneksi ke broker dalam waktu singkat.",
            ],
            [
                "Rate limiting",
                "Pembatas jumlah permintaan koneksi yang boleh masuk agar broker tidak kewalahan.",
            ],
            [
                "Packet capture",
                "Proses merekam trafik jaringan yang masuk ke broker.",
            ],
        ],
    )

    doc.add_heading("4. Masalah Pada Konsep Lama", level=1)
    doc.add_paragraph(
        "Pada konsep lama, ada host monitoring yang menjalankan sniffer dan prober terpisah dari "
        "server broker. Masalahnya, target serangan adalah broker. Jika attacker mengirim trafik "
        "langsung ke broker, host monitoring belum tentu melihat seluruh trafik itu."
    )
    doc.add_paragraph(
        "Akibatnya, data capture bisa kurang lengkap. Jika data yang dipakai analisis tidak mewakili "
        "trafik yang benar-benar diterima broker, maka hasil analisis menjadi kurang kuat."
    )

    doc.add_heading("5. Posisi Konsep Baru", level=1)
    doc.add_paragraph(
        "Pada konsep baru, server broker menjadi pusat pengamatan. Broker menjalankan Mosquitto, "
        "packet capture, dan rate limiting. Client mengirim trafik MQTT normal, sedangkan attacker "
        "mengirim SYN flood ke broker."
    )
    add_code(
        doc,
        "Windows client  -> data MQTT normal -> Server broker MQTT\n"
        "VBox attacker   -> SYN flood        -> Server broker MQTT\n\n"
        "Server broker menjalankan:\n"
        "- Mosquitto\n"
        "- packet capture\n"
        "- nftables rate limiting",
    )
    doc.add_paragraph(
        "Reasoning: karena broker adalah objek yang menerima trafik normal dan trafik serangan, "
        "maka capture yang paling masuk akal adalah capture dari sisi broker."
    )

    doc.add_heading("6. Arsitektur Fisik Untuk Demo Lokal", level=1)
    add_code(
        doc,
        "Windows laptop\n"
        "|- VBox A: Ubuntu Server broker\n"
        "|  |- Mosquitto\n"
        "|  |- tshark/tcpdump collector\n"
        "|  |- nftables rate limiting\n"
        "|\n"
        "|- VBox B: Ubuntu Server attacker\n"
        "|  |- hping3 SYN flood\n"
        "|\n"
        "|- Windows: client IoT simulasi\n"
        "   |- Python MQTT client\n"
        "   |- analisis hasil",
    )
    add_table(
        doc,
        ["Mesin", "Peran", "Yang dijalankan"],
        [
            ["VBox A", "Broker", "Mosquitto, capture, rate limiting, finalize run"],
            ["VBox B", "Attacker", "SYN flood generator"],
            ["Windows", "Client dan analisis", "MQTT test client, analyze all runs"],
        ],
    )

    doc.add_heading("7. Mode Deployment", level=1)
    add_table(
        doc,
        ["Mode", "Dipakai Untuk", "Contoh", "Catatan"],
        [
            [
                "local",
                "Demo/testbed lokal",
                "VirtualBox A/B dan Windows client",
                "Mode paling aman untuk latihan dan demo.",
            ],
            [
                "campus",
                "Broker jaringan kampus/internal/VPN",
                "Broker di jaringan Universitas Mataram",
                "Harus ada izin dari pengelola jaringan.",
            ],
            [
                "public",
                "Broker yang dapat dijangkau internet",
                "Domain/IP publik broker MQTT",
                "Harus ada izin tertulis karena melewati jaringan publik.",
            ],
        ],
    )

    doc.add_heading("8. BROKER_HOST Dan BROKER_CAPTURE_HOST", level=1)
    doc.add_paragraph(
        "Project memakai dua alamat broker agar fleksibel untuk lokal, kampus, dan publik."
    )
    add_table(
        doc,
        ["Variabel", "Dipakai Oleh", "Arti"],
        [
            [
                "BROKER_HOST",
                "Client dan attacker",
                "Alamat broker yang dipakai dari luar server broker.",
            ],
            [
                "BROKER_CAPTURE_HOST",
                "Capture dan rate limiting di server broker",
                "Alamat broker yang terlihat pada interface server broker.",
            ],
        ],
    )
    add_code(
        doc,
        "Contoh local:\n"
        "BROKER_HOST=192.168.56.10\n"
        "BROKER_CAPTURE_HOST=192.168.56.10\n\n"
        "Contoh campus/public:\n"
        "BROKER_HOST=mqtt.unram.ac.id\n"
        "BROKER_CAPTURE_HOST=10.10.10.5",
    )
    doc.add_paragraph(
        "Reasoning: pada jaringan publik atau kampus, client bisa mengakses broker lewat domain atau "
        "IP publik, tetapi server broker bisa menerima trafik pada IP internal karena NAT atau routing."
    )

    doc.add_heading("9. Skenario Pengujian", level=1)
    add_table(
        doc,
        ["Skenario", "Kondisi", "Tujuan"],
        [
            [
                "normal",
                "Client mengirim data MQTT tanpa serangan.",
                "Mengetahui kondisi dasar broker.",
            ],
            [
                "syn_flood",
                "Client mengirim data MQTT, attacker mengirim SYN flood.",
                "Melihat dampak serangan terhadap layanan MQTT.",
            ],
            [
                "syn_flood_rate_limit",
                "Client mengirim data MQTT, attacker mengirim SYN flood, rate limiting aktif.",
                "Melihat apakah rate limiting memperbaiki kondisi layanan.",
            ],
        ],
    )

    doc.add_heading("10. Mekanisme Trafik", level=1)
    doc.add_heading("10.1 Trafik Normal", level=2)
    doc.add_paragraph(
        "Windows client mengirim pesan MQTT ke broker. Isi pesannya dibuat seperti data IoT, "
        "misalnya data suhu. Data ini bukan serangan. Data ini dipakai sebagai beban normal."
    )
    add_code(doc, "Windows client -> publish data suhu -> broker MQTT")

    doc.add_heading("10.2 Trafik Serangan", level=2)
    doc.add_paragraph(
        "VBox B attacker tidak mengirim data suhu. Attacker mengirim banyak permintaan koneksi TCP "
        "ke port MQTT. Permintaan awal koneksi ini disebut SYN. Jika dikirim banyak, disebut SYN flood."
    )
    add_code(doc, "VBox B attacker -> banyak paket SYN -> broker MQTT port 1883")

    doc.add_heading("10.3 Mitigasi", level=2)
    doc.add_paragraph(
        "Pada skenario mitigasi, server broker mengaktifkan nftables untuk membatasi paket SYN "
        "yang masuk terlalu banyak. Tujuannya adalah mencegah broker kewalahan."
    )

    doc.add_heading("11. KPI Dan Cara Membacanya", level=1)
    add_table(
        doc,
        ["KPI", "Arti", "Pola Yang Diharapkan"],
        [
            [
                "mqtt_success_rate",
                "Persentase pesan MQTT yang berhasil.",
                "Normal tinggi, SYN flood bisa turun, rate limiting diharapkan naik kembali.",
            ],
            [
                "mqtt_latency_median_ms",
                "Median waktu pesan MQTT sampai diterima kembali.",
                "Normal rendah, SYN flood bisa naik, rate limiting diharapkan lebih stabil.",
            ],
            [
                "mqtt_latency_p95_ms",
                "Latency batas atas untuk mayoritas pesan.",
                "Dipakai untuk melihat lonjakan delay.",
            ],
            [
                "syn_rate_mean",
                "Rata-rata paket SYN per detik yang masuk ke broker.",
                "Harus naik saat skenario SYN flood.",
            ],
            [
                "syn_rate_peak",
                "Puncak paket SYN per detik.",
                "Menunjukkan intensitas serangan tertinggi.",
            ],
        ],
    )

    doc.add_heading("12. Struktur Project", level=1)
    add_code(
        doc,
        "01_windows_client/              client dan analisis pada Windows\n"
        "02_ubuntu_server_a_broker/      broker, capture, dan rate limiting\n"
        "03_ubuntu_server_b_attacker/    generator SYN flood\n"
        "docs/                 dokumentasi project\n"
        "legacy_old_project/   arsip kode lama\n\n"
        "Setiap folder mesin memiliki config.env.example dan README.txt sendiri.",
    )

    doc.add_heading("13. Langkah Menjalankan Mode Local", level=1)
    doc.add_paragraph(
        "Mode local dipakai untuk demo VirtualBox. Ini mode yang paling aman untuk teman yang masih belajar."
    )
    doc.add_heading("13.1 Konfigurasi VBox A Broker", level=2)
    add_code(
        doc,
        "DEPLOYMENT_MODE=local\n"
        "RUN_ID=run01_normal\n"
        "SCENARIO=normal\n"
        "BROKER_HOST=IP_VBOX_A\n"
        "BROKER_CAPTURE_HOST=IP_VBOX_A\n"
        "MQTT_PORT=1883\n"
        "BROKER_IFACE=NAMA_INTERFACE_BROKER\n"
        "EXPERIMENT_DURATION=60\n"
        "CAPTURE_DURATION=75",
    )
    add_numbered(
        doc,
        [
            "Jalankan Mosquitto di VBox A.",
            "Pada Server A, jalankan ./scripts/run_broker.sh.",
            "Setelah selesai, jalankan ./scripts/finalize_broker.sh.",
        ],
    )

    doc.add_heading("13.2 Konfigurasi Windows Client", level=2)
    add_code(
        doc,
        "Isi 01_windows_client\\config.env:\n\n"
        "DEPLOYMENT_MODE=local\n"
        "RUN_ID=run01_normal\n"
        "SCENARIO=normal\n"
        "BROKER_HOST=IP_VBOX_A\n"
        "MQTT_PORT=1883\n"
        "EXPERIMENT_DURATION=60\n\n"
        "Jalankan di PowerShell:\n"
        "powershell -ExecutionPolicy Bypass -File .\\run_client.ps1",
    )

    doc.add_heading("13.3 Konfigurasi VBox B Attacker", level=2)
    add_code(
        doc,
        "DEPLOYMENT_MODE=local\n"
        "RUN_ID=run02_syn_flood\n"
        "SCENARIO=syn_flood\n"
        "BROKER_HOST=IP_VBOX_A\n"
        "MQTT_PORT=1883\n"
        "EXPERIMENT_DURATION=60\n"
        "ATTACK_RATE=1000\n\n"
        "I_UNDERSTAND_AUTHORIZED_TESTBED=yes\n"
        "./scripts/run_attacker.sh",
    )

    doc.add_heading("14. Langkah Menjalankan Mode Campus", level=1)
    doc.add_paragraph(
        "Mode campus dipakai jika broker berada di jaringan kampus atau jaringan Universitas Mataram. "
        "Mode ini hanya boleh dipakai jika ada izin."
    )
    add_code(
        doc,
        "DEPLOYMENT_MODE=campus\n"
        "BROKER_HOST=IP_ATAU_DOMAIN_YANG_DIPAKAI_CLIENT\n"
        "BROKER_CAPTURE_HOST=IP_YANG_TERLIHAT_DI_SERVER_BROKER\n"
        "MQTT_PORT=1883",
    )
    add_bullets(
        doc,
        [
            "Capture dan rate limiting tetap dijalankan di server broker.",
            "Client menargetkan BROKER_HOST.",
            "Attacker juga menargetkan BROKER_HOST, hanya jika pengujian sudah disetujui.",
            "Untuk attacker harus menambahkan I_HAVE_WRITTEN_AUTHORIZATION=yes.",
        ],
    )
    add_code(
        doc,
        "I_UNDERSTAND_AUTHORIZED_TESTBED=yes\n"
        "I_HAVE_WRITTEN_AUTHORIZATION=yes\n"
        "./scripts/run_attacker.sh",
    )

    doc.add_heading("15. Langkah Menjalankan Mode Public", level=1)
    doc.add_paragraph(
        "Mode public dipakai jika broker benar-benar dapat diakses dari internet. Mode ini paling sensitif "
        "karena trafik serangan melewati jaringan publik. Gunakan hanya dengan izin tertulis."
    )
    add_code(
        doc,
        "DEPLOYMENT_MODE=public\n"
        "BROKER_HOST=mqtt.unram.ac.id\n"
        "BROKER_CAPTURE_HOST=IP_INTERNAL_SERVER_BROKER\n"
        "MQTT_PORT=1883",
    )
    doc.add_paragraph(
        "Reasoning: mode public disediakan agar kode siap dipakai jika penelitian benar-benar diberi "
        "akses ke broker publik. Namun untuk demo dan pengambilan data awal, mode local tetap lebih aman."
    )

    doc.add_heading("16. Urutan Praktis Per Skenario", level=1)
    add_numbered(
        doc,
        [
            "Tentukan RUN_ID dan SCENARIO.",
            "Samakan config di VBox A, VBox B, dan Windows.",
            "Jalankan broker role di VBox A.",
            "Jalankan client di Windows.",
            "Jika skenario bukan normal, jalankan attacker di VBox B.",
            "Setelah durasi selesai, finalize di VBox A.",
            "Salin hasil broker ke Windows.",
            "Gabungkan dengan mqtt_client.csv dari Windows.",
            "Pada Windows, jalankan analyze_run.ps1 untuk tiap run.",
            "Pada Windows, jalankan analyze_all.ps1 untuk membandingkan semua skenario.",
        ],
    )

    doc.add_heading("17. File Output Yang Dihasilkan", level=1)
    add_table(
        doc,
        ["File", "Asal", "Isi"],
        [
            ["mqtt_client.csv", "Windows client", "Log pesan MQTT, status, dan latency."],
            ["capture.pcapng", "VBox A broker", "Rekaman paket jaringan."],
            ["raw_flow.csv", "VBox A broker", "Hasil ekstraksi PCAP ke CSV."],
            ["metrics.json/csv", "Analyzer", "KPI per run."],
            ["summary.csv", "Analyzer semua run", "Ringkasan semua eksperimen."],
            ["comparisons.csv", "Analyzer semua run", "Mann-Whitney dan Cliff's Delta."],
            ["analysis_report.txt", "Analyzer semua run", "Laporan ringkas hasil analisis."],
        ],
    )

    doc.add_heading("18. Cara Menjelaskan Hasil", level=1)
    doc.add_paragraph(
        "Cara membaca hasil tidak perlu rumit. Cukup bandingkan tiga kondisi."
    )
    add_code(
        doc,
        "Normal:\n"
        "- success rate tinggi\n"
        "- latency rendah\n"
        "- SYN rate rendah\n\n"
        "SYN flood:\n"
        "- SYN rate naik\n"
        "- success rate turun atau latency naik\n\n"
        "SYN flood + rate limiting:\n"
        "- success rate membaik atau latency lebih stabil\n"
        "- menunjukkan mitigasi membantu broker",
    )

    doc.add_heading("19. Reasoning Desain Sistem", level=1)
    add_bullets(
        doc,
        [
            "Capture di broker dipilih karena broker adalah target yang menerima trafik.",
            "SYN flood dipilih karena sesuai dengan connection flood pada port MQTT.",
            "Rate limiting dipilih karena proposal memang membahas pembatasan trafik.",
            "Mode local dipakai untuk demo aman dan tidak mengganggu jaringan luar.",
            "Mode campus/public disiapkan agar project tetap bisa dipakai jika penelitian diberi akses ke broker sebenarnya.",
            "BROKER_HOST dan BROKER_CAPTURE_HOST dipisah agar mendukung NAT, domain publik, dan IP internal.",
            "KPI dibuat sedikit agar mudah dijelaskan dan langsung menjawab rumusan masalah.",
        ],
    )

    doc.add_heading("20. Catatan Keamanan Dan Etika", level=1)
    add_bullets(
        doc,
        [
            "Jangan menjalankan SYN flood ke IP yang bukan testbed.",
            "Jangan menyerang server publik tanpa izin tertulis.",
            "Untuk mode campus/public, minta izin dari pengelola jaringan.",
            "Gunakan mode local untuk latihan, demo, dan pengujian awal.",
            "Script attacker sudah diberi guard agar tidak berjalan tanpa konfirmasi eksplisit.",
        ],
    )

    doc.add_heading("21. Ringkasan Paling Sederhana", level=1)
    add_code(
        doc,
        "Client mengirim data MQTT normal ke broker.\n"
        "Attacker mengirim banyak permintaan koneksi ke broker.\n"
        "Broker merekam trafik yang masuk.\n"
        "Broker diberi rate limiting.\n"
        "Hasil normal, serangan, dan mitigasi dibandingkan.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
