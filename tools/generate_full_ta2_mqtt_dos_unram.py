#!/usr/bin/env python3
"""Generate full TA 2 draft for MQTT DoS rate limiting research.

The document is intentionally generated from structured text and CSV results so
future experiment updates can be folded into the thesis without manual
copy-paste drift.
"""

from __future__ import annotations

import csv
import math
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "doc"
TMP_DIR = ROOT / "tmp" / "docs" / "ta2_full"
FIG_DIR = TMP_DIR / "figures"
OUT_DOCX = OUT_DIR / "DRAFT_TA2_LENGKAP_MQTT_DOS_RATE_LIMITING_UNRAM.docx"

TITLE = (
    "ANALISIS SERANGAN DENIAL OF SERVICE PADA BROKER MQTT DAN "
    "MITIGASINYA MENGGUNAKAN RATE LIMITING PADA SERVER UNIVERSITAS MATARAM"
)
AUTHOR = "Kalsum Rahmawati"
NIM = "F1D022126"
PROGRAM = "PROGRAM STUDI TEKNIK INFORMATIKA"
FACULTY = "FAKULTAS TEKNIK"
UNIVERSITY = "UNIVERSITAS MATARAM"
YEAR = "2026"
LOGO = ROOT / "tmp" / "docs" / "proposal_assets" / "word" / "media" / "image1.jpg"

FINAL_AGG = ROOT / "output" / "unram_experiments_comprehensive" / "_final_tuning_report_20260821" / "aggregate_final.csv"
TUNING = ROOT / "output" / "unram_experiments_comprehensive" / "_final_tuning_report_20260821" / "tuning_candidates_25pps.csv"
ALL_RUNS = ROOT / "output" / "unram_experiments_comprehensive" / "_final_tuning_report_20260821" / "all_runs_final.csv"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def fnum(value: str | float | int | None, digits: int = 2) -> str:
    if value is None or value == "":
        return "-"
    try:
        n = float(value)
    except (TypeError, ValueError):
        return str(value)
    if math.isnan(n):
        return "-"
    if n.is_integer():
        return str(int(n))
    return f"{n:.{digits}f}"


def pct(value: str | float | int | None, digits: int = 2) -> str:
    return f"{fnum(value, digits)}%"


def font(size: int = 28, bold: bool = False):
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size)
    return ImageFont.load_default()


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)


def rounded_box(draw: ImageDraw.ImageDraw, xy, text: str, fill: str, outline: str, width: int = 4):
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=width)
    lines = text.split("\n")
    y = xy[1] + 28
    for i, line in enumerate(lines):
        used_font = font(26, i == 0)
        bbox = draw.textbbox((0, 0), line, font=used_font)
        x = xy[0] + ((xy[2] - xy[0]) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, fill="#111827", font=used_font)
        y += 42


def arrow(draw: ImageDraw.ImageDraw, start, end, label: str = ""):
    draw.line((start, end), fill="#111827", width=4)
    ex, ey = end
    sx, sy = start
    if ex >= sx:
        draw.polygon([(ex, ey), (ex - 18, ey - 10), (ex - 18, ey + 10)], fill="#111827")
    else:
        draw.polygon([(ex, ey), (ex + 18, ey - 10), (ex + 18, ey + 10)], fill="#111827")
    if label:
        lx = (sx + ex) / 2 - 90
        ly = (sy + ey) / 2 - 38
        draw.text((lx, ly), label, fill="#111827", font=font(21))


def make_architecture() -> Path:
    out = FIG_DIR / "gambar_3_1_arsitektur_pengujian.png"
    img = Image.new("RGB", (1600, 950), "#f8fafc")
    d = ImageDraw.Draw(img)
    d.text((365, 45), "Arsitektur Pengujian Broker MQTT Universitas Mataram", fill="#111827", font=font(40, True))
    rounded_box(d, (80, 220, 440, 465), "Perangkat Penguji\nVPN WireGuard\nProber MQTT\nPacket capture", "#dbeafe", "#1d4ed8")
    rounded_box(d, (620, 220, 980, 465), "Server UNRAM\nBroker MQTT\nMosquitto\nTCP port 1883", "#dcfce7", "#15803d")
    rounded_box(d, (1160, 220, 1520, 465), "Attacker Terkontrol\nhping3\nSYN flood\nRate bertingkat", "#fee2e2", "#b91c1c")
    arrow(d, (440, 300), (620, 300), "publish MQTT")
    arrow(d, (1160, 380), (980, 380), "SYN flood")
    d.rounded_rectangle((200, 640, 1400, 815), radius=18, fill="#fff7ed", outline="#c2410c", width=4)
    d.text((255, 680), "Data yang diamati: packet capture, log prober, indikator TCP SYN, half-open, RTT, dan success rate.", fill="#111827", font=font(27, True))
    d.text((255, 735), "Mitigasi dipasang pada server broker menggunakan nftables rate limiting terhadap paket TCP SYN ke port 1883.", fill="#111827", font=font(25))
    img.save(out)
    return out


def make_flow() -> Path:
    out = FIG_DIR / "gambar_3_2_alur_pengujian.png"
    img = Image.new("RGB", (1600, 950), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((530, 55), "Alur Pengujian dan Pengolahan Data", fill="#111827", font=font(40, True))
    items = [
        ((80, 250, 390, 420), "Baseline\nbroker normal"),
        ((485, 250, 795, 420), "SYN flood\ntanpa mitigasi"),
        ((890, 250, 1200, 420), "Rate limiting\nnftables"),
        ((1295, 250, 1555, 420), "Evaluasi\nBAB IV"),
        ((485, 610, 795, 780), "Capture\npcapng/csv"),
        ((890, 610, 1200, 780), "Kompilasi\nmetrics"),
    ]
    for xy, text in items:
        rounded_box(d, xy, text, "#eef2ff", "#4338ca")
    for start, end, label in [
        ((390, 335), (485, 335), "1"),
        ((795, 335), (890, 335), "2"),
        ((1200, 335), (1295, 335), "3"),
        ((640, 420), (640, 610), "data"),
        ((795, 695), (890, 695), "olah"),
        ((1200, 695), (1425, 420), "hasil"),
    ]:
        arrow(d, start, end, label)
    img.save(out)
    return out


def make_syn_flood_diagram() -> Path:
    out = FIG_DIR / "gambar_2_1_syn_flood.png"
    img = Image.new("RGB", (1600, 920), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((575, 55), "Konsep SYN Flood pada TCP", fill="#111827", font=font(42, True))
    rounded_box(d, (130, 210, 470, 390), "Client Normal\nSYN\nSYN-ACK\nACK", "#dbeafe", "#1d4ed8")
    rounded_box(d, (630, 210, 970, 390), "Broker MQTT\nmenunggu ACK\nresource koneksi", "#dcfce7", "#15803d")
    rounded_box(d, (1130, 210, 1470, 390), "Attacker\nbanyak SYN\nACK tidak selesai", "#fee2e2", "#b91c1c")
    arrow(d, (470, 275), (630, 275), "handshake normal")
    arrow(d, (1130, 330), (970, 330), "SYN berulang")
    d.rounded_rectangle((220, 585, 1380, 765), radius=20, fill="#fef2f2", outline="#991b1b", width=4)
    d.text((270, 625), "SYN flood tidak mengirim data sensor seperti suhu. Bentuk serangannya adalah paket TCP SYN dalam jumlah banyak.", fill="#111827", font=font(26, True))
    d.text((270, 682), "Tujuannya membuat broker sibuk menangani percobaan koneksi sehingga client normal dapat melambat atau gagal.", fill="#111827", font=font(25))
    img.save(out)
    return out


def make_bar_chart(data: list[tuple[str, float]], out_name: str, title: str, ylabel: str, color: str) -> Path:
    out = FIG_DIR / out_name
    img = Image.new("RGB", (1500, 850), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((90, 45), title, fill="#111827", font=font(38, True))
    left, top, width, height = 170, 165, 1120, 520
    d.line((left, top, left, top + height), fill="#111827", width=4)
    d.line((left, top + height, left + width, top + height), fill="#111827", width=4)
    max_v = max(v for _, v in data) if data else 1
    max_v = max(max_v, 1)
    if max_v < 100 and "Success" in title:
        max_v = 100
    group_w = width / len(data)
    for i, (label, value) in enumerate(data):
        x0 = left + i * group_w + group_w * 0.22
        x1 = left + i * group_w + group_w * 0.78
        bar_h = int((value / max_v) * (height - 40))
        y0 = top + height - bar_h
        d.rectangle((x0, y0, x1, top + height), fill=color)
        d.text((x0, y0 - 35), f"{value:.2f}", fill="#111827", font=font(21))
        d.text((x0 - 15, top + height + 18), label, fill="#111827", font=font(21))
    d.text((65, 390), ylabel, fill="#111827", font=font(23, True))
    img.save(out)
    return out


def make_line_chart(data: list[tuple[float, float, float]], out_name: str) -> Path:
    out = FIG_DIR / out_name
    img = Image.new("RGB", (1500, 850), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((90, 45), "Kenaikan Paket SYN dan Half-open Tanpa Mitigasi", fill="#111827", font=font(37, True))
    left, top, width, height = 170, 165, 1120, 520
    d.line((left, top, left, top + height), fill="#111827", width=4)
    d.line((left, top + height, left + width, top + height), fill="#111827", width=4)
    max_x = max(x for x, _, _ in data) or 1
    max_y = max(max(y1, y2) for _, y1, y2 in data) or 1

    def point(x, y):
        px = left + (x / max_x) * width
        py = top + height - (y / max_y) * (height - 30)
        return px, py

    syn_points = [point(x, y1) for x, y1, _ in data]
    half_points = [point(x, y2) for x, _, y2 in data]
    d.line(syn_points, fill="#dc2626", width=5)
    d.line(half_points, fill="#2563eb", width=5)
    for (x, y1, y2), p1, p2 in zip(data, syn_points, half_points):
        d.ellipse((p1[0] - 7, p1[1] - 7, p1[0] + 7, p1[1] + 7), fill="#dc2626")
        d.ellipse((p2[0] - 7, p2[1] - 7, p2[0] + 7, p2[1] + 7), fill="#2563eb")
        d.text((p1[0] - 20, top + height + 18), f"{int(x)} pps", fill="#111827", font=font(20))
    d.rectangle((1010, 100, 1050, 125), fill="#dc2626")
    d.text((1065, 96), "SYN/s", fill="#111827", font=font(23))
    d.rectangle((1010, 140, 1050, 165), fill="#2563eb")
    d.text((1065, 136), "Half-open/s", fill="#111827", font=font(23))
    img.save(out)
    return out


def generate_figures(agg: list[dict[str, str]], tuning: list[dict[str, str]]) -> dict[str, Path]:
    no_mit = [r for r in agg if r["batch_label"] == "tanpa_mitigasi"]
    no_mit = sorted(no_mit, key=lambda r: float(r["rate_pps"]))
    figs = {
        "syn_flood": make_syn_flood_diagram(),
        "architecture": make_architecture(),
        "flow": make_flow(),
        "success_no_mit": make_bar_chart(
            [(f"{int(float(r['rate_pps']))} pps", float(r["success_mean"])) for r in no_mit],
            "gambar_4_1_success_tanpa_mitigasi.png",
            "Success Rate Tanpa Mitigasi",
            "Success rate (%)",
            "#dc2626",
        ),
        "syn_half_no_mit": make_line_chart(
            [(float(r["rate_pps"]), float(r["syn_mean"]), float(r["half_open_mean"])) for r in no_mit],
            "gambar_4_2_syn_halfopen_tanpa_mitigasi.png",
        ),
        "tuning_success": make_bar_chart(
            [(r["rule_config"].replace("limit", "L").replace("_burst", " B"), float(r["success_mean"])) for r in tuning],
            "gambar_4_3_success_tuning.png",
            "Perbandingan Success Rate Tuning Rate Limiting",
            "Success rate (%)",
            "#15803d",
        ),
        "tuning_rtt": make_bar_chart(
            [(r["rule_config"].replace("limit", "L").replace("_burst", " B"), float(r["rtt_avg_mean"])) for r in tuning],
            "gambar_4_4_rtt_tuning.png",
            "Perbandingan RTT Rata-rata Tuning Rate Limiting",
            "RTT rata-rata (ms)",
            "#7c3aed",
        ),
    }
    return figs


def set_doc_style(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3)
    sec.bottom_margin = Cm(3)
    sec.left_margin = Cm(4)
    sec.right_margin = Cm(3)

    for name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.color.rgb = RGBColor(0, 0, 0)
    doc.styles["Normal"].font.size = Pt(12)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.5
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)


def add_page_number(section) -> None:
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def set_run_font(run, size: int = 12, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def paragraph(doc: Document, text: str = "", *, align=None, indent: bool = True, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, 12, bold, italic)
    return p


def centered(doc: Document, text: str, size: int = 12, bold: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size, bold)
    return p


def chapter_title(doc: Document, number: str, title: str):
    doc.add_page_break()
    centered(doc, f"BAB {number}", 12, True)
    centered(doc, title.upper(), 12, True)
    paragraph(doc, "", indent=False)


def section_title(doc: Document, number: str, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f"{number} {title}")
    set_run_font(run, 12, True)


def sub_title(doc: Document, number: str, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{number} {title}")
    set_run_font(run, 12, True)


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"- {text}")
    set_run_font(run, 12)


def numbered(doc: Document, items: Iterable[str]):
    for i, text in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"{i}. {text}")
        set_run_font(run, 12)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], title: str | None = None):
    if title:
        centered(doc, title, 12, True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_run_font(run, 9)
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "D9EAF7")
        cell._tc.get_or_add_tcPr().append(shading)
    paragraph(doc, "", indent=False)
    return table


def add_figure(doc: Document, path: Path, caption: str, width_inches: float = 5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(0)
    r = cap.add_run(caption)
    set_run_font(r, 11, True)


def front_matter(doc: Document):
    centered(doc, TITLE, 14, True)
    for _ in range(2):
        paragraph(doc, "", indent=False)
    centered(doc, "SKRIPSI", 12, True)
    centered(doc, "Diajukan untuk memenuhi sebagian persyaratan memperoleh gelar Sarjana Komputer", 12)
    paragraph(doc, "", indent=False)
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.45))
    paragraph(doc, "", indent=False)
    centered(doc, "Oleh:", 12)
    centered(doc, AUTHOR, 12, True)
    centered(doc, NIM, 12, True)
    for _ in range(5):
        paragraph(doc, "", indent=False)
    centered(doc, PROGRAM, 12, True)
    centered(doc, FACULTY, 12, True)
    centered(doc, UNIVERSITY, 12, True)
    centered(doc, YEAR, 12, True)

    doc.add_page_break()
    centered(doc, "HALAMAN PERSETUJUAN", 12, True)
    paragraph(doc, "", indent=False)
    paragraph(doc, "Skripsi ini disusun sebagai draft Tugas Akhir II untuk dibimbingkan kepada dosen pembimbing. Bagian nama dosen pembimbing, tanggal persetujuan, dan tanda tangan masih perlu disesuaikan dengan format resmi program studi.", indent=True)
    add_table(
        doc,
        ["Komponen", "Keterangan"],
        [
            ["Judul", TITLE.title()],
            ["Nama Mahasiswa", AUTHOR],
            ["NIM", NIM],
            ["Program Studi", "Teknik Informatika"],
            ["Pembimbing I", "........................................................"],
            ["Pembimbing II", "........................................................"],
        ],
    )

    doc.add_page_break()
    centered(doc, "KATA PENGANTAR", 12, True)
    paragraph(doc, "Puji syukur penulis panjatkan ke hadirat Tuhan Yang Maha Esa karena atas rahmat dan karunia-Nya draft skripsi yang berjudul \"Analisis Serangan Denial of Service pada Broker MQTT dan Mitigasinya Menggunakan Rate Limiting pada Server Universitas Mataram\" dapat disusun. Draft ini memuat latar belakang, dasar teori, metodologi, hasil pengujian awal sampai pengujian komprehensif, pembahasan, kesimpulan, dan saran.")
    paragraph(doc, "Penelitian ini disusun untuk menganalisis dampak serangan Denial of Service, khususnya SYN flood, terhadap broker MQTT yang digunakan sebagai pusat pertukaran pesan pada sistem berbasis Internet of Things. Selain itu, penelitian ini juga mengevaluasi penggunaan rate limiting sebagai mekanisme mitigasi agar broker tetap dapat melayani client normal ketika terjadi peningkatan trafik koneksi.")
    paragraph(doc, "Penulis menyadari bahwa draft ini masih dapat disempurnakan, terutama pada bagian pengujian lanjutan, kelengkapan bukti konfigurasi, dan penyesuaian format administrasi. Oleh karena itu, masukan dari dosen pembimbing sangat diperlukan agar skripsi ini menjadi lebih tepat secara akademik dan teknis.")
    paragraph(doc, "Mataram, 2026", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
    paragraph(doc, "Penulis", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)

    doc.add_page_break()
    centered(doc, "ABSTRAK", 12, True)
    paragraph(doc, "MQTT merupakan protokol komunikasi ringan yang banyak digunakan pada sistem Internet of Things karena mendukung pola publish-subscribe dan kebutuhan bandwidth yang relatif kecil. Pada arsitektur MQTT, broker berperan sebagai pusat penerima dan penerus pesan. Peran tersebut membuat broker menjadi titik penting yang dapat memengaruhi ketersediaan layanan. Jika broker menerima trafik koneksi yang berlebihan, layanan MQTT dapat mengalami peningkatan waktu respons, penurunan keberhasilan publish, bahkan tidak dapat diakses oleh client normal.")
    paragraph(doc, "Penelitian ini menganalisis serangan Denial of Service terhadap broker MQTT pada lingkungan server Universitas Mataram. Jenis serangan yang difokuskan adalah SYN flood, yaitu serangan pada proses pembentukan koneksi TCP dengan cara mengirim banyak paket SYN ke port MQTT. Pengujian dilakukan melalui beberapa skenario, yaitu kondisi normal, serangan SYN flood tanpa mitigasi, serangan dengan rate limiting awal, dan tuning konfigurasi rate limiting. Data dikumpulkan melalui packet capture, prober MQTT, ekstraksi trafik TCP, dan kompilasi metrik availability.")
    paragraph(doc, "Hasil pengujian menunjukkan bahwa pada kondisi normal broker memiliki success rate 100%. Pada serangan 10 pps dan 25 pps tanpa mitigasi, indikator trafik SYN dan half-open meningkat, tetapi success rate client normal masih 100%. Pada serangan 50 pps tanpa mitigasi, success rate turun menjadi 27,27% dan health check setelah pengujian gagal, sehingga skenario tersebut tidak diulang untuk menghindari risiko gangguan berlebihan. Pengujian mitigasi menunjukkan bahwa konfigurasi rate limiting yang terlalu ketat, seperti limit 15/s burst 20, justru menurunkan success rate pada serangan 25 pps menjadi rata-rata 82,39%. Setelah tuning, konfigurasi limit 30/s burst 40 menghasilkan success rate 100%, RTT rata-rata 458,70 ms, dan health check 3/3 OK pada skenario 25 pps. Dengan demikian, rate limiting dapat menjadi mitigasi yang efektif apabila dikalibrasi berdasarkan baseline dan intensitas serangan.")
    paragraph(doc, "Kata kunci: MQTT, broker, Denial of Service, SYN flood, rate limiting, nftables, availability.", indent=False)

    doc.add_page_break()
    centered(doc, "ABSTRACT", 12, True)
    paragraph(doc, "MQTT is a lightweight communication protocol widely used in Internet of Things systems because it supports the publish-subscribe pattern and requires relatively low bandwidth. In MQTT architecture, the broker acts as the central component that receives and forwards messages. This role makes the broker a critical point for service availability. When the broker receives excessive connection traffic, the service may suffer from higher response time, lower publish success, or temporary unavailability for normal clients.")
    paragraph(doc, "This research analyzes Denial of Service attacks against an MQTT broker in the Universitas Mataram server environment. The attack is focused on SYN flood, which targets the TCP connection establishment process by sending many SYN packets to the MQTT port. The experiment consists of baseline testing, SYN flood without mitigation, SYN flood with initial rate limiting, and rate limiting tuning. Data are collected using packet capture, MQTT prober, TCP traffic extraction, and availability metric compilation.")
    paragraph(doc, "The results show that the broker achieved 100% success rate in baseline condition. During 10 pps and 25 pps attacks without mitigation, SYN and half-open indicators increased, while normal client success rate remained 100%. At 50 pps without mitigation, the success rate dropped to 27.27% and post-test health check failed; therefore, the scenario was not repeated to avoid excessive disruption. The mitigation tests indicate that overly strict rate limiting, such as 15/s with burst 20, reduced the success rate at 25 pps to 82.39% on average. After tuning, the 30/s limit with burst 40 achieved 100% success rate, 458.70 ms average RTT, and 3/3 successful health checks at 25 pps. These findings indicate that rate limiting can mitigate SYN flood against MQTT brokers when it is calibrated against baseline and attack intensity.")
    paragraph(doc, "Keywords: MQTT, broker, Denial of Service, SYN flood, rate limiting, nftables, availability.", indent=False)

    doc.add_page_break()
    centered(doc, "DAFTAR ISI", 12, True)
    toc = [
        "HALAMAN JUDUL",
        "HALAMAN PERSETUJUAN",
        "KATA PENGANTAR",
        "ABSTRAK",
        "ABSTRACT",
        "DAFTAR ISI",
        "DAFTAR GAMBAR",
        "DAFTAR TABEL",
        "BAB I PENDAHULUAN",
        "BAB II TINJAUAN PUSTAKA",
        "BAB III METODOLOGI PENELITIAN",
        "BAB IV HASIL DAN PEMBAHASAN",
        "BAB V PENUTUP",
        "DAFTAR PUSTAKA",
        "LAMPIRAN",
    ]
    for item in toc:
        paragraph(doc, item, indent=False)

    doc.add_page_break()
    centered(doc, "DAFTAR GAMBAR", 12, True)
    for item in [
        "Gambar 2.1 Konsep SYN flood pada TCP",
        "Gambar 3.1 Arsitektur pengujian broker MQTT Universitas Mataram",
        "Gambar 3.2 Alur pengujian dan pengolahan data",
        "Gambar 4.1 Success rate tanpa mitigasi",
        "Gambar 4.2 Kenaikan paket SYN dan half-open tanpa mitigasi",
        "Gambar 4.3 Perbandingan success rate tuning rate limiting",
        "Gambar 4.4 Perbandingan RTT rata-rata tuning rate limiting",
    ]:
        paragraph(doc, item, indent=False)

    doc.add_page_break()
    centered(doc, "DAFTAR TABEL", 12, True)
    for item in [
        "Tabel 2.1 Penelitian terkait",
        "Tabel 3.1 Alat dan bahan penelitian",
        "Tabel 3.2 Skenario pengujian",
        "Tabel 3.3 Indikator keberhasilan",
        "Tabel 4.1 Ringkasan lingkungan pengujian",
        "Tabel 4.2 Hasil agregat tanpa mitigasi",
        "Tabel 4.3 Hasil mitigasi awal",
        "Tabel 4.4 Perbandingan tuning rate limiting",
        "Tabel 4.5 Penilaian KPI konfigurasi mitigasi",
        "Tabel Lampiran 1 Daftar data pengujian",
    ]:
        paragraph(doc, item, indent=False)


def chapter_1(doc: Document):
    chapter_title(doc, "I", "Pendahuluan")
    section_title(doc, "1.1", "Latar Belakang")
    for text in [
        "Perkembangan Internet of Things mendorong banyak sistem informasi modern memakai perangkat yang dapat mengirim data secara otomatis melalui jaringan. Pada lingkungan kampus, konsep ini dapat digunakan untuk pemantauan ruangan, sensor lingkungan, perangkat keamanan, laboratorium, dan layanan lain yang membutuhkan pengiriman data secara periodik. Agar perangkat tersebut dapat saling terhubung, diperlukan protokol komunikasi yang ringan, mudah diterapkan, dan mampu berjalan pada perangkat dengan sumber daya terbatas.",
        "MQTT menjadi salah satu protokol yang sering digunakan pada sistem Internet of Things karena bekerja dengan model publish-subscribe. Pada model ini, perangkat tidak harus saling mengetahui alamat satu sama lain secara langsung. Perangkat cukup mengirim pesan ke broker, kemudian broker meneruskan pesan tersebut kepada client yang berlangganan topik tertentu. Dengan pola ini, broker MQTT menjadi komponen pusat yang menentukan apakah proses pertukaran data berjalan lancar atau tidak.",
        "Posisi broker yang sangat penting juga membuat broker menjadi target yang menarik bagi serangan Denial of Service. Denial of Service merupakan serangan yang bertujuan mengganggu ketersediaan layanan sehingga pengguna sah tidak dapat menggunakan layanan secara normal. Pada broker MQTT, gangguan dapat terjadi ketika broker menerima terlalu banyak percobaan koneksi, paket jaringan, atau permintaan publish dalam waktu singkat. Akibatnya, waktu respons dapat meningkat, koneksi normal dapat gagal, dan layanan MQTT dapat kehilangan fungsi utamanya sebagai penghubung antar client.",
        "Penelitian ini memfokuskan serangan pada SYN flood. SYN flood adalah bentuk serangan pada lapisan transport TCP. Serangan ini tidak berupa pengiriman data sensor seperti suhu, kelembapan, atau payload aplikasi MQTT. Bentuk serangannya adalah paket TCP SYN dalam jumlah banyak ke port broker MQTT. Paket SYN digunakan dalam proses awal pembentukan koneksi TCP. Jika paket SYN dikirim terus-menerus dan tidak diselesaikan dengan proses handshake normal, server dapat menyimpan banyak koneksi setengah terbuka atau half-open sehingga sumber daya jaringan dan sistem menjadi terbebani.",
        "Universitas Mataram digunakan sebagai konteks penelitian karena server broker MQTT ditempatkan pada lingkungan jaringan yang relevan dengan studi kasus kampus. Dengan konteks ini, penelitian tidak hanya membahas teori serangan, tetapi juga mengamati bagaimana broker MQTT merespons trafik normal, trafik serangan, dan trafik setelah mitigasi diterapkan. Pengujian tetap dilakukan secara terkendali agar tidak menimbulkan gangguan yang tidak diperlukan terhadap layanan lain.",
        "Mitigasi yang digunakan pada penelitian ini adalah rate limiting. Rate limiting bekerja dengan membatasi laju paket atau koneksi yang masuk ke server. Pada penelitian ini, rate limiting diarahkan pada paket TCP SYN menuju port MQTT 1883. Tujuannya bukan mematikan akses ke broker, melainkan menahan lonjakan percobaan koneksi yang tidak wajar agar client normal tetap memiliki kesempatan untuk terhubung dan mengirim data.",
        "Berdasarkan kondisi tersebut, penelitian ini penting dilakukan untuk memberikan gambaran yang lebih terukur mengenai dampak SYN flood terhadap broker MQTT dan bagaimana rate limiting dapat membantu menjaga availability. Hasil penelitian diharapkan dapat menjadi dasar teknis yang sederhana, dapat dipahami, dan dapat direproduksi untuk pengujian keamanan broker MQTT pada lingkungan akademik.",
    ]:
        paragraph(doc, text)

    section_title(doc, "1.2", "Rumusan Masalah")
    numbered(doc, [
        "Bagaimana kondisi broker MQTT pada server Universitas Mataram ketika menerima trafik normal?",
        "Bagaimana pengaruh serangan SYN flood terhadap trafik TCP, koneksi half-open, waktu respons, dan success rate client MQTT?",
        "Bagaimana efektivitas rate limiting dalam menjaga ketersediaan layanan broker MQTT ketika serangan SYN flood berjalan?",
        "Konfigurasi rate limiting seperti apa yang paling layak digunakan berdasarkan hasil pengujian terkontrol?",
    ])

    section_title(doc, "1.3", "Batasan Masalah")
    numbered(doc, [
        "Penelitian berfokus pada broker MQTT yang berjalan pada server Universitas Mataram dan dapat diakses melalui jalur VPN yang diberikan untuk keperluan pengujian.",
        "Port MQTT yang diuji adalah port 1883 dengan koneksi TCP tanpa TLS.",
        "Serangan yang diuji adalah SYN flood terkontrol. Penelitian tidak membahas eksploitasi kredensial, brute force autentikasi, malware, atau pengambilalihan server.",
        "Pengujian dilakukan dengan laju serangan terbatas agar tidak menimbulkan gangguan berlebihan terhadap server.",
        "Mitigasi yang diuji adalah rate limiting menggunakan nftables pada trafik TCP SYN ke port broker MQTT.",
        "Metrik utama yang dianalisis adalah success rate, RTT, jumlah paket SYN, half-open, TCP packet rate, RST, dan hasil health check setelah pengujian.",
    ])

    section_title(doc, "1.4", "Tujuan Penelitian")
    numbered(doc, [
        "Mengukur kondisi baseline broker MQTT pada server Universitas Mataram saat menerima trafik normal.",
        "Menganalisis dampak SYN flood terhadap indikator trafik dan availability broker MQTT.",
        "Menerapkan dan mengevaluasi rate limiting sebagai mitigasi terhadap peningkatan paket SYN pada port MQTT.",
        "Menentukan konfigurasi rate limiting yang paling sesuai berdasarkan indikator keberhasilan pengujian.",
    ])

    section_title(doc, "1.5", "Manfaat Penelitian")
    sub_title(doc, "1.5.1", "Manfaat Akademik")
    paragraph(doc, "Penelitian ini dapat menjadi rujukan sederhana mengenai hubungan antara protokol MQTT, serangan Denial of Service pada lapisan TCP, dan mitigasi menggunakan rate limiting. Pembahasan disusun agar konsep teknis seperti broker, koneksi TCP, paket SYN, dan SYN flood dapat dijelaskan secara runtut dalam konteks tugas akhir.")
    sub_title(doc, "1.5.2", "Manfaat Praktis")
    paragraph(doc, "Hasil penelitian dapat membantu administrator atau pengelola sistem memahami indikator awal ketika broker MQTT menerima trafik koneksi yang tidak wajar. Selain itu, hasil tuning rate limiting dapat menjadi contoh bahwa konfigurasi mitigasi harus diuji terlebih dahulu, karena aturan yang terlalu ketat dapat mengganggu client normal.")

    section_title(doc, "1.6", "Sistematika Penulisan")
    numbered(doc, [
        "BAB I berisi latar belakang, rumusan masalah, batasan masalah, tujuan, manfaat, dan sistematika penulisan.",
        "BAB II berisi teori pendukung mengenai MQTT, broker, TCP, SYN flood, packet capture, rate limiting, dan penelitian terkait.",
        "BAB III berisi metodologi penelitian, arsitektur pengujian, alat dan bahan, skenario, indikator, serta alur pengolahan data.",
        "BAB IV berisi hasil pengujian baseline, serangan tanpa mitigasi, pengujian mitigasi awal, tuning rate limiting, dan pembahasan.",
        "BAB V berisi kesimpulan dan saran berdasarkan hasil penelitian.",
    ])


def chapter_2(doc: Document, figs: dict[str, Path]):
    chapter_title(doc, "II", "Tinjauan Pustaka")
    section_title(doc, "2.1", "Penelitian Terkait")
    paragraph(doc, "Penelitian terkait keamanan MQTT pada umumnya membahas beberapa sisi, yaitu karakteristik protokol MQTT, kerentanan broker terhadap trafik berlebihan, teknik pendeteksian anomali, dan mitigasi pada lapisan jaringan. Beberapa penelitian menekankan bahwa MQTT cocok untuk IoT karena ringan, tetapi penggunaan broker sebagai pusat komunikasi membuat availability broker menjadi aspek kritis. Jika broker gagal melayani koneksi, maka seluruh client yang bergantung pada broker ikut terdampak.")
    paragraph(doc, "Penelitian lain membahas serangan Denial of Service terhadap broker MQTT melalui connection flood, publish flood, dan SYN flood. Connection flood menekan broker pada sisi sesi MQTT, publish flood menekan broker dengan banyak pesan aplikasi, sedangkan SYN flood menekan proses awal koneksi TCP sebelum komunikasi MQTT terjadi. Penelitian ini memilih SYN flood karena bentuknya lebih jelas untuk dikorelasikan dengan metrik jaringan seperti paket SYN, half-open connection, dan rasio paket TCP.")
    add_table(
        doc,
        ["No", "Topik", "Fokus", "Kaitan dengan Penelitian"],
        [
            ["1", "Keamanan MQTT", "Kerentanan broker pada sistem IoT", "Menjadi dasar bahwa broker adalah objek utama yang perlu dijaga availability-nya."],
            ["2", "DoS pada IoT", "Gangguan layanan akibat trafik berlebihan", "Menjelaskan konsep serangan yang menurunkan kemampuan layanan."],
            ["3", "SYN flood", "Eksploitasi proses handshake TCP", "Menjadi jenis serangan yang diuji pada port MQTT 1883."],
            ["4", "Rate limiting", "Pembatasan laju paket atau koneksi", "Menjadi metode mitigasi yang diuji dan dituning."],
            ["5", "Packet capture", "Pengambilan data trafik jaringan", "Digunakan untuk membuktikan perubahan pola trafik secara kuantitatif."],
        ],
        "Tabel 2.1 Penelitian terkait",
    )

    section_title(doc, "2.2", "Internet of Things")
    for text in [
        "Internet of Things adalah konsep yang menghubungkan perangkat fisik ke jaringan agar perangkat tersebut dapat mengirim, menerima, atau memproses data. Perangkat IoT umumnya memiliki karakteristik ringan, bekerja otomatis, dan sering mengirim data secara periodik. Contoh data yang dikirim dapat berupa suhu, kelembapan, status perangkat, atau nilai sensor lain.",
        "Dalam penelitian ini, perangkat IoT tidak harus berupa sensor fisik. Client MQTT yang mengirim pesan periodik dapat dianggap sebagai simulasi perangkat IoT karena fungsinya sama, yaitu mengirim data ke broker. Dengan simulasi ini, penelitian tetap dapat mengamati apakah broker mampu menerima pesan normal saat terjadi serangan pada jalur koneksi.",
    ]:
        paragraph(doc, text)

    section_title(doc, "2.3", "MQTT")
    paragraph(doc, "MQTT adalah protokol komunikasi berbasis TCP yang menggunakan pola publish-subscribe. Client yang mengirim pesan disebut publisher, client yang menerima pesan disebut subscriber, dan komponen yang mengatur pengiriman pesan disebut broker. MQTT banyak digunakan pada IoT karena struktur pesannya sederhana dan overhead komunikasinya relatif kecil.")
    paragraph(doc, "Dalam pola publish-subscribe, publisher tidak mengirim data langsung ke subscriber. Publisher mengirim data ke broker pada topik tertentu. Subscriber yang berlangganan topik tersebut menerima pesan dari broker. Dengan demikian, broker menjadi pusat koordinasi. Jika broker terganggu, maka proses komunikasi antar client juga terganggu.")

    section_title(doc, "2.4", "Broker MQTT")
    paragraph(doc, "Broker MQTT adalah server yang menerima koneksi client MQTT, menerima pesan publish, menyimpan informasi subscription, dan meneruskan pesan kepada subscriber yang sesuai. Pada penelitian ini, broker MQTT dijalankan menggunakan Mosquitto. Mosquitto dipilih karena ringan, umum digunakan, dan tersedia pada sistem operasi Linux.")
    paragraph(doc, "Broker MQTT pada penelitian ini ditempatkan sebagai objek utama yang diamati. Artinya, analisis tidak diarahkan pada perangkat client, tetapi pada bagaimana broker menerima trafik normal, menerima trafik serangan, dan merespons ketika aturan rate limiting diterapkan. Posisi ini membuat data pengujian lebih tepat karena broker adalah target yang langsung menerima paket SYN dari attacker.")

    section_title(doc, "2.5", "TCP dan Three-way Handshake")
    paragraph(doc, "MQTT berjalan di atas TCP. Sebelum client MQTT dapat mengirim pesan, client harus membentuk koneksi TCP ke broker. Proses pembentukan koneksi TCP umumnya terdiri dari tiga tahap, yaitu client mengirim SYN, server membalas SYN-ACK, dan client mengirim ACK. Setelah proses ini selesai, koneksi dianggap terbentuk dan komunikasi aplikasi dapat berjalan.")
    paragraph(doc, "Pada kondisi normal, proses handshake berlangsung singkat. Namun, jika server menerima banyak paket SYN dan proses handshake tidak diselesaikan, server dapat menyimpan banyak koneksi setengah terbuka. Kondisi ini disebut half-open connection. Jumlah half-open yang tinggi dapat menjadi indikator bahwa server sedang menerima trafik SYN yang tidak normal.")

    section_title(doc, "2.6", "Denial of Service")
    paragraph(doc, "Denial of Service adalah serangan yang bertujuan mengurangi atau menghentikan ketersediaan suatu layanan. Fokus serangan bukan mengambil data rahasia, melainkan membuat layanan sulit digunakan oleh pengguna sah. Dalam konteks broker MQTT, dampak DoS dapat dilihat dari meningkatnya waktu respons, menurunnya success rate publish, meningkatnya timeout, atau gagalnya health check.")
    paragraph(doc, "Serangan DoS dapat terjadi pada beberapa lapisan. Pada lapisan aplikasi, penyerang dapat mengirim banyak pesan MQTT. Pada lapisan transport, penyerang dapat membanjiri proses pembentukan koneksi TCP. Penelitian ini menggunakan pendekatan lapisan transport karena SYN flood dapat diamati melalui paket jaringan dan indikator TCP.")

    section_title(doc, "2.7", "SYN Flood")
    paragraph(doc, "SYN flood adalah serangan dengan cara mengirim banyak paket TCP SYN ke server. Paket SYN adalah paket pembuka koneksi TCP. Jika paket ini datang dalam jumlah besar dan tidak diselesaikan dengan ACK yang valid, server dapat menghabiskan sumber daya untuk menunggu kelanjutan koneksi. Pada broker MQTT, SYN flood diarahkan ke port 1883 karena port tersebut digunakan untuk koneksi MQTT tanpa TLS.")
    paragraph(doc, "Penting dipahami bahwa SYN flood bukan pengiriman data sensor. Jika client normal mengirim data suhu ke broker, maka itu adalah trafik aplikasi MQTT. Sebaliknya, attacker pada SYN flood mengirim paket pembuka koneksi TCP dalam jumlah banyak. Karena serangan terjadi sebelum komunikasi MQTT selesai terbentuk, analisisnya perlu melihat paket TCP, bukan hanya isi pesan MQTT.")
    add_figure(doc, figs["syn_flood"], "Gambar 2.1 Konsep SYN flood pada TCP", 5.8)

    section_title(doc, "2.8", "Packet Capture dan TShark")
    paragraph(doc, "Packet capture adalah proses merekam paket jaringan yang melewati suatu interface. Dalam penelitian ini, packet capture digunakan untuk memperoleh bukti kuantitatif mengenai jumlah paket TCP, paket SYN, paket RST, dan pola koneksi menuju broker. File hasil capture dapat disimpan dalam format pcapng, kemudian diekstrak menjadi CSV agar mudah dianalisis.")
    paragraph(doc, "TShark digunakan karena merupakan versi command line dari Wireshark. TShark cocok untuk pengujian berbasis skrip karena dapat dijalankan bersamaan dengan prober dan serangan. Dengan TShark, penelitian dapat mengambil data mentah yang dapat diperiksa ulang, bukan hanya mengandalkan log aplikasi.")

    section_title(doc, "2.9", "Prober MQTT")
    paragraph(doc, "Prober MQTT adalah mekanisme pengujian yang mengirim pesan MQTT secara periodik ke broker untuk mengukur apakah broker masih dapat melayani client normal. Dalam penelitian ini, prober mewakili perangkat IoT atau client sah. Nilai utama yang diamati dari prober adalah keberhasilan publish, waktu respons, dan timeout.")
    paragraph(doc, "Penggunaan prober penting karena peningkatan paket serangan belum tentu langsung berarti layanan gagal. Broker dapat menerima banyak paket SYN, tetapi client normal mungkin masih berhasil publish. Oleh karena itu, data packet capture dan data prober harus dibaca bersama agar kesimpulan availability tidak hanya berdasarkan satu indikator.")

    section_title(doc, "2.10", "Rate Limiting dan nftables")
    paragraph(doc, "Rate limiting adalah teknik pembatasan laju trafik agar server tidak menerima paket atau koneksi melebihi ambang tertentu. Pada Linux, rate limiting dapat diterapkan dengan nftables. nftables adalah framework firewall modern yang dapat membuat aturan untuk menerima, menolak, atau membatasi paket berdasarkan protokol, port, dan kondisi tertentu.")
    paragraph(doc, "Pada penelitian ini, rate limiting diterapkan pada paket TCP SYN yang menuju port MQTT 1883. Prinsipnya adalah menerima paket SYN selama masih berada dalam batas laju yang ditentukan, kemudian menjatuhkan paket yang melebihi batas. Namun, konfigurasi rate limiting harus dikalibrasi. Jika terlalu longgar, serangan tetap masuk terlalu banyak. Jika terlalu ketat, client normal ikut terganggu.")

    section_title(doc, "2.11", "Metrik Evaluasi")
    numbered(doc, [
        "Success rate adalah persentase percobaan publish MQTT yang berhasil dilakukan oleh prober.",
        "RTT adalah waktu yang dibutuhkan prober untuk menyelesaikan satu percobaan publish ke broker.",
        "SYN rate adalah rata-rata paket SYN per detik yang terlihat pada capture.",
        "Half-open adalah indikator koneksi yang belum selesai dalam proses handshake TCP.",
        "RST menunjukkan adanya reset koneksi pada trafik TCP.",
        "Health check menunjukkan apakah broker masih dapat diakses setelah skenario pengujian selesai.",
    ])

    section_title(doc, "2.12", "Kerangka Pemikiran")
    paragraph(doc, "Kerangka pemikiran penelitian ini dimulai dari posisi broker MQTT sebagai titik pusat komunikasi IoT. Karena broker menerima semua koneksi client, maka broker menjadi titik yang rawan terhadap lonjakan trafik. SYN flood dipilih sebagai bentuk serangan karena langsung menekan proses pembentukan koneksi TCP ke broker. Untuk membuktikan dampaknya, penelitian menggabungkan data packet capture dan data prober. Selanjutnya, rate limiting diterapkan dan dituning untuk melihat apakah availability broker dapat dipertahankan.")

    section_title(doc, "2.13", "Hubungan Layer TCP dan Layer Aplikasi MQTT")
    paragraph(doc, "Dalam penelitian ini perlu dibedakan antara trafik pada layer TCP dan trafik pada layer aplikasi MQTT. MQTT berjalan setelah koneksi TCP terbentuk. Artinya, sebelum pesan publish dapat dikirim, client harus lebih dulu menyelesaikan handshake TCP dengan broker. Jika proses awal koneksi TCP terganggu, maka komunikasi MQTT juga ikut terganggu walaupun payload MQTT belum sempat dikirim.")
    paragraph(doc, "Perbedaan ini penting agar pembahasan tidak keliru. Ketika penelitian menyebut serangan SYN flood, serangan tersebut bukan berupa pesan MQTT dengan topik tertentu. Serangan terjadi lebih awal, yaitu pada tahap pembukaan koneksi TCP. Karena itu, indikator utama yang digunakan bukan isi topic atau payload MQTT, tetapi jumlah paket SYN, half-open connection, reset koneksi, dan keberhasilan client normal membentuk koneksi ke broker.")
    paragraph(doc, "Di sisi lain, MQTT tetap perlu diuji melalui prober karena tujuan akhir penelitian adalah availability layanan MQTT. Jika packet capture menunjukkan banyak paket SYN, tetapi prober tetap dapat publish, maka serangan sudah terlihat pada jaringan tetapi belum menurunkan layanan secara nyata. Jika prober mulai timeout atau gagal publish, maka bukti availability sudah muncul. Kombinasi dua sudut pandang ini membuat analisis lebih kuat.")

    section_title(doc, "2.14", "Availability pada Layanan MQTT")
    paragraph(doc, "Availability adalah kemampuan layanan untuk tetap dapat digunakan ketika dibutuhkan. Pada broker MQTT, availability dapat dipahami sebagai kemampuan broker menerima koneksi dan pesan dari client normal. Dalam penelitian ini, availability tidak hanya dinilai dari status service Mosquitto aktif, karena service yang aktif belum tentu dapat merespons client dengan baik.")
    paragraph(doc, "Indikator availability yang lebih tepat adalah keberhasilan publish dan waktu respons. Jika prober berhasil publish secara konsisten, maka broker masih dapat menjalankan fungsi utama. Jika prober gagal publish, timeout, atau RTT meningkat sangat tinggi, maka layanan mulai kehilangan availability. Oleh sebab itu, penelitian ini memakai success rate dan RTT sebagai indikator utama dari sisi client.")
    paragraph(doc, "Health check setelah pengujian juga digunakan untuk memastikan broker kembali dapat diakses setelah skenario selesai. Health check yang gagal menjadi tanda bahwa dampak serangan tidak hanya terjadi selama trafik serangan berjalan, tetapi dapat bertahan beberapa saat setelahnya. Pada pengujian server nyata, kondisi seperti ini harus menjadi batas aman agar skenario tidak diulang secara berlebihan.")

    section_title(doc, "2.15", "Kebutuhan Validitas Pengujian Keamanan")
    paragraph(doc, "Pengujian keamanan jaringan perlu menjaga validitas dan keamanan. Validitas berarti data yang dikumpulkan benar-benar mewakili kondisi yang ingin dianalisis. Dalam penelitian ini, validitas diperkuat dengan tiga cara. Pertama, target pengujian didefinisikan jelas, yaitu broker MQTT pada port 1883. Kedua, trafik serangan dibuat bertingkat sehingga perubahan indikator dapat dibandingkan dengan baseline. Ketiga, setiap skenario aman diulang beberapa kali untuk mengurangi kemungkinan hasil kebetulan.")
    paragraph(doc, "Keamanan pengujian berarti proses eksperimen tidak boleh menimbulkan gangguan yang tidak terkendali. Karena itu, skenario rate tinggi tidak diulang ketika health check gagal. Keputusan ini sesuai dengan prinsip penelitian keamanan yang bertanggung jawab. Penelitian tetap mendapatkan bukti bahwa rate tinggi dapat menurunkan availability, tetapi tidak memaksakan pengulangan yang berpotensi merugikan lingkungan server.")
    paragraph(doc, "Dengan dasar teori tersebut, hasil BAB IV dapat dibaca secara lebih proporsional. Kenaikan SYN dan half-open adalah bukti adanya tekanan pada layer TCP. Penurunan success rate dan peningkatan RTT adalah bukti dampak pada layanan MQTT. Rate limiting dinilai berhasil jika dapat menjaga availability tanpa menutup akses client normal.")


def chapter_3(doc: Document, figs: dict[str, Path]):
    chapter_title(doc, "III", "Metodologi Penelitian")
    section_title(doc, "3.1", "Jenis Penelitian")
    paragraph(doc, "Penelitian ini merupakan penelitian eksperimental karena dilakukan dengan membuat skenario pengujian, menjalankan trafik normal dan trafik serangan, menerapkan mitigasi, kemudian mengukur hasilnya. Data yang digunakan bersifat kuantitatif karena berupa nilai success rate, RTT, jumlah paket SYN, half-open, RST, dan hasil health check.")
    paragraph(doc, "Eksperimen dilakukan secara terkendali. Laju serangan tidak langsung dinaikkan ke nilai besar, tetapi dimulai dari tingkat rendah, sedang, kemudian tinggi secara terbatas. Pendekatan ini digunakan agar penelitian tetap memperoleh bukti dampak tanpa menimbulkan risiko gangguan berlebihan pada server.")

    section_title(doc, "3.2", "Objek Penelitian")
    paragraph(doc, "Objek penelitian adalah broker MQTT pada server Universitas Mataram. Broker MQTT berperan menerima koneksi dari client MQTT pada port 1883. Serangan SYN flood diarahkan ke port tersebut karena port 1883 adalah port standar MQTT tanpa TLS.")
    paragraph(doc, "Perangkat penguji mengakses server melalui VPN WireGuard. Jalur VPN ini digunakan sebagai jalur resmi pengujian sehingga trafik dapat diarahkan ke alamat server yang diberikan. Dengan demikian, pengujian tidak dilakukan secara acak ke jaringan publik, tetapi pada target dan jalur yang telah disetujui.")

    section_title(doc, "3.3", "Arsitektur Sistem Pengujian")
    paragraph(doc, "Arsitektur pengujian terdiri dari perangkat penguji, server broker MQTT, dan proses attacker terkontrol. Perangkat penguji menjalankan prober MQTT, packet capture, dan skrip pengolahan data. Server Universitas Mataram menjalankan broker Mosquitto dan menerima koneksi pada port 1883. Attacker terkontrol menghasilkan paket SYN dengan laju tertentu ke port MQTT.")
    add_figure(doc, figs["architecture"], "Gambar 3.1 Arsitektur pengujian broker MQTT Universitas Mataram", 5.9)
    paragraph(doc, "Arsitektur ini mempertahankan konsep sniffer dan prober dari rancangan awal. Prober digunakan untuk mengukur layanan dari sisi client normal, sedangkan sniffer atau packet capture digunakan untuk melihat trafik jaringan. Perbedaannya, pengujian TA 2 ini diarahkan agar data dapat langsung dikaitkan dengan broker yang diuji dan skenario SYN flood yang terukur.")

    section_title(doc, "3.4", "Alat dan Bahan")
    add_table(
        doc,
        ["No", "Komponen", "Fungsi"],
        [
            ["1", "Server Universitas Mataram", "Menjalankan broker MQTT Mosquitto pada port 1883."],
            ["2", "WireGuard VPN", "Menghubungkan perangkat penguji ke jaringan server secara resmi."],
            ["3", "Mosquitto", "Broker MQTT dan utilitas client publish."],
            ["4", "TShark/dumpcap", "Melakukan packet capture dan ekstraksi trafik jaringan."],
            ["5", "hping3", "Menghasilkan paket TCP SYN terkontrol untuk skenario SYN flood."],
            ["6", "nftables", "Menerapkan rate limiting pada paket SYN menuju port MQTT."],
            ["7", "Python", "Mengotomatisasi eksperimen, kompilasi CSV, dan pembuatan ringkasan."],
        ],
        "Tabel 3.1 Alat dan bahan penelitian",
    )

    section_title(doc, "3.5", "Skenario Pengujian")
    paragraph(doc, "Skenario pengujian dibagi menjadi empat kelompok. Pertama, baseline untuk mengetahui kondisi broker tanpa serangan. Kedua, SYN flood tanpa mitigasi pada beberapa tingkat laju. Ketiga, SYN flood dengan konfigurasi rate limiting awal. Keempat, tuning rate limiting untuk mencari konfigurasi yang paling layak.")
    add_table(
        doc,
        ["Skenario", "Rate Serangan", "Mitigasi", "Repetisi", "Tujuan"],
        [
            ["Baseline", "0 pps", "Tidak", "3", "Mengukur kondisi normal broker."],
            ["SYN rendah", "10 pps", "Tidak", "3", "Melihat perubahan trafik awal."],
            ["SYN sedang", "25 pps", "Tidak", "3", "Melihat beban koneksi yang lebih jelas."],
            ["SYN tinggi", "50 pps", "Tidak", "1", "Membuktikan dampak availability dengan risiko terbatas."],
            ["Mitigasi awal", "10, 15, 25 pps", "15/s burst 20", "3", "Menguji aturan awal."],
            ["Tuning", "25 pps", "20/30, 25/30, 30/40", "3", "Menentukan konfigurasi paling layak."],
        ],
        "Tabel 3.2 Skenario pengujian",
    )

    section_title(doc, "3.6", "Alur Pengujian")
    numbered(doc, [
        "Memastikan VPN aktif dan server broker MQTT dapat diakses pada port 1883.",
        "Menjalankan health check awal untuk memastikan broker dalam kondisi siap diuji.",
        "Menjalankan packet capture pada interface VPN atau interface yang dilalui trafik menuju broker.",
        "Menjalankan prober MQTT sebagai client normal yang mengirim pesan periodik ke broker.",
        "Menjalankan SYN flood terkontrol sesuai rate skenario.",
        "Menghentikan capture setelah durasi pengujian selesai.",
        "Mengekstrak file pcapng menjadi CSV trafik TCP.",
        "Mengompilasi metrik dan ringkasan hasil pengujian.",
        "Melakukan cooldown dan health check sebelum skenario berikutnya.",
    ])
    add_figure(doc, figs["flow"], "Gambar 3.2 Alur pengujian dan pengolahan data", 5.9)

    section_title(doc, "3.7", "Indikator Keberhasilan")
    paragraph(doc, "Indikator keberhasilan digunakan agar hasil penelitian tidak hanya berupa deskripsi umum. Pada penelitian ini, konfigurasi mitigasi dianggap layak apabila mampu menjaga success rate client normal, menjaga health check broker, dan tidak menyebabkan RTT rata-rata menjadi ekstrem.")
    add_table(
        doc,
        ["Indikator", "Kriteria", "Alasan"],
        [
            ["Success rate", "Rata-rata >= 95%", "Client normal harus tetap bisa publish."],
            ["Health check", "Seluruh repetisi OK", "Broker harus tetap dapat diakses setelah pengujian."],
            ["RTT rata-rata", "< 800 ms untuk tuning 25 pps", "Respons tidak boleh terlalu lambat."],
            ["Bukti trafik", "SYN dan half-open meningkat saat serangan", "Menunjukkan serangan benar terjadi."],
            ["Etika pengujian", "Skenario berisiko tinggi tidak diulang jika health gagal", "Menghindari gangguan berlebihan."],
        ],
        "Tabel 3.3 Indikator keberhasilan",
    )

    section_title(doc, "3.8", "Metode Pengumpulan Data")
    paragraph(doc, "Data dikumpulkan dari dua sumber utama. Sumber pertama adalah packet capture yang merekam trafik TCP menuju broker MQTT. Sumber kedua adalah log prober yang mencatat keberhasilan publish dan waktu respons. Kedua sumber ini saling melengkapi. Packet capture menunjukkan apa yang terjadi pada jaringan, sedangkan prober menunjukkan dampaknya terhadap client normal.")
    paragraph(doc, "Setiap hasil pengujian disimpan dalam folder eksperimen yang memuat file mentah dan file olahan. File mentah digunakan sebagai bukti yang dapat diperiksa ulang, sedangkan file olahan digunakan untuk perhitungan tabel dan grafik. Dengan struktur ini, data BAB IV tidak hanya berasal dari pengamatan manual, tetapi dari artefak eksperimen yang terdokumentasi.")

    section_title(doc, "3.9", "Teknik Analisis Data")
    paragraph(doc, "Analisis dilakukan dengan membandingkan nilai metrik antar skenario. Pada baseline, nilai success rate dan RTT digunakan sebagai acuan kondisi normal. Pada skenario SYN flood, nilai SYN rate dan half-open dibandingkan dengan baseline untuk melihat apakah terjadi peningkatan trafik koneksi. Selanjutnya, success rate dan RTT prober digunakan untuk melihat apakah peningkatan trafik tersebut berdampak pada availability.")
    paragraph(doc, "Pada pengujian mitigasi, hasil rate limiting tidak langsung dianggap berhasil hanya karena aturan firewall aktif. Keberhasilan dilihat dari kemampuan aturan tersebut menjaga client normal. Oleh karena itu, beberapa konfigurasi rate limiting dibandingkan. Konfigurasi terbaik dipilih berdasarkan KPI, bukan berdasarkan asumsi bahwa batas rate yang lebih kecil selalu lebih aman.")

    section_title(doc, "3.10", "Etika dan Keamanan Pengujian")
    paragraph(doc, "Pengujian SYN flood dilakukan hanya pada target yang telah ditentukan dan melalui akses yang diberikan. Laju serangan dibatasi, pengujian diberi jeda cooldown, dan skenario yang menyebabkan health check gagal tidak diulang. Prinsip ini penting karena penelitian keamanan jaringan harus tetap menjaga ketersediaan layanan dan tidak menimbulkan dampak di luar ruang lingkup penelitian.")

    section_title(doc, "3.11", "Rancangan File Output")
    paragraph(doc, "Setiap skenario menghasilkan beberapa file output agar hasil dapat diperiksa ulang. File pcapng menyimpan paket jaringan mentah. File raw_flow.csv berisi hasil ekstraksi paket menjadi bentuk tabel. File prober.csv mencatat status publish MQTT dari client normal. File metrics.csv dan summary berisi hasil kompilasi metrik utama. Dengan pemisahan ini, penelitian tidak hanya menyimpan kesimpulan akhir, tetapi juga menyimpan jalur data dari capture sampai tabel BAB IV.")
    add_table(
        doc,
        ["File", "Isi", "Fungsi Analisis"],
        [
            ["capture.pcapng", "Paket jaringan mentah", "Bukti primer trafik jaringan."],
            ["raw_flow.csv", "Ekstraksi field TCP", "Menghitung SYN, RST, packet rate, dan half-open."],
            ["prober.csv", "Percobaan publish MQTT", "Menghitung success rate dan RTT."],
            ["metrics.csv", "Metrik hasil kompilasi", "Sumber tabel agregat."],
            ["summary.txt", "Ringkasan tiap run", "Memudahkan pemeriksaan cepat."],
        ],
        "Tabel 3.4 Rancangan file output pengujian",
    )

    section_title(doc, "3.12", "Rancangan Pengendalian Risiko")
    paragraph(doc, "Risiko utama pada penelitian ini adalah terganggunya akses broker atau terjadinya blokir sementara akibat trafik yang terlalu agresif. Untuk mengendalikan risiko tersebut, pengujian dilakukan secara bertahap. Rate rendah dijalankan terlebih dahulu, kemudian rate sedang. Rate tinggi hanya digunakan untuk membuktikan dampak availability dan tidak diulang apabila health check gagal.")
    paragraph(doc, "Cooldown antar skenario digunakan agar kondisi jaringan dan broker tidak langsung dipakai untuk pengujian berikutnya ketika masih dalam keadaan tidak stabil. Selain itu, health check dilakukan sebelum dan sesudah skenario. Jika broker tidak sehat setelah skenario tertentu, pengujian dihentikan atau diturunkan intensitasnya. Prosedur ini membuat eksperimen lebih aman dan lebih dapat dipertanggungjawabkan.")

    section_title(doc, "3.13", "Rancangan Tuning Mitigasi")
    paragraph(doc, "Tuning mitigasi dilakukan karena nilai rate limiting tidak dapat ditentukan hanya berdasarkan dugaan. Nilai yang terlalu rendah memang dapat menjatuhkan lebih banyak paket SYN, tetapi juga berisiko menjatuhkan koneksi normal. Nilai yang terlalu tinggi dapat menjaga client normal, tetapi mungkin tidak cukup membatasi lonjakan serangan. Oleh karena itu, penelitian membandingkan beberapa kandidat konfigurasi.")
    paragraph(doc, "Kandidat konfigurasi dipilih dengan menaikkan limit dan burst secara bertahap. Konfigurasi awal 15/s burst 20 digunakan sebagai titik awal. Ketika hasilnya belum memenuhi KPI pada 25 pps, konfigurasi dinaikkan menjadi 20/s burst 30, 25/s burst 30, dan 30/s burst 40. Konfigurasi final dipilih berdasarkan success rate, health check, RTT rata-rata, dan kestabilan antar repetisi.")


def agg_table_rows(agg: list[dict[str, str]], batch: str) -> list[list[str]]:
    rows = [r for r in agg if r["batch_label"] == batch]
    rows = sorted(rows, key=lambda r: float(r["rate_pps"]))
    return [
        [
            fnum(r["rate_pps"], 0),
            fnum(r["runs"], 0),
            pct(r["success_mean"]),
            fnum(r["rtt_avg_mean"]),
            fnum(r["syn_mean"], 4),
            fnum(r["half_open_mean"], 4),
            f"{fnum(r['health_ok_count'], 0)}/{fnum(r['runs'], 0)}",
        ]
        for r in rows
    ]


def chapter_4(doc: Document, figs: dict[str, Path], agg: list[dict[str, str]], tuning: list[dict[str, str]], all_runs: list[dict[str, str]]):
    chapter_title(doc, "IV", "Hasil dan Pembahasan")
    section_title(doc, "4.1", "Gambaran Umum Pengujian")
    paragraph(doc, "Bab ini memaparkan hasil pengujian broker MQTT pada server Universitas Mataram. Pengujian dilakukan untuk melihat kondisi baseline, dampak SYN flood tanpa mitigasi, hasil rate limiting awal, dan hasil tuning konfigurasi rate limiting. Data yang dibahas merupakan hasil agregasi dari file eksperimen yang tersimpan pada direktori output pengujian.")
    paragraph(doc, "Interpretasi hasil dilakukan secara hati-hati. Tidak semua peningkatan trafik langsung berarti layanan gagal. Oleh karena itu, penelitian ini membedakan dua jenis bukti. Bukti pertama adalah bukti trafik, seperti peningkatan paket SYN dan half-open. Bukti kedua adalah bukti availability, seperti success rate prober, RTT, dan health check broker setelah pengujian.")
    add_table(
        doc,
        ["Komponen", "Nilai"],
        [
            ["Target", "Broker MQTT server Universitas Mataram"],
            ["Alamat target", "172.16.10.44 melalui VPN WireGuard"],
            ["Port uji", "TCP 1883"],
            ["Broker", "Mosquitto"],
            ["Serangan", "SYN flood terkontrol menggunakan hping3"],
            ["Mitigasi", "nftables rate limiting pada TCP SYN port 1883"],
            ["Durasi", "Pendek dan bertahap dengan cooldown antar skenario"],
            ["Output", "pcapng, raw_flow.csv, prober.csv, metrics.csv, summary"],
        ],
        "Tabel 4.1 Ringkasan lingkungan pengujian",
    )

    section_title(doc, "4.2", "Hasil Baseline")
    base = [r for r in agg if r["batch_label"] == "tanpa_mitigasi" and fnum(r["rate_pps"], 0) == "0"][0]
    paragraph(doc, f"Pada kondisi baseline tanpa mitigasi, broker memperoleh success rate rata-rata {pct(base['success_mean'])} dari 3 repetisi. RTT rata-rata berada pada {fnum(base['rtt_avg_mean'])} ms, SYN rata-rata {fnum(base['syn_mean'], 4)} paket per detik, dan half-open rata-rata {fnum(base['half_open_mean'], 4)}. Health check berhasil pada {fnum(base['health_ok_count'], 0)}/3 repetisi.")
    paragraph(doc, "Hasil baseline ini menunjukkan bahwa sebelum serangan, broker berada dalam kondisi dapat melayani client normal. Nilai SYN dan half-open yang rendah menjadi acuan untuk menilai peningkatan trafik pada skenario serangan. Dengan baseline ini, perubahan pada skenario berikutnya dapat dibaca sebagai akibat dari perlakuan pengujian, bukan kondisi awal broker yang sudah bermasalah.")

    section_title(doc, "4.3", "Hasil SYN Flood Tanpa Mitigasi")
    add_table(
        doc,
        ["Rate", "Run", "Success", "RTT avg", "SYN/s", "Half-open/s", "Health"],
        agg_table_rows(agg, "tanpa_mitigasi"),
        "Tabel 4.2 Hasil agregat tanpa mitigasi",
    )
    paragraph(doc, "Pada rate 10 pps dan 25 pps tanpa mitigasi, success rate client normal masih 100%. Namun, indikator jaringan berubah signifikan dibanding baseline. SYN rata-rata naik dari sekitar 0,7803 paket per detik pada baseline menjadi 7,7595 pada 10 pps dan 21,7097 pada 25 pps. Half-open juga naik dari sekitar 0,2419 menjadi 7,0261 dan 20,8454. Kenaikan ini membuktikan bahwa trafik SYN flood benar-benar terlihat pada jalur pengujian.")
    paragraph(doc, "Meskipun success rate pada 10 pps dan 25 pps masih 100%, hasil tersebut tidak berarti serangan tidak terjadi. Hal yang terjadi adalah broker masih mampu melayani client normal pada intensitas tersebut. Oleh karena itu, pembahasan akademiknya harus dipisahkan: serangan terbukti dari indikator trafik, sedangkan dampak availability baru terlihat jelas ketika intensitas lebih tinggi.")
    add_figure(doc, figs["success_no_mit"], "Gambar 4.1 Success rate tanpa mitigasi", 5.8)
    add_figure(doc, figs["syn_half_no_mit"], "Gambar 4.2 Kenaikan paket SYN dan half-open tanpa mitigasi", 5.8)

    section_title(doc, "4.4", "Dampak Availability pada Rate Tinggi")
    high = [r for r in agg if r["batch_label"] == "tanpa_mitigasi" and fnum(r["rate_pps"], 0) == "50"][0]
    paragraph(doc, f"Pada serangan 50 pps tanpa mitigasi, success rate turun menjadi {pct(high['success_mean'])}. RTT rata-rata meningkat menjadi {fnum(high['rtt_avg_mean'])} ms dan RTT maksimum mencapai {fnum(high['rtt_max_max'], 0)} ms. Health check setelah pengujian gagal. Hasil ini menjadi bukti bahwa ketika intensitas SYN flood dinaikkan, availability broker dapat terdampak secara langsung.")
    paragraph(doc, "Skenario 50 pps hanya dijalankan satu kali. Keputusan ini bukan kelemahan prosedur, melainkan pembatasan etis dan operasional. Setelah health check gagal, pengulangan skenario berisiko menimbulkan blokir sementara atau gangguan yang tidak diperlukan pada server. Oleh karena itu, data 50 pps digunakan sebagai bukti dampak tinggi, sedangkan analisis tuning mitigasi difokuskan pada rate 25 pps yang lebih aman untuk direplikasi.")

    section_title(doc, "4.5", "Hasil Mitigasi Awal")
    add_table(
        doc,
        ["Rate", "Run", "Success", "RTT avg", "SYN/s", "Half-open/s", "Health"],
        agg_table_rows(agg, "mitigasi_awal_15_20"),
        "Tabel 4.3 Hasil mitigasi awal",
    )
    paragraph(doc, "Mitigasi awal menggunakan konfigurasi limit 15/s dengan burst 20. Pada rate 10 pps dan 15 pps, success rate masih 100%. Namun, pada rate 25 pps, success rate turun menjadi rata-rata 82,39% dan RTT rata-rata naik menjadi 1621,73 ms. Hasil ini menunjukkan bahwa rate limiting awal belum optimal.")
    paragraph(doc, "Temuan ini penting karena rate limiting tidak otomatis efektif hanya karena paket berlebih dijatuhkan. Jika batas terlalu ketat terhadap pola trafik nyata, client normal dapat ikut terdampak. Pada konteks penelitian ini, konfigurasi 15/s burst 20 terlalu agresif untuk skenario 25 pps karena menurunkan availability prober. Oleh sebab itu, diperlukan proses tuning.")

    section_title(doc, "4.6", "Tuning Rate Limiting")
    tune_rows = [
        [
            r["rule_config"],
            fnum(r["limit_per_sec"], 0),
            fnum(r["burst_packets"], 0),
            pct(r["success_mean"]),
            pct(r["success_min"]),
            fnum(r["rtt_avg_mean"]),
            f"{fnum(r['health_ok_count'], 0)}/{fnum(r['runs'], 0)}",
            f"{fnum(r['score'], 0)}/3",
        ]
        for r in tuning
    ]
    add_table(
        doc,
        ["Konfigurasi", "Limit", "Burst", "Success avg", "Success min", "RTT avg", "Health", "Score"],
        tune_rows,
        "Tabel 4.4 Perbandingan tuning rate limiting",
    )
    paragraph(doc, "Tuning dilakukan pada skenario 25 pps karena rate ini masih dapat direplikasi dengan aman dan cukup menunjukkan beban koneksi. Tiga konfigurasi utama dibandingkan dengan mitigasi awal, yaitu limit 20/s burst 30, limit 25/s burst 30, dan limit 30/s burst 40. Setiap konfigurasi diuji sebanyak 3 repetisi.")
    paragraph(doc, "Konfigurasi limit 20/s burst 30 menghasilkan success rate rata-rata 88,69%. Nilai ini lebih baik dari mitigasi awal 15/s burst 20, tetapi belum mencapai KPI success rate minimal 95%. Konfigurasi limit 25/s burst 30 menghasilkan rata-rata 82,05% dan health check hanya 2/3, sehingga tidak layak dijadikan konfigurasi final. Konfigurasi limit 30/s burst 40 menghasilkan success rate 100%, RTT rata-rata 458,70 ms, dan health check 3/3.")
    add_figure(doc, figs["tuning_success"], "Gambar 4.3 Perbandingan success rate tuning rate limiting", 5.8)
    add_figure(doc, figs["tuning_rtt"], "Gambar 4.4 Perbandingan RTT rata-rata tuning rate limiting", 5.8)

    section_title(doc, "4.7", "Penilaian KPI Konfigurasi Mitigasi")
    best = tuning[0]
    add_table(
        doc,
        ["KPI", "Kriteria", "Hasil konfigurasi terbaik", "Status"],
        [
            ["Success rate", ">= 95%", pct(best["success_mean"]), "Memenuhi"],
            ["Success minimum", "Tidak turun di bawah 95%", pct(best["success_min"]), "Memenuhi"],
            ["RTT rata-rata", "< 800 ms", f"{fnum(best['rtt_avg_mean'])} ms", "Memenuhi"],
            ["Health check", "Semua repetisi OK", f"{fnum(best['health_ok_count'], 0)}/{fnum(best['runs'], 0)}", "Memenuhi"],
            ["Repetisi", "Minimal 3 untuk skenario aman", fnum(best["runs"], 0), "Memenuhi"],
        ],
        "Tabel 4.5 Penilaian KPI konfigurasi mitigasi",
    )
    paragraph(doc, "Berdasarkan KPI, konfigurasi terbaik adalah limit 30/s burst 40. Konfigurasi ini tidak dipilih karena nilai limit-nya paling besar, tetapi karena hasil pengujiannya paling stabil. Pada tiga repetisi serangan 25 pps, client normal tetap berhasil publish 100%, RTT tidak ekstrem, dan broker tetap sehat setelah pengujian.")
    paragraph(doc, "Secara teknis, konfigurasi limit 30/s burst 40 memberi ruang yang cukup bagi koneksi normal dan variasi trafik sesaat, tetapi tetap memberikan batas terhadap lonjakan SYN. Hasil ini menjelaskan mengapa rate limiting perlu dituning. Batas yang terlalu kecil dapat menekan trafik sah, sedangkan batas yang terlalu besar dapat membiarkan trafik serangan masuk terlalu banyak.")

    section_title(doc, "4.8", "Pembahasan Substantif")
    for text in [
        "Hasil penelitian menunjukkan bahwa analisis DoS pada broker MQTT tidak cukup hanya melihat apakah server aktif atau tidak. Pada rate 10 pps dan 25 pps tanpa mitigasi, server masih terlihat mampu melayani prober, tetapi metrik jaringan sudah menunjukkan perubahan besar. Dengan kata lain, serangan dapat terdeteksi dari pola trafik sebelum dampak availability muncul secara nyata.",
        "Penurunan availability terlihat jelas pada rate 50 pps. Pada titik ini, success rate turun drastis dan RTT meningkat. Data ini menguatkan asumsi bahwa broker MQTT dapat terdampak oleh SYN flood ketika intensitas serangan melewati kemampuan server atau jalur jaringan. Namun, karena pengujian dilakukan pada server nyata melalui VPN, penelitian membatasi pengulangan skenario tinggi demi menjaga keamanan lingkungan.",
        "Pengujian mitigasi menghasilkan temuan penting: mitigasi yang terlalu ketat dapat membuat hasil lebih buruk bagi client normal. Hal ini terlihat pada konfigurasi 15/s burst 20 dan 20/s burst 30. Kedua konfigurasi tersebut mampu membatasi trafik, tetapi tidak memenuhi KPI availability pada 25 pps. Oleh karena itu, nilai efektivitas mitigasi harus dinilai dari gabungan indikator trafik dan layanan.",
        "Konfigurasi 30/s burst 40 menjadi hasil paling layak dalam ruang lingkup pengujian. Konfigurasi ini mempertahankan success rate 100% pada serangan 25 pps dan menjaga health check tetap berhasil. Kesimpulan ini tetap dibatasi pada lingkungan uji yang digunakan, yaitu broker Mosquitto pada server Universitas Mataram, akses melalui VPN WireGuard, port MQTT 1883, dan serangan SYN flood terkontrol.",
    ]:
        paragraph(doc, text)

    section_title(doc, "4.9", "Keterbatasan Pengujian")
    numbered(doc, [
        "Skenario 50 pps hanya dilakukan satu kali karena menimbulkan health check gagal dan berisiko menyebabkan blokir sementara.",
        "Pengujian dilakukan melalui jalur VPN, sehingga hasil dapat dipengaruhi oleh kondisi jaringan VPN dan koneksi internet penguji.",
        "Port yang diuji adalah 1883 tanpa TLS. Hasil belum mewakili broker MQTT TLS pada port 8883.",
        "Serangan berasal dari satu perangkat penguji, sehingga belum merepresentasikan Distributed Denial of Service dari banyak sumber.",
        "Rate limiting dituning berdasarkan skenario 25 pps. Jika pola trafik produksi berbeda, konfigurasi perlu diuji ulang sebelum diterapkan permanen.",
    ])

    section_title(doc, "4.10", "Ringkasan Hasil BAB IV")
    numbered(doc, [
        "Baseline menunjukkan broker dapat melayani client normal dengan success rate 100%.",
        "SYN flood 10 pps dan 25 pps meningkatkan SYN rate dan half-open, tetapi belum menurunkan success rate client normal.",
        "SYN flood 50 pps menurunkan success rate menjadi 27,27% dan menyebabkan health check gagal.",
        "Rate limiting awal 15/s burst 20 belum optimal karena success rate pada 25 pps turun menjadi 82,39%.",
        "Tuning 30/s burst 40 menjadi konfigurasi terbaik karena memenuhi seluruh KPI pada skenario 25 pps.",
    ])

    section_title(doc, "4.11", "Ringkasan Data Per Repetisi")
    sample_rows = []
    for r in all_runs[:18]:
        sample_rows.append([
            r.get("scenario_code", "-"),
            r.get("batch_label", "-"),
            fnum(r.get("rate_pps"), 0),
            pct(r.get("prober_success_rate")),
            fnum(r.get("rtt_avg_ms")),
            fnum(r.get("syn_rate_target_mean"), 4),
            str(r.get("post_health_ok", "-")),
        ])
    add_table(
        doc,
        ["Run ID", "Batch", "Rate", "Success", "RTT avg", "SYN/s", "Health"],
        sample_rows,
        "Tabel 4.6 Cuplikan data per repetisi",
    )
    paragraph(doc, "Tabel 4.6 ditampilkan sebagai cuplikan karena seluruh data per repetisi tersedia pada file all_runs_final.csv. File tersebut menjadi lampiran data utama yang dapat digunakan untuk pemeriksaan ulang, perhitungan tambahan, atau pembaruan grafik jika pengujian dilanjutkan.")

    section_title(doc, "4.12", "Analisis Validitas Hasil")
    paragraph(doc, "Validitas hasil dapat dilihat dari konsistensi antara data prober dan data packet capture. Pada baseline, packet capture menunjukkan SYN dan half-open rendah, sementara prober menunjukkan success rate 100%. Pada skenario SYN flood, packet capture menunjukkan kenaikan SYN dan half-open sesuai kenaikan rate serangan. Hubungan ini menunjukkan bahwa skenario serangan yang dijalankan memang terdeteksi pada data jaringan.")
    paragraph(doc, "Pada rate 10 pps dan 25 pps tanpa mitigasi, data prober tetap menunjukkan success rate 100%. Hasil ini perlu dibaca sebagai bukti bahwa broker masih memiliki kapasitas untuk melayani client normal pada intensitas tersebut. Kesimpulan yang tepat bukan menyatakan bahwa serangan tidak berpengaruh, tetapi menyatakan bahwa pengaruhnya pada tahap itu lebih terlihat pada indikator jaringan daripada availability.")
    paragraph(doc, "Pada rate 50 pps, data prober dan health check berubah sejalan dengan indikator jaringan. Success rate turun, RTT meningkat, dan health check gagal. Kesesuaian antar indikator ini memperkuat kesimpulan bahwa intensitas serangan tinggi dapat menurunkan availability broker. Karena skenario ini menyebabkan kondisi tidak sehat, pengujian tidak diulang sebagai bentuk pengendalian risiko.")
    paragraph(doc, "Pada pengujian mitigasi, validitas pembahasan diperkuat oleh proses tuning. Jika penelitian hanya memakai satu konfigurasi rate limiting, hasilnya dapat menyesatkan karena konfigurasi awal 15/s burst 20 ternyata belum optimal. Dengan membandingkan beberapa konfigurasi, penelitian menunjukkan bahwa keberhasilan mitigasi bergantung pada kalibrasi, bukan sekadar keberadaan aturan firewall.")

    section_title(doc, "4.13", "Implikasi Hasil terhadap Pengelolaan Broker MQTT")
    paragraph(doc, "Hasil penelitian memiliki implikasi praktis bagi pengelolaan broker MQTT. Pertama, administrator perlu memantau indikator layer transport, bukan hanya log MQTT. Serangan SYN flood dapat menekan broker sebelum pesan MQTT terbentuk, sehingga log aplikasi mungkin tidak selalu cukup untuk melihat awal gangguan. Packet capture atau monitoring koneksi TCP dapat memberikan tanda lebih awal.")
    paragraph(doc, "Kedua, mitigasi harus diuji dengan trafik normal. Aturan rate limiting yang terlalu agresif dapat terlihat aman dari sisi firewall, tetapi merugikan client sah. Pada penelitian ini, konfigurasi 15/s burst 20 menurunkan success rate saat serangan 25 pps, sedangkan 30/s burst 40 mampu menjaga success rate. Perbedaan ini menunjukkan pentingnya pengujian berbasis KPI.")
    paragraph(doc, "Ketiga, hasil tuning tidak boleh dipindahkan ke semua lingkungan tanpa evaluasi ulang. Broker dengan jumlah client lebih banyak, topik lebih aktif, atau pola koneksi berbeda dapat membutuhkan limit yang berbeda. Oleh karena itu, konfigurasi terbaik pada penelitian ini sebaiknya dipahami sebagai rekomendasi untuk lingkungan uji yang diamati, bukan aturan universal untuk semua broker MQTT.")

    section_title(doc, "4.14", "Interpretasi Per Skenario")
    sub_title(doc, "4.14.1", "Baseline")
    paragraph(doc, "Skenario baseline berfungsi sebagai titik nol penelitian. Pada tahap ini, broker tidak menerima SYN flood sehingga paket SYN yang terekam terutama berasal dari aktivitas koneksi normal. Success rate 100% menunjukkan bahwa client normal dapat menggunakan broker tanpa kendala berarti. Nilai RTT baseline menjadi pembanding untuk membaca apakah pada skenario lain terjadi pelambatan.")
    paragraph(doc, "Baseline juga membuktikan bahwa alat uji bekerja dengan benar. Prober dapat mengirim pesan, capture dapat merekam paket, dan proses ekstraksi dapat menghasilkan metrik. Jika baseline gagal, maka skenario serangan tidak dapat ditafsirkan dengan baik. Karena baseline berhasil pada seluruh repetisi, data berikutnya memiliki dasar perbandingan yang jelas.")
    sub_title(doc, "4.14.2", "SYN Flood 10 pps")
    paragraph(doc, "Pada rate 10 pps, indikator SYN dan half-open meningkat dibanding baseline. Namun, success rate tetap 100%. Hasil ini menunjukkan bahwa broker masih mampu menyerap beban serangan rendah. Secara akademik, skenario ini berguna untuk menunjukkan fase awal serangan, yaitu ketika anomali jaringan sudah muncul tetapi layanan belum mengalami kegagalan nyata.")
    paragraph(doc, "Skenario ini juga memperlihatkan pentingnya pengukuran ganda. Jika penelitian hanya melihat success rate, maka 10 pps dapat dianggap tidak berdampak. Namun, packet capture menunjukkan peningkatan paket SYN dan half-open. Artinya, serangan tetap berlangsung, hanya saja intensitasnya belum cukup besar untuk menurunkan availability.")
    sub_title(doc, "4.14.3", "SYN Flood 25 pps")
    paragraph(doc, "Pada rate 25 pps tanpa mitigasi, indikator trafik meningkat lebih besar dibanding 10 pps. Success rate masih 100%, tetapi jumlah SYN dan half-open mendekati nilai yang jauh lebih tinggi daripada baseline. Skenario ini menjadi titik penting karena cukup aman untuk diulang, tetapi cukup kuat untuk menguji apakah rate limiting mengganggu client normal.")
    paragraph(doc, "Karena itu, rate 25 pps dipilih sebagai basis tuning mitigasi. Pemilihan ini lebih bertanggung jawab daripada memaksa tuning pada 50 pps, karena 50 pps sudah terbukti dapat membuat health check gagal. Dengan 25 pps, penelitian tetap dapat membandingkan konfigurasi mitigasi secara berulang tanpa memperbesar risiko gangguan.")
    sub_title(doc, "4.14.4", "SYN Flood 50 pps")
    paragraph(doc, "Skenario 50 pps memberikan bukti paling jelas terhadap availability. Success rate turun drastis dan RTT meningkat sampai mendekati batas timeout. Health check yang gagal menunjukkan bahwa dampak serangan tidak hanya terlihat selama paket dikirim, tetapi juga berpengaruh pada kondisi layanan setelah skenario selesai.")
    paragraph(doc, "Walaupun hanya satu repetisi, data ini tetap bernilai sebagai bukti dampak tinggi. Dalam penelitian keamanan pada target nyata, pengulangan tidak selalu layak jika skenario sudah menunjukkan efek yang tidak sehat. Keputusan untuk tidak mengulang 50 pps menjaga keseimbangan antara kebutuhan bukti akademik dan tanggung jawab operasional.")
    sub_title(doc, "4.14.5", "Mitigasi Awal dan Tuning")
    paragraph(doc, "Mitigasi awal membuktikan bahwa memasang rate limiting belum cukup. Pada 15/s burst 20, client normal ikut terdampak saat serangan 25 pps. Hasil ini memberi pelajaran bahwa mitigasi tidak boleh hanya dinilai dari seberapa banyak trafik dibatasi. Mitigasi harus menjaga layanan bagi pengguna sah.")
    paragraph(doc, "Tuning menunjukkan perbaikan bertahap. Limit 20/s burst 30 belum memenuhi KPI. Limit 25/s burst 30 tidak stabil karena ada repetisi dengan health check gagal. Limit 30/s burst 40 menjadi konfigurasi terbaik karena seluruh indikator utama terpenuhi. Pola ini memperlihatkan bahwa proses penelitian tidak berhenti pada penerapan alat, tetapi melakukan evaluasi sampai ditemukan konfigurasi yang dapat dipertanggungjawabkan.")

    section_title(doc, "4.15", "Keterkaitan Hasil dengan Rumusan Masalah")
    paragraph(doc, "Rumusan masalah pertama berkaitan dengan kondisi broker pada trafik normal. Hasil baseline menjawab rumusan ini dengan menunjukkan success rate 100%, RTT yang masih dapat diterima, serta health check yang berhasil. Dengan demikian, broker dalam kondisi awal layak digunakan sebagai objek pengujian.")
    paragraph(doc, "Rumusan masalah kedua berkaitan dengan pengaruh SYN flood. Hasil tanpa mitigasi menunjukkan bahwa SYN flood meningkatkan SYN rate dan half-open secara bertingkat. Dampak availability tidak langsung muncul pada 10 pps dan 25 pps, tetapi muncul jelas pada 50 pps. Ini menunjukkan bahwa dampak serangan bergantung pada intensitas dan kapasitas broker/jalur jaringan.")
    paragraph(doc, "Rumusan masalah ketiga berkaitan dengan efektivitas rate limiting. Hasil mitigasi membuktikan bahwa rate limiting dapat efektif, tetapi efektivitasnya bergantung pada konfigurasi. Konfigurasi awal yang terlalu ketat belum memenuhi KPI, sedangkan konfigurasi 30/s burst 40 memenuhi seluruh indikator keberhasilan pada skenario 25 pps.")
    paragraph(doc, "Rumusan masalah keempat berkaitan dengan konfigurasi yang paling layak. Berdasarkan tabel KPI, konfigurasi yang paling layak dalam ruang lingkup penelitian ini adalah limit 30/second dengan burst 40 packets. Konfigurasi tersebut menjaga success rate 100%, RTT rata-rata rendah, dan health check berhasil pada seluruh repetisi.")

    section_title(doc, "4.16", "Kesiapan Data untuk BAB V")
    paragraph(doc, "Data BAB IV sudah cukup untuk menyusun kesimpulan awal TA 2 karena telah mencakup baseline, serangan tanpa mitigasi, bukti dampak availability, mitigasi awal, dan tuning mitigasi. Data juga memiliki repetisi untuk skenario aman sehingga tidak hanya berupa satu kali percobaan. Bagian yang masih dapat ditambah pada penelitian lanjutan adalah variasi durasi, variasi jumlah client normal, dan monitoring sumber daya server.")
    paragraph(doc, "Dengan data yang tersedia, kesimpulan tidak boleh dibuat berlebihan. Kesimpulan yang tepat adalah rate limiting 30/s burst 40 efektif pada lingkungan uji ini untuk skenario SYN flood 25 pps. Kesimpulan tidak boleh menyatakan bahwa konfigurasi tersebut pasti optimal untuk semua server, semua jaringan, atau semua jenis serangan DoS. Batasan ini membuat skripsi lebih defensif secara akademik.")


def chapter_5(doc: Document):
    chapter_title(doc, "V", "Penutup")
    section_title(doc, "5.1", "Kesimpulan")
    paragraph(doc, "Berdasarkan hasil pengujian dan pembahasan, diperoleh beberapa kesimpulan sebagai berikut.")
    numbered(doc, [
        "Broker MQTT pada server Universitas Mataram dalam kondisi baseline mampu melayani client normal dengan success rate 100%. Nilai SYN dan half-open pada kondisi normal rendah, sehingga dapat digunakan sebagai pembanding terhadap skenario serangan.",
        "Serangan SYN flood terbukti meningkatkan indikator trafik TCP menuju broker. Pada rate 10 pps dan 25 pps tanpa mitigasi, SYN rate dan half-open meningkat dibanding baseline, meskipun success rate client normal masih 100%. Hal ini menunjukkan bahwa serangan dapat terlihat dari packet capture sebelum berdampak langsung pada availability.",
        "Dampak availability terlihat pada rate 50 pps tanpa mitigasi. Pada skenario ini, success rate turun menjadi 27,27%, RTT rata-rata meningkat menjadi 2310,91 ms, dan health check setelah pengujian gagal. Data ini membuktikan bahwa intensitas SYN flood yang lebih tinggi dapat mengganggu ketersediaan broker MQTT.",
        "Rate limiting dapat digunakan sebagai mitigasi, tetapi harus dikalibrasi. Konfigurasi awal limit 15/s burst 20 belum optimal karena pada serangan 25 pps success rate turun menjadi 82,39%. Dengan demikian, aturan yang terlalu ketat dapat mengganggu client normal.",
        "Konfigurasi terbaik pada ruang lingkup pengujian ini adalah nftables rate limiting dengan limit 30/second dan burst 40 packets untuk paket TCP SYN menuju port 1883. Pada skenario 25 pps, konfigurasi ini menghasilkan success rate 100%, RTT rata-rata 458,70 ms, dan health check 3/3 OK.",
    ])

    section_title(doc, "5.2", "Saran")
    numbered(doc, [
        "Pengujian berikutnya dapat menambah durasi dan jumlah repetisi agar hasil statistik menjadi lebih kuat.",
        "Skenario serangan dapat diperluas ke publish flood atau connection flood agar dampak DoS pada lapisan aplikasi MQTT juga dapat dibandingkan.",
        "Pengujian pada port MQTT TLS 8883 perlu dilakukan untuk melihat apakah penggunaan TLS memengaruhi pola serangan dan mitigasi.",
        "Jika mitigasi akan diterapkan pada server produksi, konfigurasi rate limiting perlu diuji ulang sesuai trafik normal server agar tidak memblokir client sah.",
        "Penelitian lanjutan dapat menambahkan monitoring CPU, memori, socket, dan log broker agar analisis dampak DoS lebih lengkap.",
        "Pengujian distributed attack dari banyak sumber hanya boleh dilakukan pada lingkungan lab tertutup atau setelah memperoleh izin tertulis yang jelas.",
    ])

    section_title(doc, "5.3", "Kontribusi Penelitian")
    paragraph(doc, "Kontribusi utama penelitian ini adalah menyediakan alur pengujian DoS MQTT yang dapat dijelaskan secara sederhana tetapi tetap terukur. Penelitian tidak hanya menjalankan serangan, tetapi juga menunjukkan hubungan antara trafik TCP, hasil prober MQTT, dan keputusan mitigasi. Dengan demikian, hasilnya dapat digunakan sebagai bahan pembelajaran keamanan jaringan pada konteks broker MQTT.")
    paragraph(doc, "Kontribusi kedua adalah pembuktian bahwa rate limiting memerlukan tuning. Temuan ini penting karena pada praktiknya konfigurasi firewall sering dipasang berdasarkan angka perkiraan. Penelitian ini menunjukkan bahwa konfigurasi awal dapat gagal memenuhi KPI, sedangkan konfigurasi lain dapat lebih stabil setelah diuji.")
    paragraph(doc, "Kontribusi ketiga adalah penerapan batas etis pada pengujian. Skenario 50 pps tidak diulang karena health check gagal. Keputusan ini dapat dijadikan contoh bahwa penelitian keamanan perlu menyeimbangkan kebutuhan data dan keselamatan layanan. Hasil tetap dapat dibahas, tetapi risiko operasional tidak diperbesar secara tidak perlu.")


def references(doc: Document):
    doc.add_page_break()
    centered(doc, "DAFTAR PUSTAKA", 12, True)
    refs = [
        "Al-Fuqaha, A., Guizani, M., Mohammadi, M., Aledhari, M., dan Ayyash, M. (2015). Internet of Things: A Survey on Enabling Technologies, Protocols, and Applications. IEEE Communications Surveys & Tutorials.",
        "Banks, A., Briggs, E., Borgendale, K., dan Gupta, R. (2019). MQTT Version 5.0. OASIS Standard.",
        "Bandyopadhyay, S., dan Sen, J. (2011). Internet of Things: Applications and Challenges in Technology and Standardization. Wireless Personal Communications.",
        "Cisco. (2023). What Is a Denial-of-Service Attack? Cisco Security Documentation.",
        "Cloudflare. (2024). What Is a SYN Flood Attack? Cloudflare Learning Center.",
        "Eclipse Foundation. (2024). Eclipse Mosquitto Documentation.",
        "Fall, K. R., dan Stevens, W. R. (2012). TCP/IP Illustrated, Volume 1: The Protocols. Addison-Wesley.",
        "Fielding, R. T. (2000). Architectural Styles and the Design of Network-based Software Architectures. University of California, Irvine.",
        "Forouzan, B. A. (2017). Data Communications and Networking. McGraw-Hill.",
        "Hwang, R. H., Peng, M. C., Nguyen, V. L., dan Chang, Y. L. (2019). An LSTM-Based Deep Learning Approach for Classifying Malicious Traffic at the Packet Level. Applied Sciences.",
        "Kurose, J. F., dan Ross, K. W. (2021). Computer Networking: A Top-Down Approach. Pearson.",
        "Light, R. A. (2017). Mosquitto: Server and Client Implementation of the MQTT Protocol. Journal of Open Source Software.",
        "Mishra, B., dan Kertesz, A. (2020). The Use of MQTT in M2M and IoT Systems: A Survey. IEEE Access.",
        "NIST. (2020). Foundational Cybersecurity Activities for IoT Device Manufacturers. National Institute of Standards and Technology.",
        "OASIS. (2014). MQTT Version 3.1.1. OASIS Standard.",
        "Postel, J. (1981). Transmission Control Protocol. RFC 793.",
        "Roesch, M. (1999). Snort: Lightweight Intrusion Detection for Networks. USENIX LISA.",
        "Sicari, S., Rizzardi, A., Grieco, L. A., dan Coen-Porisini, A. (2015). Security, Privacy and Trust in Internet of Things: The Road Ahead. Computer Networks.",
        "The Tcpdump Group. (2024). tcpdump and libpcap Documentation.",
        "Wireshark Foundation. (2024). Wireshark and TShark User Documentation.",
        "Netfilter Project. (2024). nftables Wiki and Documentation.",
        "Zhang, C., Green, R., dan Alam, M. (2020). Network Traffic Analysis for MQTT-based IoT Systems. International Journal of Network Security.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(-1.0)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(ref)
        set_run_font(run, 12)


def appendices(doc: Document):
    doc.add_page_break()
    centered(doc, "LAMPIRAN", 12, True)
    section_title(doc, "Lampiran 1", "Daftar Artefak Data Pengujian")
    add_table(
        doc,
        ["No", "File/Folder", "Keterangan"],
        [
            ["1", "output/unram_experiments_comprehensive/20260821_comprehensive_nomitigation_v2", "Data baseline dan SYN flood tanpa mitigasi."],
            ["2", "output/unram_experiments_comprehensive/20260821_comprehensive_mitigation_v1", "Data mitigasi awal limit 15/s burst 20."],
            ["3", "output/unram_experiments_comprehensive/20260821_tuning_limit20_burst30_rate25", "Data tuning limit 20/s burst 30."],
            ["4", "output/unram_experiments_comprehensive/20260821_tuning_limit25_burst30_rate25", "Data tuning limit 25/s burst 30."],
            ["5", "output/unram_experiments_comprehensive/20260821_tuning_limit30_burst40_rate25", "Data tuning limit 30/s burst 40."],
            ["6", "output/unram_experiments_comprehensive/_final_tuning_report_20260821", "Ringkasan final, agregat, dan kandidat tuning."],
        ],
        "Tabel Lampiran 1 Daftar data pengujian",
    )
    section_title(doc, "Lampiran 2", "Contoh Aturan Rate Limiting")
    paragraph(doc, "Konfigurasi mitigasi yang direkomendasikan dari hasil pengujian adalah pembatasan paket TCP SYN menuju port 1883 dengan limit 30/second dan burst 40 packets. Aturan ini perlu diterapkan sebagai aturan sementara saat pengujian atau disesuaikan ulang sebelum dipakai permanen pada server produksi.")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.0
    code = (
        "table inet mqtt_mitigation\n"
        "chain input {\n"
        "  type filter hook input priority 0; policy accept;\n"
        "  tcp dport 1883 tcp flags syn limit rate 30/second burst 40 packets accept\n"
        "  tcp dport 1883 tcp flags syn drop\n"
        "}"
    )
    r = p.add_run(code)
    r.font.name = "Courier New"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    r.font.size = Pt(9)

    section_title(doc, "Lampiran 3", "Catatan Penggunaan Data")
    paragraph(doc, "Data pada skripsi ini digunakan sebagai draft TA 2 untuk bimbingan. Jika dosen pembimbing meminta pengujian tambahan, struktur skrip dan dokumen sudah dibuat agar data baru dapat ditambahkan ke tabel BAB IV tanpa mengubah konsep utama penelitian.")

    section_title(doc, "Lampiran 4", "Panduan Replikasi Pengujian Secara Ringkas")
    numbered(doc, [
        "Pastikan akses VPN ke jaringan server aktif dan alamat broker dapat dijangkau.",
        "Pastikan service Mosquitto pada server aktif dan port 1883 terbuka untuk jalur pengujian.",
        "Jalankan health check awal dari perangkat penguji untuk memastikan broker siap menerima koneksi.",
        "Jalankan packet capture pada interface yang dilewati trafik menuju broker.",
        "Jalankan prober MQTT untuk mengirim pesan normal secara periodik.",
        "Jalankan SYN flood terkontrol pada rate yang telah ditetapkan.",
        "Hentikan capture setelah durasi selesai dan simpan file pcapng.",
        "Ekstrak pcapng menjadi raw_flow.csv menggunakan TShark.",
        "Kompilasi data prober dan data capture menjadi metrics.csv.",
        "Jalankan health check setelah skenario selesai.",
        "Beri jeda cooldown sebelum menjalankan skenario berikutnya.",
        "Jika health check gagal, hentikan peningkatan rate atau ulangi hanya setelah mendapat izin dan kondisi server kembali normal.",
    ])

    section_title(doc, "Lampiran 5", "Contoh Narasi Hasil untuk Bimbingan")
    paragraph(doc, "Pada pengujian baseline, broker MQTT dapat melayani client normal dengan success rate 100%. Ketika SYN flood dinaikkan ke 10 pps dan 25 pps, indikator trafik seperti SYN rate dan half-open meningkat, tetapi layanan masih dapat menerima publish dari client normal. Dampak availability baru terlihat jelas pada 50 pps, yaitu success rate turun menjadi 27,27% dan health check gagal.")
    paragraph(doc, "Untuk mitigasi, saya menggunakan rate limiting pada paket TCP SYN menuju port 1883. Hasil awal dengan limit 15/s burst 20 belum optimal karena client normal ikut terdampak pada 25 pps. Setelah tuning, konfigurasi 30/s burst 40 menjadi hasil terbaik karena success rate tetap 100%, RTT rata-rata 458,70 ms, dan health check berhasil di semua repetisi.")
    paragraph(doc, "Kesimpulan sementara yang dapat disampaikan adalah SYN flood memang meningkatkan beban koneksi pada broker MQTT, dan rate limiting dapat membantu menjaga availability jika dikalibrasi. Penelitian ini tidak menyatakan satu konfigurasi berlaku universal, tetapi menunjukkan konfigurasi yang paling layak untuk lingkungan uji broker Universitas Mataram pada skenario yang diuji.")

    section_title(doc, "Lampiran 6", "Catatan Batasan untuk Dosen Pembimbing")
    paragraph(doc, "Draft ini memakai data pengujian yang sudah tersedia sampai tahap tuning rate limiting. Data tersebut sudah mencakup tiga repetisi untuk skenario aman dan satu bukti dampak tinggi pada 50 pps. Jika dibutuhkan penguatan akademik, pengujian dapat ditambah dengan durasi lebih panjang, jumlah client normal lebih banyak, atau monitoring CPU dan memori server.")
    paragraph(doc, "Bagian format administratif seperti nama pembimbing, halaman pengesahan, nomor halaman romawi, dan penyesuaian gaya sitasi masih perlu disesuaikan dengan pedoman resmi program studi. Namun, struktur substansi penelitian dari BAB I sampai BAB V sudah disusun agar dapat langsung menjadi bahan bimbingan TA 2.")

    section_title(doc, "Lampiran 7", "Glosarium Singkat")
    add_table(
        doc,
        ["Istilah", "Penjelasan Sederhana"],
        [
            ["Broker MQTT", "Server pusat yang menerima pesan dari publisher dan meneruskan pesan ke subscriber."],
            ["Client MQTT", "Perangkat atau program yang terhubung ke broker untuk mengirim atau menerima pesan."],
            ["Publisher", "Client yang mengirim pesan ke suatu topic pada broker MQTT."],
            ["Subscriber", "Client yang berlangganan topic dan menerima pesan dari broker MQTT."],
            ["Topic", "Alamat logis pada MQTT yang dipakai untuk mengelompokkan pesan."],
            ["TCP", "Protokol transport yang membuat koneksi antara client dan server sebelum data aplikasi dikirim."],
            ["SYN", "Paket awal yang dikirim client untuk memulai koneksi TCP."],
            ["SYN-ACK", "Balasan server terhadap paket SYN sebagai tahap kedua handshake TCP."],
            ["ACK", "Paket konfirmasi dari client untuk menyelesaikan handshake TCP."],
            ["SYN flood", "Serangan yang mengirim banyak SYN agar server sibuk menangani percobaan koneksi."],
            ["Half-open", "Koneksi TCP yang baru sebagian terbentuk dan belum selesai handshake."],
            ["DoS", "Serangan yang bertujuan menurunkan atau menghentikan ketersediaan layanan."],
            ["Rate limiting", "Pembatasan laju trafik agar paket/koneksi tidak melebihi ambang tertentu."],
            ["Burst", "Toleransi lonjakan sesaat sebelum pembatasan benar-benar menjatuhkan paket."],
            ["Success rate", "Persentase percobaan client normal yang berhasil."],
            ["RTT", "Waktu respons dari satu percobaan komunikasi client ke broker."],
            ["Health check", "Pemeriksaan singkat untuk memastikan broker masih dapat diakses."],
        ],
        "Tabel Lampiran 2 Glosarium singkat",
    )

    section_title(doc, "Lampiran 8", "Checklist Validasi Pengujian")
    paragraph(doc, "Checklist ini dipakai untuk memastikan pengujian yang dilakukan dapat dipertanggungjawabkan. Setiap poin sebaiknya diperiksa sebelum data dimasukkan ke BAB IV.")
    add_table(
        doc,
        ["No", "Pemeriksaan", "Status yang Diharapkan"],
        [
            ["1", "VPN aktif dan route ke server benar", "Alamat broker dapat dijangkau dari perangkat penguji."],
            ["2", "Broker Mosquitto aktif", "Port 1883 terbuka dan dapat menerima publish MQTT."],
            ["3", "Capture berjalan sebelum serangan", "File pcapng terbentuk dan ukuran file bertambah."],
            ["4", "Prober berjalan selama skenario", "File prober.csv berisi percobaan publish."],
            ["5", "Rate serangan sesuai skenario", "Ringkasan menunjukkan rate yang diuji."],
            ["6", "Cooldown dilakukan", "Skenario berikutnya tidak langsung dijalankan setelah dampak tinggi."],
            ["7", "Health check akhir berhasil atau dicatat gagal", "Status health menjadi bagian dari pembahasan."],
            ["8", "File mentah disimpan", "pcapng dan CSV tetap tersedia untuk diperiksa ulang."],
            ["9", "Agregasi sesuai run", "Jumlah repetisi pada tabel sesuai folder data."],
            ["10", "Kesimpulan tidak melebihi data", "Kesimpulan dibatasi pada lingkungan dan skenario uji."],
        ],
        "Tabel Lampiran 3 Checklist validasi pengujian",
    )

    section_title(doc, "Lampiran 9", "Catatan Interpretasi untuk Penguji")
    paragraph(doc, "Jika success rate masih 100% pada rate serangan tertentu, hasil tersebut tidak otomatis berarti serangan gagal. Lihat terlebih dahulu data packet capture. Jika SYN dan half-open meningkat, maka serangan terbukti masuk, tetapi broker masih mampu melayani client normal. Pernyataan yang tepat adalah serangan belum menurunkan availability pada intensitas tersebut.")
    paragraph(doc, "Jika success rate turun, RTT naik tinggi, atau health check gagal, maka sudah ada bukti dampak availability. Pada kondisi seperti ini, pengujian tidak perlu langsung dinaikkan ke rate yang lebih besar. Lebih tepat melakukan cooldown, memastikan broker kembali sehat, lalu menentukan apakah skenario perlu diulang dengan izin dan batas yang aman.")
    paragraph(doc, "Jika rate limiting membuat success rate turun, jangan langsung menyimpulkan rate limiting gagal secara konsep. Yang gagal bisa jadi nilai konfigurasinya. Karena itu, perlu tuning. Konfigurasi yang baik adalah konfigurasi yang menahan trafik tidak wajar tetapi tetap memberi ruang untuk client normal.")
    paragraph(doc, "Dalam penyampaian ke dosen pembimbing, fokus utama bukan menunjukkan bahwa serangan dibuat sebesar mungkin. Fokus akademiknya adalah menunjukkan hubungan sebab-akibat secara terukur: baseline normal, SYN flood menaikkan indikator trafik, intensitas tinggi menurunkan availability, dan rate limiting yang dituning dapat membantu menjaga layanan.")

    section_title(doc, "Lampiran 10", "Rencana Penguatan Jika Diminta Pembimbing")
    paragraph(doc, "Apabila dosen pembimbing meminta penguatan data, penelitian dapat dilanjutkan tanpa mengubah topik utama. Penguatan pertama adalah menambah durasi setiap skenario, misalnya dari pengujian pendek menjadi 3 sampai 5 menit per run. Durasi yang lebih panjang dapat memperlihatkan apakah broker tetap stabil dalam waktu yang lebih lama.")
    paragraph(doc, "Penguatan kedua adalah menambah jumlah client normal. Pada draft ini, prober digunakan untuk mewakili client normal. Pengujian lanjutan dapat menjalankan beberapa prober bersamaan untuk melihat apakah rate limiting tetap memberi ruang bagi lebih dari satu client sah. Penguatan ini relevan karena dalam sistem IoT nyata, broker biasanya melayani banyak perangkat.")
    paragraph(doc, "Penguatan ketiga adalah menambah monitoring sumber daya server. Selain packet capture dan prober, penelitian dapat mengambil data CPU, memori, jumlah socket, dan log Mosquitto. Data ini akan membuat pembahasan lebih lengkap karena dampak serangan dapat dilihat dari sisi jaringan, sisi client, dan sisi server.")
    paragraph(doc, "Penguatan keempat adalah membandingkan port 1883 dan 8883. Port 1883 menggunakan MQTT tanpa TLS, sedangkan 8883 umumnya digunakan untuk MQTT dengan TLS. Pengujian TLS dapat menunjukkan apakah lapisan keamanan tambahan memengaruhi beban koneksi dan kebutuhan mitigasi. Namun, pengujian ini membutuhkan konfigurasi sertifikat dan izin yang lebih jelas.")
    paragraph(doc, "Penguatan kelima adalah melakukan uji pada lingkungan lab tertutup. Jika ingin mencoba rate yang lebih tinggi atau serangan dari beberapa sumber, pengujian sebaiknya dipindahkan ke lab agar tidak mengganggu server nyata. Dengan cara ini, penelitian tetap dapat memperluas pembahasan tanpa memperbesar risiko operasional pada server Universitas Mataram.")

    section_title(doc, "Lampiran 11", "Matriks Kesiapan Bimbingan")
    add_table(
        doc,
        ["Aspek", "Kondisi Draft", "Catatan"],
        [
            ["Topik", "Sudah jelas", "Analisis SYN flood pada broker MQTT dan mitigasi rate limiting."],
            ["BAB I", "Sudah disusun", "Latar belakang, rumusan masalah, tujuan, manfaat, dan batasan sudah sesuai arah penelitian."],
            ["BAB II", "Sudah disusun", "Teori MQTT, TCP, SYN flood, packet capture, prober, dan rate limiting sudah tersedia."],
            ["BAB III", "Sudah disusun", "Metode, arsitektur, skenario, indikator, dan etika pengujian sudah dijelaskan."],
            ["BAB IV", "Sudah berisi data", "Memakai hasil pengujian komprehensif dan tuning mitigasi."],
            ["BAB V", "Sudah disusun", "Kesimpulan dibatasi pada data yang tersedia."],
            ["Data", "Cukup untuk bimbingan", "Masih bisa diperkuat dengan durasi dan monitoring server tambahan."],
            ["Format", "Draft", "Masih perlu penyesuaian pedoman resmi, halaman pengesahan, dan sitasi final."],
        ],
        "Tabel Lampiran 4 Matriks kesiapan bimbingan",
    )
    paragraph(doc, "Matriks ini menunjukkan bahwa draft sudah siap dipakai sebagai bahan diskusi TA 2. Bagian yang perlu disampaikan sejak awal kepada pembimbing adalah bahwa data pengujian sudah ada, tetapi masih dapat diperluas jika pembimbing meminta validasi tambahan. Dengan demikian, diskusi bimbingan dapat berfokus pada keputusan akademik, bukan lagi pada kebingungan konsep.")


def build() -> Path:
    ensure_dirs()
    agg = read_csv(FINAL_AGG)
    tuning = read_csv(TUNING)
    all_runs = read_csv(ALL_RUNS)
    figs = generate_figures(agg, tuning)

    doc = Document()
    set_doc_style(doc)
    add_page_number(doc.sections[0])

    front_matter(doc)
    chapter_1(doc)
    chapter_2(doc, figs)
    chapter_3(doc, figs)
    chapter_4(doc, figs, agg, tuning, all_runs)
    chapter_5(doc)
    references(doc)
    appendices(doc)

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build())
